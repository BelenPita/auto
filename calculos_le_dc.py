
import math
import numpy as np
from math import pi
from math import sqrt as raiz
from math import log
import cmath
# Coeficientes
mu0 = 4e-7 * pi # H/m
epsilon0 = 8.854e-12  # Permisividad del vacío en F/m


# Valores linea
coef_temp = 0.00403 # 1/°C
frecuencia = 50 # Hz
longitud = 6.5320 # km
resistividad = 250 # Ohmios metro
tension_nominal = 220 # kV
cos_phi = 0.928
potencia_transportada = 165 #MVA
tiempo_accionamiento_proteccion = 0.5 # s
altitud_media = 800

# Caraterísticas conductor
material = "Aluminio-Acero"
composicion = "54+7"  
diametro_alambre_ext = 3.08  # mm
seccion = 454.5 # mm2
diametro = 27.7 # mm
temperatura_invierno = 4 # °C
temperatura_verano = 19 # °C
Tc = 85 # °C
resistencia = 0.0719 # Ohmios


emisividad_conductor = 0.5
coeficiente_absorcion = 0.5
radiacion_invierno = 79 # W/m2
radiacion_verano = 261 # W/m2
velocidad_viento = 0.6 # m/s

# Valores cable tierra
resistencia_tierra = 0.33 # Ohmios
seccion_tierra = 155.5 # mm2
diametro_tierra = 18 # mm
coef_temp_tierra = 14.4e-6 # 1/°C

# Puntos
tres = (4.1,36.63737812)
dos = (4.3,31.137781)
uno = (4.1,25.63771977)
seis = (-4.1,36.63737812)
cinco = (-4.3,31.137781)
cuatro = (-4.1,25.63771977)
tierra = (0,45.51358267)


# Cálculo de la resistencia a 85ºC y en ca
def resistencia_a_temp(resistencia_20C, temp):
    return resistencia_20C * (1 + coef_temp * (temp - 20))
resistencia_85C = resistencia_a_temp(resistencia, 85)
print(f"Resistencia a 85ºC: {resistencia_85C:.6f} Ohmios")

reactancia_pelicular_85C = 10**-7 * 8 * pi * frecuencia * (1 / (resistencia_a_temp(resistencia, 85) / 1000))
print(f"Reactancia por efecto pelicular a 85ºC: {reactancia_pelicular_85C:.6f} Ohmios")
ys = reactancia_pelicular_85C ** 2 / (192 + 0.8*reactancia_pelicular_85C**2)
resistencia_ca = resistencia_85C * (1 + ys)
print(f"Resistencia en CA a 85ºC: {resistencia_ca:.6f} Ohmios")
r_ca_longitud = resistencia_ca * longitud
print(f"Resistencia en CA para {longitud} km: {r_ca_longitud:.6f} Ohmios")


# Cálculo matriz de impedancias
penetracion_terreno = raiz(resistividad/(pi*frecuencia*mu0))
print(f"Penetración terreno: {penetracion_terreno:.6f} m")

# Matriz de distancias
puntos = [uno, dos, tres, cuatro, cinco, seis, tierra]
n_puntos = len(puntos)
matriz_distancias = [[0]*n_puntos for _ in range(n_puntos)]

for i in range(n_puntos):
    for j in range(n_puntos):
        x1, y1 = puntos[i]
        x2, y2 = puntos[j]
        distancia = raiz((x2-x1)**2 + (y2-y1)**2)
        matriz_distancias[i][j] = distancia

print("\nMatriz de distancias (km):")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
etiquetas = ["Punto1", "Punto2", "Punto3", "Punto4", "Punto5", "Punto6", "Punto7"]
for i, fila in enumerate(matriz_distancias):
    print(f"{etiquetas[i]}: ", end="")
    for dist in fila:
        print(f"{dist:9.2f}", end="  ")
    print()

# Matriz D_prima: distancia entre puntos y sus espejos respecto al suelo
# El espejo de un punto (x, y) respecto al suelo es (x, -y)
matriz_D_prima = [[0]*n_puntos for _ in range(n_puntos)]

for i in range(n_puntos):
    for j in range(n_puntos):
        x1, y1 = puntos[i]
        # Espejo del punto j respecto al suelo
        x2_espejo, y2_espejo = puntos[j][0], -puntos[j][1]
        distancia = raiz((x2_espejo-x1)**2 + (y2_espejo-y1)**2)
        matriz_D_prima[i][j] = distancia

print("\nMatriz D' (distancias entre puntos y espejos) (km):")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
for i, fila in enumerate(matriz_D_prima):
    print(f"{etiquetas[i]}: ", end="")
    for dist in fila:
        print(f"{dist:9.4f}", end="  ")
    print()
# kij = raiz(2)*D_prima_ij / penetracion_terreno
matriz_kij = [[0]*n_puntos for _ in range(n_puntos)]
for i in range(n_puntos):
    for j in range(n_puntos):
        matriz_kij[i][j] = raiz(2) * matriz_D_prima[i][j] / penetracion_terreno  
print("\nMatriz k_ij:")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
for i, fila in enumerate(matriz_kij):
    print(f"{etiquetas[i]}: ", end="")
    for k in fila:
        print(f"{k:9.2f}", end="  ")
    print()
# P_ij = pi/8 - k_ij*cos(tetha_ij)/(3*raiz(2)) siendo cos(tetha_ij) = (h_i + h_j )/ D_prima_ij
matriz_Pij = [[0]*n_puntos for _ in range(n_puntos)]
for i in range(n_puntos):
    for j in range(n_puntos):
        h_i = puntos[i][1]
        h_j = puntos[j][1]
        D_prima_ij = matriz_D_prima[i][j]
        if D_prima_ij != 0:
            cos_tetha_ij = (h_i + h_j) / D_prima_ij
        else:
            cos_tetha_ij = 0
        k_ij = matriz_kij[i][j]
        P_ij = (pi / 8) - (k_ij * cos_tetha_ij) / (3 * raiz(2))
        matriz_Pij[i][j] = P_ij
print("\nMatriz P_ij:")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
for i, fila in enumerate(matriz_Pij):
    print(f"{etiquetas[i]}: ", end="")
    for P in fila:
        print(f"{P:9.6f}", end="  ")
    print()

# Q_ij = 0.5*log neperiano(1.85138/k_ij) + k_ij*cos(tetha_ij)/(3*raiz(2))
matriz_Qij = [[0]*n_puntos for _ in range(n_puntos)]
for i in range(n_puntos):
    for j in range(n_puntos):
        k_ij = matriz_kij[i][j]
        h_i = puntos[i][1]
        h_j = puntos[j][1]
        D_prima_ij = matriz_D_prima[i][j]
        if D_prima_ij != 0:
            cos_tetha_ij = (h_i + h_j) / D_prima_ij
        else:
            cos_tetha_ij = 0
        if k_ij != 0:
            Q_ij = 0.5 * log(1.85138 / k_ij) + (k_ij * cos_tetha_ij) / (3 * raiz(2))
            
        else:
            Q_ij = 0
        matriz_Qij[i][j] = Q_ij
print("\nMatriz Q_ij:")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
for i, fila in enumerate(matriz_Qij):
    print(f"{etiquetas[i]}: ", end="")
    for Q in fila:
        print(f"{Q:9.6f}", end="  ")
    print()


# Resistencia de cada uno de los conductores (ultimo a tierra con dato de resistencia a tierra directamente)
resistencias_conductores = [resistencia_ca, resistencia_ca, resistencia_ca, resistencia_ca, resistencia_ca, resistencia_ca, resistencia_tierra]
print("\nResistencias de los conductores (Ohmios/km):")
for i, R in enumerate(resistencias_conductores):
    print(f"{etiquetas[i]}: {R:.6f} Ohmios/km")

"""
# Radio equivalente de cada conductor
def radio_equivalente(diametro):
    return math.exp(-0.25) * diametro/2  # en mm
radio_eq_conductores = [
    radio_equivalente(diametro),
    radio_equivalente(diametro),
    radio_equivalente(diametro),
    radio_equivalente(diametro_tierra)
]
print("\nRadio equivalente de los conductores (m):")
for i, r_eq in enumerate(radio_eq_conductores):
    print(f"{etiquetas[i]}: {r_eq:.6f} m")
"""

# Radio de cada conductor en mm
radio_conductores_mm = [
    diametro / 2,
    diametro / 2,
    diametro / 2,
    diametro / 2,
    diametro / 2,
    diametro / 2,
    diametro_tierra / 2
]
print("\nRadio de los conductores (mm):")
for i, r in enumerate(radio_conductores_mm):
    print(f"{etiquetas[i]}: {r:.6f} mm")

# Matriz de impedancias con:
# Z_ii = R_i + j*mu0*frecuencia*((1/(4*n_i)+log(D_prima_ii/r_i)))+mu0*2*frecuencia*(P_ii + j*Q_ii) 
# Z_ij = j*mu0*frecuencia*(log(D_prima_ij/D_ij)) + mu0*2*frecuencia*(P_ij + j*Q_ij) para i != j
matriz_impedancias = [[0]*n_puntos for _ in range(n_puntos)]
for i in range(n_puntos):
    for j in range(n_puntos):
        R_i = resistencias_conductores[i]
        D_ij = matriz_distancias[i][j]
        D_prima_ij = matriz_D_prima[i][j]
        P_ij = matriz_Pij[i][j]
        Q_ij = matriz_Qij[i][j]
        r_i = radio_conductores_mm[i] / 1000  # Convertir mm a m
        if i == j:
            n_i = 1  # Número de conductores en fase, asumir 1 para simplificar
            Z_ii = R_i + (1j * mu0 * frecuencia * ((1 / (4 * n_i)) + log(D_prima_ij / r_i)) + mu0 * 2 * frecuencia * (P_ij + 1j * Q_ij))*1000
            matriz_impedancias[i][j] = Z_ii
        else:
            Z_ij = (1j * mu0 * frecuencia * (log(D_prima_ij / D_ij)) + mu0 * 2 * frecuencia * (P_ij + 1j * Q_ij))*1000
            matriz_impedancias[i][j] = Z_ij
print("\nMatriz de impedancias (Ohmios/km):")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
for i, fila in enumerate(matriz_impedancias):
    print(f"{etiquetas[i]}: ", end="")
    for Z in fila:
        print(f"{Z:18.4f}", end="  ")
    print()

""" Como la linea posee cable de tierra, es necesario realizar un análisis matricial para eliminarlos y obtener una matriz 3*3
que representa las impedancias por fase
Z=[Zf, Zft; Ztf, Zt]
Zfas = Zf - Zft * inv(Zt) * Ztf
"""

Z = np.array(matriz_impedancias)
Zf = Z[0:6, 0:6]
Zt = Z[6:7, 6:7]
Zft = Z[0:6, 6:7]
Ztf = Z[6:7, 0:6]
Zt_inv = np.linalg.inv(Zt)
Zfas = Zf - Zft @ Zt_inv @ Ztf
print("\nMatriz de impedancias por fase (Ohmios/km):")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
for i, fila in enumerate(Zfas):
    print(f"{etiquetas[i]}: ", end="")
    for Z in fila:
        print(f"{Z:18.3f}", end="  ")
    print()
# Finalmente, multiplicar por la longitud de la línea para obtener las impedancias totales
Zfas_total = Zfas * longitud
print("\nMatriz de impedancias por fase total para la longitud dada (Ohmios):")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
for i, fila in enumerate(Zfas_total):
    print(f"{etiquetas[i]}: ", end="")
    for Z in fila:
        print(f"{Z:18.3f}", end="  ")
    print()


# Impedancias de secuencia
A = np.array([[1, 1, 1, 0, 0, 0],
              [complex(-0.5, -raiz(3)/2), complex(-0.5, raiz(3)/2), 1, 0, 0, 0],
              [complex(-0.5, raiz(3)/2), complex(-0.5, -raiz(3)/2), 1, 0, 0, 0],
              [0, 0, 0, 1, 1, 1],
              [0, 0, 0, complex(-0.5, -raiz(3)/2), complex(-0.5, raiz(3)/2), 1],
              [0, 0, 0, complex(-0.5, raiz(3)/2), complex(-0.5, -raiz(3)/2), 1]])
A_inv = np.linalg.inv(A)
Z_seq = A_inv @ Zfas @ A
print("\nMatriz de impedancias de secuencia (Ohmios):")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
for i, fila in enumerate(Z_seq):
    print(f"Fila {i}: ", end="")
    for Z in fila:
        print(f"{Z:18.3f}", end="  ")
    print()

# Matriz anterior de impedancias de secuencia pero con argumento y angulo
print("\nMatriz de impedancias de secuencia con magnitud y ángulo (Ohmios):")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
for i, fila in enumerate(Z_seq):
    print(f"Fila {i}: ", end="")
    for Z in fila:
        magnitud = abs(Z)
        angulo = math.degrees(math.atan2(Z.imag, Z.real))
        print(f"{magnitud:12.4f} ∠ {angulo:8.2f}°", end="  ")
    print()

# Impedancia homopolar de la linea (Z0)
Z0 = Z_seq[2, 2]
print(f"\nImpedancia homopolar de la línea (Z0): {Z0:.3f} Ohmios/km")
# Impedancia directa e inversa de la linea (Z1)
Z1 = Z_seq[1, 1]
print(f"Impedancia directa e inversa de la línea (Z1): {Z1:.3f} Ohmios/km")



#Teniendo en cuenta la longitud de la línea
Z_seq_total = Z_seq * longitud
print("\nMatriz de impedancias de secuencia total para la longitud dada (Ohmios):")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
for i, fila in enumerate(Z_seq_total):
    print(f"Fila {i}: ", end="")
    for Z in fila:
        print(f"{Z:18.3f}", end="  ")
    print()
# Matriz anterior de impedancias de secuencia pero con argumento y angulo
print("\nMatriz de impedancias de secuencia total con magnitud y ángulo para la longitud dada (Ohmios):")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
for i, fila in enumerate(Z_seq_total):
    print(f"Fila {i}: ", end="")
    for Z in fila:
        magnitud = abs(Z)
        angulo = math.degrees(math.atan2(Z.imag, Z.real))
        print(f"{magnitud:12.3f} ∠ {angulo:8.2f}°", end="  ")
    print()

# Impedancia homopolar de la linea (Z0)
Z0_total = Z_seq_total[2, 2]
print(f"\nImpedancia homopolar de la línea total (Z0): {Z0_total:.3f} Ohmios")
# Impedancia directa e inversa de la linea (Z1)
Z1_total = Z_seq_total[1, 1]
print(f"Impedancia directa e inversa de la línea total (Z1): {Z1_total:.3f} Ohmios")
 


# Resistencias R0/R1
R0 = Z0.real
R1 = Z1.real
print(f"\nR0/R1: {R0/R1:.3f}")
# Reactancias X0/X1
X0 = Z0.imag
X1 = Z1.imag
print(f"X0/X1: {X0/X1:.3f}")




#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# MATRIZ DE CAPACIDADES
#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# Coeficientes de potencial
# C_ii = (1/(2*pi*epsilon0))*log(D_prima_ii/r_i)
# C_ij = (1/(2*pi*epsilon0))*log(D_prima_ij/D_ij) para i != j
matriz_capacidades = [[0]*n_puntos for _ in range(n_puntos)]
for i in range(n_puntos):
    for j in range(n_puntos):
        D_ij = matriz_distancias[i][j]/1000   # Convertir m a km
        D_prima_ij = matriz_D_prima[i][j]/1000   # Convertir m a km
        r_i = radio_conductores_mm[i] / 1000000  # Convertir mm a km
        if i == j:
            C_ii = ((1 / (2 * pi * epsilon0)) * log(D_prima_ij / r_i))/1000000000
            matriz_capacidades[i][j] = C_ii
        else:
            C_ij = (1 / (2 * pi * epsilon0)) * log(D_prima_ij / D_ij)/1000000000
            matriz_capacidades[i][j] = C_ij
print("\nMatriz de capacidades (km/uF):")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
for i, fila in enumerate(matriz_capacidades):
    print(f"{etiquetas[i]}: ", end="")
    for C in fila:
        print(f"{C:18.2f}", end="  ")
    print()
# Matriz de capacitancias por fase (eliminando tierra)
C = np.array(matriz_capacidades)
Cf = C[0:6, 0:6]
Ct = C[6:7, 6:7]
Cft = C[0:6, 6:7]
Ctf = C[6:7, 0:6]
Ct_inv = np.linalg.inv(Ct)
Cfas = Cf - Cft @ Ct_inv @ Ctf
print("\nMatriz de capacidades por fase (km/uF):")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7")
for i, fila in enumerate(Cfas):
    print(f"{etiquetas[i]}: ", end="")
    for C in fila:
        print(f"{C:18.2f}", end="  ")
    print()
#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# MATRIZ DE SUSCEPTANCIAS
#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# Matriz de susceptancias por km (B=j*2*pi*frecuencia*C^-1) solo parte imaginaria, no mostrando la parte real
Cfas_inv = np.linalg.inv(Cfas)
Bfas = 1j * 2 * pi * frecuencia * Cfas_inv
print("\nMatriz de susceptancias por fase (uS/km):")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
for i, fila in enumerate(Bfas):
    print(f"{etiquetas[i]}: ", end="")
    for B in fila:
        print(f"{B:18.3f}", end="  ")
    print()

# Matriz de susceptancias total para la longitud dada
Bfas_total = Bfas * longitud
print("\nMatriz de susceptancias por fase total para la longitud dada (uS):")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7    Punto8")
for i, fila in enumerate(Bfas_total):
    print(f"{etiquetas[i]}: ", end="")
    for B in fila:
        print(f"{B:18.2f}", end="  ")
    print()
# Susceptancias de secuencia
B_seq = A_inv @ Bfas @ A
print("\nMatriz de susceptancias de secuencia (uS):")
print("     Punto1    Punto2    Punto3    Punto4    Punto6    Punto7 ")
for i, fila in enumerate(B_seq):
    print(f"Fila {i}: ", end="")
    for B in fila:
        print(f"{B:18.6f}", end="  ")
    print()
# Susceptancia homopolar de la linea (B0)
B0 = B_seq[2, 2]
print(f"\nSusceptancia homopolar de la línea (B0): {B0:.3f} uS/km")
# Susceptancia directa e inversa de la linea (B1)
B1 = B_seq[1, 1]
print(f"Susceptancia directa e inversa de la línea (B1): {B1:.3f} uS/KM")
# Susceptancia homopolar de la linea total (B0)
B0_total = B_seq[2, 2] * longitud
print(f"\nSusceptancia homopolar de la línea total (B0): {B0_total:.3f} uS")
# Susceptancia directa e inversa de la linea total (B1)
B1_total = B_seq[1, 1] * longitud
print(f"Susceptancia directa e inversa de la línea total (B1): {B1_total:.3f} uS")

#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# CÁLCULO IMPEDANCIA CARACTERÍSTICA Y CONSTANTE DE PROPAGACIÓN
#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# Impedancia característica (Zc=raiz((R+jX)/(jB)))
Zc = cmath.sqrt(Z1 / (1 * B1*1e-6))
print(f"Zc: {Zc:.3f}")
magnitud_Zc = abs(Zc)
angulo_Zc = math.degrees(math.atan2(Zc.imag, Zc.real))
print(f"\nImpedancia característica de secuencia directa (Zc1): {magnitud_Zc:.4f} ∠ {angulo_Zc:.2f}° Ohmios")

# Constante de propagación (gamma=raiz((R+jX)*(jB)))
gamma = cmath.sqrt(Z1 * (1 * B1*1e-6))
print("Gamma:", gamma)
magnitud_gamma = abs(gamma)
angulo_gamma = math.degrees(math.atan2(gamma.imag, gamma.real))
print(f"\nConstante de propagación de secuencia directa (γ1): {magnitud_gamma:.6f} ∠ {angulo_gamma:.2f}° 1/km")
# Constante de propagación total para la longitud dada
gamma_total = gamma * longitud
magnitud_gamma_total = abs(gamma_total)
angulo_gamma_total = math.degrees(math.atan2(gamma_total.imag, gamma_total.real))
print(f"\nConstante de propagación para la longitud dada: {gamma_total} = {magnitud_gamma_total:.6f} ∠ {angulo_gamma_total:.2f}°")

#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# POTENCIA CARACTERÍSTICA
#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# Pc = tension_nominal**2 / Zc
Pc = (tension_nominal) ** 2 / Zc
magnitud_Pc = abs(Pc)
angulo_Pc = math.degrees(math.atan2(Pc.imag, Pc.real))
print(f"\nPotencia característica: {Pc} = {magnitud_Pc:.4f} ∠ {angulo_Pc:.2f}° kVA")


#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# CAIDA DE TENSIÓN
#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# AU=potencia_transportada(MW)*longitud*(R*cos(phi)+X*sin(phi))/((tension_nominal**2)*10*cos(phi))
potencia_transportada_MW = potencia_transportada * cos_phi  # Convertir MVA a MW usando el factor de potencia
AU = potencia_transportada_MW * 1000 * longitud * (Z1.real * cos_phi + Z1.imag * math.sin(math.acos(cos_phi))) / ((tension_nominal ** 2) * 10 * cos_phi)
print(f"\nCaída de tensión AU: {AU:.4f} %")
# Poner si la caida de tension es inferior al 5%: La caída de tensión es inferior al 5%; y si no lo es: La caída de tensión es superior al 5%
if AU < 5:
    print("La caída de tensión es inferior al 5%")
else:
    print("La caída de tensión es superior al 5%")

#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# DENSIDAD MÁXIMA DE CORRIENTE E INTENSIDAD MÁXIMA POR CABLE
#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
""" Valores de densidades de corriente (A/mm2) para diferentes secciones y luego dependiendo de la seccion se hará una regla de 3 para sacar el valor correspondiente:
10 mm2: -
15 mm2: 6 A/mm2
25 mm2: 5 A/mm2
35 mm2: 4.55 A/mm2
50 mm2: 4.0 A/mm2
70 mm2: 3.55 A/mm2
95 mm2: 3.2 A/mm2
125 mm2: 2.9 A/mm2
160 mm2: 2.7 A/mm2
200 mm2: 2.5 A/mm2
250 mm2: 2.3 A/mm2
300 mm2: 2.15 A/mm2
400 mm2: 1.95 A/mm2
500 mm2: 1.8 A/mm2
600 mm2: 1.65 A/mm2
"""
secciones = [15, 25, 35, 50, 70, 95, 125, 160, 200, 250, 300, 400, 500, 600]
densidades = [6, 5, 4.55, 4.0, 3.55, 3.2, 2.9, 2.7, 2.5, 2.3, 2.15, 1.95, 1.8, 1.65]  # A/mm2

def densidad_corriente(seccion):
    if seccion < secciones[0] or seccion > secciones[-1]:
        return None  # Fuera de rango
    for i in range(len(secciones) - 1):
        if secciones[i] <= seccion <= secciones[i + 1]:
            # Regla de tres
            densidad = densidades[i] + (densidades[i + 1] - densidades[i]) * (seccion - secciones[i]) / (secciones[i + 1] - secciones[i])
            return densidad
    return None
densidad_maxima = densidad_corriente(seccion)
print(f"\nDensidad máxima de corriente para la sección {seccion} mm2: {densidad_maxima:.4f} A/mm2")

""" Coeficiente reductor para diferentes composiciones de Al-Ac:
30+7: 0.916
6+1: 0.937
26+7: 0.937
54+7: 0.95
45+7: 0.97
"""
composiciones = ["30+7", "6+1", "26+7", "54+7", "45+7"]
coef_reductores = [0.916, 0.937, 0.937, 0.95, 0.97]
def coeficiente_reductor(composicion):
    if composicion in composiciones:
        index = composiciones.index(composicion)
        return coef_reductores[index]
    return None
coef_reductor = coeficiente_reductor(composicion)
print(f"Coeficiente reductor para la composición {composicion}: {coef_reductor:.4f}")

densidad_max_con_reduccion = densidad_maxima * coef_reductor
print(f"Densidad máxima de corriente con coeficiente reductor: {densidad_max_con_reduccion:.4f} A/mm2")

intensidad_maxima_conductor = densidad_max_con_reduccion * seccion  # A
print(f"Corriente máxima admisible del conductor: {intensidad_maxima_conductor:.2f} A")

#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# POTENCIA MÁXIMA ADMISIBLE POR INTENSIDAD
#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
potencia_maxima_admisible = (intensidad_maxima_conductor * tension_nominal * raiz(3)) * cos_phi 
print(f"\nPotencia máxima admisible por intensidad: {potencia_maxima_admisible/1000:.2f} MW")
# Comparar potencia máxima admisible con potencia transportada
if potencia_maxima_admisible/1000 >= potencia_transportada_MW:
    print("La potencia máxima admisible por intensidad es mayor que la potencia transportada.")
else:
    print("La potencia máxima admisible por intensidad es menor que la potencia transportada.")

#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# POTENCIA DE TRANSPORTE EN FUNCIÓN DE CONDICIONES METEOROLÓGICAS
#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# Calor aportado por radiación solar para invierno y verano Qs = coeficiente_absorcion * radiacion * diametro/1000
Qs_invierno = coeficiente_absorcion * radiacion_invierno * diametro/1000 # W/m
Qs_verano = coeficiente_absorcion * radiacion_verano * diametro/1000 # W/m
print(f"\nCalor aportado por radiación solar en invierno: {Qs_invierno:.3f} W/m")
print(f"Calor aportado por radiación solar en verano: {Qs_verano:.3f} W/m")

# Calor cedido por radiación Qr = pi * diametro/1000 * emisividad_conductor * cte_boltzmann * ((Tc+273.15)^4 - (Ta+273.15)^4)
cte_boltzmann = 5.6704e-8 # W/m2K4
Qr_invierno = pi * (diametro/1000) * emisividad_conductor * cte_boltzmann * ((Tc + 273.15)**4 - (temperatura_invierno + 273.15)**4) # W/m
Qr_verano = pi * (diametro/1000) * emisividad_conductor * cte_boltzmann * ((Tc + 273.15)**4 - (temperatura_verano + 273.15)**4) # W/m
print(f"\nCalor cedido por radiación en invierno: {Qr_invierno:.3f} W/m")
print(f"Calor cedido por radiación en verano: {Qr_verano:.3f} W/m")

# Calor cedido por convección Qc = pi * conductividad_termica * (Tc - Ta) * Nu
# 1.  Convección natural: Nu = A * (Gr*Pr)^m
""" Valores para A y m dependiendo del rango de Gr*Pr:
0.1 <= Gr*Pr < 100: A=1.02, m=0.148
100 <= Gr*Pr < 10000: A=0.85, m=0.188
10000 <= Gr*Pr < 1e7: A=0.48, m=0.25
1e7 <= Gr*Pr < 1e12: A=0.125, m=0.333
"""
def coeficientes_conveccion_natural(Gr_Pr):
    if 0.1 <= Gr_Pr < 100:
        return 1.02, 0.148
    elif 100 <= Gr_Pr < 10000:
        return 0.85, 0.188
    elif 10000 <= Gr_Pr < 1e7:
        return 0.48, 0.25
    elif 1e7 <= Gr_Pr < 1e12:
        return 0.125, 0.333
    else:
        return None, None
# Gr = (diametro/1000)^3 * (Tc - Ta) * g / ((Tav + 27.15) * (viscosidad_cinematica)^2)
g = 9.81  # m/s2
# viscosidad_cinematica = 1.32e-5 + 9.5e-8 * Tav
def viscosidad_cinematica(Tav):
    return 1.32e-5 + 9.5e-8 * Tav
viscosidad_cinematica_invierno = viscosidad_cinematica((Tc + temperatura_invierno) / 2)
viscosidad_cinematica_verano = viscosidad_cinematica((Tc + temperatura_verano) / 2)
print(f"\nViscosidad cinemática en invierno: {viscosidad_cinematica_invierno:.8f} m2/s")
print(f"Viscosidad cinemática en verano: {viscosidad_cinematica_verano:.8f} m2/s")
# densidad_relativa_aire = exp(-1.16e-4* h)

def Gr(Ta, viscosidad_cinematica):
    Tav = (Tc + Ta) / 2
    return ((diametro/1000)**3 * (Tc - Ta) * g) / ((Tav + 273.15) * (viscosidad_cinematica**2))
Gr_invierno = Gr(temperatura_invierno, viscosidad_cinematica_invierno)
Gr_verano = Gr(temperatura_verano, viscosidad_cinematica_verano)
print(f"\nNúmero de Grashof en invierno: {Gr_invierno:.2e}")
print(f"Número de Grashof en verano: {Gr_verano:.2e}")

# Pr = calor_esp_aire * viscosidad_dinamica / conductividad_termica_aire
calor_esp_aire = 1005  # J/kgK
# viscosidad_dinamica = (17.239 + 4.625e-2 * Tav - 2.03e-5 * Tav **2)e-6 kg/m s
def viscosidad_dinamica(Tav):
    return (17.239 + 4.625e-2 * Tav - 2.03e-5 * Tav **2) * 1e-6  # kg/m s
viscosidad_dinamica_invierno = viscosidad_dinamica((Tc + temperatura_invierno) / 2)
viscosidad_dinamica_verano = viscosidad_dinamica((Tc + temperatura_verano) / 2)
print(f"\nViscosidad dinámica en invierno: {viscosidad_dinamica_invierno:.6e} kg/m s")
print(f"Viscosidad dinámica en verano: {viscosidad_dinamica_verano:.6e} kg/m s")
# conductividad_termica_aire = 2.368e-2 + 7.23e-5 * Tav - 2.763e-8 * Tav**2 W/Km
def conductividad_termica_aire(Tav):
    return 2.368e-2 + 7.23e-5 * Tav - 2.763e-8 * Tav**2  # W/Km
conductividad_termica_aire_invierno = conductividad_termica_aire((Tc + temperatura_invierno) / 2)
conductividad_termica_aire_verano = conductividad_termica_aire((Tc + temperatura_verano) / 2)
print(f"\nConductividad térmica del aire en invierno: {conductividad_termica_aire_invierno:.6e} W/Km")
print(f"Conductividad térmica del aire en verano: {conductividad_termica_aire_verano:.6e} W/Km")

Pr_invierno = calor_esp_aire * viscosidad_dinamica_invierno / conductividad_termica_aire_invierno
Pr_verano = calor_esp_aire * viscosidad_dinamica_verano / conductividad_termica_aire_verano
print(f"\nNúmero de Prandtl en invierno: {Pr_invierno:.4f}")
print(f"Número de Prandtl en verano: {Pr_verano:.4f}")

Gr_Pr_invierno = Gr_invierno * Pr_invierno
Gr_Pr_verano = Gr_verano * Pr_verano
Nu_invierno = 0
Nu_verano = 0
A_invierno, m_invierno = coeficientes_conveccion_natural(Gr_Pr_invierno)
if A_invierno is not None:
    Nu_invierno = A_invierno * (Gr_Pr_invierno ** m_invierno)
A_verano, m_verano = coeficientes_conveccion_natural(Gr_Pr_verano)
if A_verano is not None:
    Nu_verano = A_verano * (Gr_Pr_verano ** m_verano)
print(f"\nNúmero de Nusselt en invierno: {Nu_invierno:.4f}")
print(f"Número de Nusselt en verano: {Nu_verano:.4f}")

# 2. Convección forzada: Nu = B1 * Re^n
# Re = (diametro/1000) * velocidad_viento / viscosidad_cinematica
# Rugosidad: Rf = diametro_alambre_ext / (2*(diametro-diametro_alambe_ext))
Rf = diametro_alambre_ext / (2 * (diametro - diametro_alambre_ext))
print(f"Rf:{Rf:.3f}")
def Re(viscosidad_cinematica):
    return (diametro/1000) * velocidad_viento / viscosidad_cinematica
Re_invierno = Re(viscosidad_cinematica_invierno)
Re_verano = Re(viscosidad_cinematica_verano)
print(f"\nNúmero de Reynolds en invierno: {Re_invierno:.2e}")
print(f"Número de Reynolds en verano: {Re_verano:.2e}")







Qc_invierno = 100.13
Qc_verano =78.74






# Resultados corriente máxima: I = raiz((Qr+Qc-Qs)/resistencia_ca/1000)
print(f"\nResistencia del conductor a la temperatura de cálculo: {resistencia_ca:.6f} Ohmios/km")
I_max_invierno = raiz((Qr_invierno + Qc_invierno - Qs_invierno)/ (resistencia_ca * 1e-3))
I_max_verano = raiz((Qr_verano + Qc_verano - Qs_verano) / (resistencia_ca * 1e-3))
print(f"\nCorriente máxima admisible en invierno según condiciones meteorológicas: {I_max_invierno:.2f} A")
print(f"Corriente máxima admisible en verano según condiciones meteorológicas: {I_max_verano:.2f} A")

#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# POTENCIA MÁXIMA DE TRANSPORTE SEGÚN CONDICIONES METEOROLÓGICAS
#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
potencia_maxima_invierno = (I_max_invierno * tension_nominal * raiz(3)) * cos_phi
potencia_maxima_verano = (I_max_verano * tension_nominal * raiz(3)) * cos_phi
print(f"\nPotencia máxima de transporte en invierno según condiciones meteorológicas: {potencia_maxima_invierno/1000:.2f} MVA")
print(f"Potencia máxima de transporte en verano según condiciones meteorológicas: {potencia_maxima_verano/1000:.2f} MVA")












#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# PÉRDIDAS DE POTENCIA
#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------

perdidas_potencia = ((potencia_transportada_MW * resistencia_ca * longitud) / (tension_nominal**2 * cos_phi**2))*100
print(f"\nPérdidas de potencia en la línea: {perdidas_potencia:.5f} %")
# En valor absoluto
perdidas_potencia_valor = (potencia_transportada_MW * resistencia_ca * longitud) / (tension_nominal**2 * cos_phi**2) * potencia_transportada_MW
print(f"Pérdidas de potencia en la línea en valor absoluto: {perdidas_potencia_valor:.5f} MW")

#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# CORTOCIRCUITO MÁXIMO
#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
""" 
Valores de c_conductor, densidad_conductor, cte_conductor, alpha_conductor en función de si el material es Cobre, Aluminio-Acero o Acero:
Material: Cobre
c_conductor = 390 J/kgK
densidad_conductor = 8900 kg/m3
cte_conductor = 56e6 1/ohmio m
alpha_conductor = 0.0039 1/ºC
Material: Aluminio-Acero
c_conductor = 910 J/kgK
densidad_conductor = 2700 kg/m3
cte_conductor = 34.8e6 1/ohmio m
alpha_conductor = 0.004 1/ºC
Material: Acero
c_conductor = 480 J/kgK
densidad_conductor = 7850 kg/m3
cte_conductor = 7.25e6 1/ohmio m
alpha_conductor = 0.0045 1/ºC
"""
materiales = ["Cobre", "Aluminio-Acero", "Acero"]
c_conductores = [390, 910, 480]  # J/kgK
densidades_conductores = [8900, 2700, 7850]  # kg/m3
ctes_conductores = [56e6, 34.8e6, 7.25e6]  # 1/ohmio m
alphas_conductores = [0.0039, 0.004, 0.0045]  # 1/ºC
def propiedades_material(material):
    if material in materiales:
        index = materiales.index(material)
        return (c_conductores[index], densidades_conductores[index], ctes_conductores[index], alphas_conductores[index])
    return (None, None, None, None)
c_conductor, densidad_conductor, cte_conductor, alpha_conductor = propiedades_material(material)
print(f"c_conductor, densidad_conductor, cte_conductor, alpha_conductor", c_conductor, densidad_conductor, cte_conductor, alpha_conductor )
""" 
Temperatura max recomendada durante un circuito segun el tipo de material
Acero: 300ºC
Resto de materiales: 200ºC
"""
if material == "Acero":
    temperatura_max_recomendada = 300
else:
    temperatura_max_recomendada = 200



# Factor K
multiplicacion = log ((1+alpha_conductor*(temperatura_max_recomendada-20))/(1+alpha_conductor*(Tc-20)))
print(f"multiplicación:", multiplicacion)
K = (raiz (cte_conductor*c_conductor*densidad_conductor*multiplicacion/alpha_conductor))/1000000
print(f"Factor K: {K} Araiz(s)/mm2")

Icc_max = K * seccion / raiz(tiempo_accionamiento_proteccion)
print(f"Icc max: {Icc_max} A")



#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# EFECTO CORONA
#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
beta = 1 # Coeficiente reductor para conductores múltiples
presion_barometrica = 1/(log(log(76)-altitud_media/18330))*100 
print(f"Presión barométrica {presion_barometrica} cmHg")
# Factor corrección densidad aire: delta
def delta (temperatura):
    return 3.92*presion_barometrica/(273+temperatura)
delta_invierno = delta (temperatura_invierno)
delta_verano = delta (temperatura_verano)
print(f"Factor corrección densidad aire invierno {delta_invierno}")
print(f"Factor corrección densidad aire verano {delta_verano}")


DMG = 100*(matriz_distancias[0][1]*matriz_distancias[1][2]*matriz_distancias[2][0])**(1/6)
print(f"DMG: {DMG}")

DMG = 807
print(f"DMG: {DMG}")

# Tensión crítica dieléctrica
mc = 0.83
mt_invierno = 0.8
mt_verano = 1
Ep = 21.21

def Uc (mt,delta):
    return raiz(3)*mc*mt*Ep*delta*diametro/20*1/beta*log(DMG/(diametro/20))
Uc_invierno = Uc(mt_invierno, delta_invierno)
Uc_verano = Uc(mt_verano,delta_verano)
print(f"Tensión critica invierno {Uc_invierno} kV")
print(f"Tensión crítica verano {Uc_verano} kV")


if Uc_invierno < tension_nominal:
    print(f"Hay pérdidas por efecto corona en invierno")
    Perdidas_efecto_corona_invierno = 241/delta_invierno * (frecuencia + 25) * raiz (diametro/20/DMG) * (tension_nominal - Uc_invierno)/raiz(3) * 1e-5
    print(f"Pérdidas por efecto corona en invierno: {Perdidas_efecto_corona_invierno} kW·km/fase")

else: print(f"No hay pérdidas por efecto corona en invierno")

if Uc_verano < tension_nominal:
    print(f"Hay pérdidas por efecto corona en verano")
    Perdidas_efecto_corona_verano = 241/delta_verano * (frecuencia + 25) * raiz (diametro/20/DMG) * (tension_nominal - Uc_verano)/raiz(3) * 1e-5
    print(f"Pérdidas por efecto corona en invierno: {Perdidas_efecto_corona_verano} kW·km/fase")
else: print(f"No hay pérdidas por efecto corona en verano")




#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# CAMPO MAGNÉTICO
#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

# Matriz coeficientes de potencial
print("\nMatriz de capacidades (km/uF):")
print("           Punto1               Punto2               Punto3               Punto4")
for i, fila in enumerate(matriz_capacidades):
    print(f"{etiquetas[i]}: ", end="")
    for C in fila:
        print(f"{C:18.2f}", end="  ")
    print()
# Matriz coeficientes de potencial inversa
matriz_capacidades_inversa = np.linalg.inv(matriz_capacidades) * 1000
print("\nMatriz de capacidades inversa(nF/km):")
print("           Punto1               Punto2               Punto3               Punto4")
for i, fila in enumerate(matriz_capacidades_inversa):
    print(f"{etiquetas[i]}: ", end="")
    for C in fila:
        print(f"{C:18.2f}", end="  ")
    print()

# Vector de potencial

d12 = raiz(14**2+7**2)
d23 = 14
d31 = 7
dmg = (d12*d23*d31)**(1/3)
print(f"{dmg}")













# ==========================================================
# GENERACIÓN AUTOMÁTICA DE DOCUMENTO WORD CON RESULTADOS
# ==========================================================

from docx import Document
import os
import re
import numpy as np

# ----------------------------------------------------------
# RUTAS
# ----------------------------------------------------------

# Carpeta donde está este .py
base_dir = os.path.dirname(os.path.abspath(__file__))

# Carpeta CE
carpeta_ce = os.path.join(base_dir, "CE")

# Plantilla Word
plantilla = os.path.join(carpeta_ce, "plantilla.docx")

# Documento Word de salida
salida = os.path.join(carpeta_ce, "CE_relleno.docx")

# Comprobación
if not os.path.exists(plantilla):
    raise FileNotFoundError(f"No se encuentra la plantilla Word en:\n{plantilla}")

# Abrir Word
doc = Document(plantilla)

# ----------------------------------------------------------
# FUNCIÓN DE REEMPLAZO AUTOMÁTICO
# ----------------------------------------------------------

def reemplazar_placeholders(doc):
    """
    Reemplaza:
    - {{variable}} -> valor escalar
    - {{matriz}}   -> tabla Word (si es matriz 2D)
    usando las variables globales del script.
    """

    patron = re.compile(r"\{\{(\w+)\}\}")

    def es_matriz(valor):
        return isinstance(valor, (list, np.ndarray)) and np.array(valor).ndim == 2

    def formatear_valor(valor):
        if isinstance(valor, complex):
            return f"{valor.real:.6f} + j{valor.imag:.6f}"
        elif isinstance(valor, float):
            return f"{valor:.6f}"
        else:
            return str(valor)

    def insertar_tabla(parrafo, matriz):
        matriz = np.array(matriz)
        filas, columnas = matriz.shape

        tabla = doc.add_table(rows=filas, cols=columnas)
#        tabla.style = "Table Grid"

        for i in range(filas):
            for j in range(columnas):
                tabla.cell(i, j).text = formatear_valor(matriz[i, j])

        # Eliminar {{matriz}}
        parrafo.text = ""

        # Insertar tabla justo debajo del párrafo
        parrafo._p.addnext(tabla._tbl)

    # ------------------------
    # PÁRRAFOS NORMALES
    # ------------------------
    for p in doc.paragraphs:
        texto = p.text
        matches = patron.findall(texto)

        for nombre in matches:
            if nombre in globals():
                valor = globals()[nombre]

                # MATRIZ → TABLA
                if es_matriz(valor):
                    insertar_tabla(p, valor)

                # ESCALAR
                else:
                    texto = texto.replace(
                        f"{{{{{nombre}}}}}",
                        formatear_valor(valor)
                    )
                    p.text = texto

    # ------------------------
    # TABLAS EXISTENTES
    # ------------------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto = cell.text
                matches = patron.findall(texto)

                for nombre in matches:
                    if nombre in globals():
                        valor = globals()[nombre]
                        texto = texto.replace(
                            f"{{{{{nombre}}}}}",
                            formatear_valor(valor)
                        )
                        cell.text = texto

# ----------------------------------------------------------
# EJECUCIÓN
# ----------------------------------------------------------

reemplazar_placeholders(doc)
doc.save(salida)

print("✅ Documento Word generado correctamente:")
print(salida)
