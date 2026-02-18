
import math
import numpy as np
from math import pi
from math import sqrt as raiz
from math import log
import cmath
import io
import sys

buffer = io.StringIO()
sys.stdout = buffer

# Coeficientes
mu0 = 4e-7 * pi # H/m
epsilon0 = 8.854e-12  # Permisividad del vacío en F/m



# Valores linea
coef_temp = 0.0039 # 1/°C
frecuencia = 50 # Hz
longitud = 1.53 # km
resistividad = 250 # Ohmios metro
tension_nominal = 220 # kV
cos_phi = 0.928
potencia_transportada = 160 #MVA
tiempo_accionamiento_proteccion = 0.5 # s
altitud_media = 843

# Caraterísticas conductor
conductor = "LA-380"
material = "Aluminio-Acero"
composicion = "54+7"  
diametro_alambre_ext = 2.82  # mm
seccion = 381 # mm2
diametro = 25.38 # mm
temperatura_invierno = 5 # °C
temperatura_verano = 24 # °C
Tc = 85 # °C
resistencia = 0.0857 # Ohmios
n_i = 2  # Número de conductores en fase
distancia_conductores = 400 #mm

emisividad_conductor = 0.5
coeficiente_absorcion = 0.5
radiacion_invierno = 110 # W/m²
radiacion_verano = 330 # W/m²
velocidad_viento = 0.6 # m/s

# Valores cable tierra
resistencia_tierra = 0.33 # Ω
seccion_tierra = 155.5 # mm²
diametro_tierra = 18 # mm
coef_temp_tierra = 14.4e-6 # 1/°C

# Puntos
tres=(4.3,24)
dos=(-4.1,27.3)
uno=(4.1,30.6)
cuatro=(0,35)


# Valores para cálculo del efecto corona
mc = 0.85
mt_invierno = 0.8
mt_verano = 1


# DMG y RMG
RMG = raiz (diametro/2 * distancia_conductores)
RMG_prima = raiz(distancia_conductores * (diametro/2) * math.exp(-1/4))


# Comparación los valores de diametro y resistencia de la parte superior con los valores de la normativa a partir del dato de conductor
diametro_normativa = {
    "LA-30": 7.14,
    "LA-56": 9.45,
    "LA-78": 11.34,
    "LA-110": 14.0,
    "LA-145": 15.75,
    "LA-180": 17.5,
    "LA-280": 21.8,
    "LA-380": 25.38,
    "LA-455": 27.72,
    "LA-545": 30.42,
    "LA-635": 32.85
}
resistencia_normativa = {
    "LA-30": 1.0749,
    "LA-56": 0.6136,
    "LA-78": 0.4261,
    "LA-110": 0.3066,
    "LA-145": 0.2422,
    "LA-180": 0.1962,
    "LA-280": 0.1194,
    "LA-380": 0.0857,
    "LA-455": 0.0718,
    "LA-545": 0.0596,
    "LA-635": 0.0511
}
diametro_alambre_ext_normativa = {
    "LA-30": 2.38,
    "LA-56": 3.15,
    "LA-78": 3.78,
    "LA-110": 2,
    "LA-145": 2.25,
    "LA-180": 2.5,
    "LA-280": 2.68,
    "LA-380": 2.82,
    "LA-455": 3.08,
    "LA-545": 3.38,
    "LA-635": 3.65
}
seccion_normativa = {
    "LA-30": 31.1,
    "LA-56": 54.6,
    "LA-78": 78.6,
    "LA-110": 116.2,
    "LA-145": 147.1,
    "LA-180": 181.6,
    "LA-280": 281.1,
    "LA-380": 381,
    "LA-455": 454.5,
    "LA-545": 547.3,
    "LA-635": 636.6
}
composicion_normativa = {
    "LA-30": "6+1",
    "LA-56": "6+1",
    "LA-78": "6+1",
    "LA-110": "30+7",
    "LA-145": "30+7",
    "LA-180": "30+7",
    "LA-280": "26+7",
    "LA-380": "54+7",
    "LA-455": "54+7",
    "LA-545": "54+7",
    "LA-635": "54+19"
}

diametro_calculado = diametro_normativa.get(conductor, None)
resistencia_calculada = resistencia_normativa.get(conductor, None)
diametro_alambre_ext_calculado = diametro_alambre_ext_normativa.get(conductor, None)
seccion_calculada = seccion_normativa.get(conductor, None)
composicion_calculada = composicion_normativa.get(conductor, None)
print(f"Diámetro para {conductor}: {diametro_calculado} mm")
print(f"Resistencia para {conductor}: {resistencia_calculada} Ω/km")
print(f"Diámetro alambre exterior para {conductor}: {diametro_alambre_ext_calculado} mm")
print(f"Sección para {conductor}: {seccion_calculada} mm²")
print(f"Composición para {conductor}: {composicion_calculada}")





# Cálculo de la resistencia a 85ºC y en ca
print("\n1. RESISTENCIA ELÉCTRICA DE LA LÍNEA")
def resistencia_a_temp(resistencia_20C, temp):
    return resistencia_20C * (1 + coef_temp * (temp - 20))/n_i
resistencia_85C = resistencia_a_temp(resistencia, 85)
print(f"Resistencia a 85ºC: {resistencia_85C:.6f} Ω/km")
reactancia_pelicular_85C = 10**-7 * 8 * pi * frecuencia * (1 / (resistencia_a_temp(resistencia, 85) / 1000))
print(f"Reactancia por efecto pelicular a 85ºC: {reactancia_pelicular_85C:.6f} Ω")
ys = reactancia_pelicular_85C ** 2 / (192 + 0.8*reactancia_pelicular_85C**2)
print(f"ys: {ys:.6f}")
resistencia_ca = resistencia_85C * (1 + ys)
print(f"Resistencia con efecto skin a 85ºC: {resistencia_ca:.6f} Ω/km")
# print(f"Resistencia en CA total a 85ºC: {resistencia_ca/n_i:.6f} Ω/km")
r_ca_longitud = resistencia_ca * longitud
print(f"Resistencia en CA para {longitud} km: {r_ca_longitud:.6f} Ω")


resistencia_ca = resistencia_a_temp(resistencia, 85) * (1 + 7.5*frecuencia**2*(diametro/10)**4*1e-7)
print(f"\nResistencia con efecto skin a 85ºC: {resistencia_ca:.6f} Ω/km")
r_ca_longitud = resistencia_ca * longitud
print(f"Resistencia en CA para {longitud} km: {r_ca_longitud:.6f} Ω")


resistencia_ca = resistencia_a_temp(resistencia, 20) * (1 + 7.5*frecuencia**2*(diametro/10)**4*1e-7)
print(f"\nResistencia con efecto skin a 20ºC: {resistencia_ca:.6f} Ω/km")
r_ca_longitud = resistencia_ca * longitud
print(f"Resistencia en CA para {longitud} km: {r_ca_longitud:.6f} Ω")



puntos = [uno, dos, tres, cuatro]
n_puntos = len(puntos)
matriz_distancias = [[0]*n_puntos for _ in range(n_puntos)]
for i in range(n_puntos):
    for j in range(n_puntos):
        x1, y1 = puntos[i]
        x2, y2 = puntos[j]
        distancia = raiz((x2-x1)**2 + (y2-y1)**2)
        matriz_distancias[i][j] = distancia

print(f"\nRMG: {RMG:.4f} mm")
print(f"RMG': {RMG_prima:.4f} mm")
DMG = (matriz_distancias[0][1]*matriz_distancias[1][2]*matriz_distancias[2][0]) ** (1/3) *1000
print(f"DMG: {DMG:.4f} mm")




# Calculo reactancia
print("\nCALCULO REACTANCIA Y SUSCEPTANCIA SIMPLIFICADO")
L = 2e-4 * math.log(DMG/RMG_prima)
print(f"\nL={L:.6f} H/km")
X = 2*pi*frecuencia*L
print(f"X={X:.6f} Ω/km")
C = 0.0556 / (math.log(DMG/RMG)) 
B = 2*pi*frecuencia*C
print(f"\nC={C:.6e} μF/km")
print(f"B={B:.6e} μS/km")

"""
I=potencia_transportada*1e6/(raiz(3)*tension_nominal*1e3)
print(f"\nIntensidad línea: {I:.2f} A")


potencia_transportada_MW = potencia_transportada * cos_phi  

Δ𝑈 = raiz(3) * I * longitud * (resistencia_ca * cos_phi + X * math.sin(math.acos(cos_phi))) 
print(f"\nCaída de tensión Δ𝑈: {Δ𝑈:.4f} V")
Δ𝑈 = potencia_transportada_MW * 1000 * longitud * (resistencia_ca * cos_phi + X * math.sin(math.acos(cos_phi))) / ((tension_nominal ** 2) * 10 * cos_phi)
print(f"Caída de tensión Δ𝑈: {Δ𝑈:.4f} %")


Δ𝑃=3*resistencia_ca*I**2*longitud
print(f"\nPérdidas en línea Δ𝑃: {Δ𝑃/1e6:.4f} MW")
Δ𝑃=Δ𝑃*100/(potencia_transportada_MW*1e6)
print(f"Pérdidas en línea Δ𝑃: {Δ𝑃:.4f} %")
"""
# Cálculo matriz de impedancias
print("\n2. MATRIZ DE IMPEDANCIAS")
penetracion_terreno = raiz(resistividad/(pi*frecuencia*mu0))
print(f"Penetración terreno: {penetracion_terreno:.6f} m")
# Matriz de distancias
puntos = [uno, dos, tres, cuatro]
n_puntos = len(puntos)
matriz_distancias = [[0]*n_puntos for _ in range(n_puntos)]
for i in range(n_puntos):
    for j in range(n_puntos):
        x1, y1 = puntos[i]
        x2, y2 = puntos[j]
        distancia = raiz((x2-x1)**2 + (y2-y1)**2)
        matriz_distancias[i][j] = distancia
print("\nMatriz de distancias (km):")
etiquetas = ["Punto1", "Punto2", "Punto3", "Punto4"]
for i, fila in enumerate(matriz_distancias):
    print(f"", end="")
    for dist in fila:
        print(f"{dist:9.4f}", end="  ")
    print()
# Matriz D_prima: distancia entre puntos y sus espejos respecto al suelo
# El espejo de un punto (x, y) respecto al suelo es (x, -y)
matriz_D_prima = [[0]*n_puntos for _ in range(n_puntos)]
for i in range(n_puntos):
    for j in range(n_puntos):
        x1, y1 = puntos[i]
        # Espejo del punto j respecto al suelo
        x2_espejo, y2_espejo = puntos[j][0], -puntos[j][1]
        distancia = raiz((x2_espejo-x1)**2 + (y2_espejo-y1)**2)
        matriz_D_prima[i][j] = distancia
print("\nMatriz D' (distancias entre puntos y espejos) (km):")
for i, fila in enumerate(matriz_D_prima):
    print(f"", end="")
    for dist in fila:
        print(f"{dist:9.4f}", end="  ")
    print()
# kij = raiz(2)*D_prima_ij / penetracion_terreno
matriz_kij = [[0]*n_puntos for _ in range(n_puntos)]
for i in range(n_puntos):
    for j in range(n_puntos):
        matriz_kij[i][j] = raiz(2) * matriz_D_prima[i][j] / penetracion_terreno  
print("\nMatriz k_ij:")
for i, fila in enumerate(matriz_kij):
    print(f"", end="")
    for k in fila:
        print(f"{k:9.4f}", end="  ")
    print()
# P_ij = pi/8 - k_ij*cos(tetha_ij)/(3*raiz(2)) siendo cos(tetha_ij) = (h_i + h_j )/ D_prima_ij
matriz_Pij = [[0]*n_puntos for _ in range(n_puntos)]
for i in range(n_puntos):
    for j in range(n_puntos):
        h_i = puntos[i][1]
        h_j = puntos[j][1]
        D_prima_ij = matriz_D_prima[i][j]
        if D_prima_ij != 0:
            cos_tetha_ij = (h_i + h_j) / D_prima_ij
        else:
            cos_tetha_ij = 0
        k_ij = matriz_kij[i][j]
        P_ij = (pi / 8) - (k_ij * cos_tetha_ij) / (3 * raiz(2))
        matriz_Pij[i][j] = P_ij
print("\nMatriz P_ij:")
for i, fila in enumerate(matriz_Pij):
    print(f" ", end="")
    for P in fila:
        print(f"{P:9.4f}", end="  ")
    print()
# Q_ij = 0.5*log neperiano(1.85138/k_ij) + k_ij*cos(tetha_ij)/(3*raiz(2))
matriz_Qij = [[0]*n_puntos for _ in range(n_puntos)]
for i in range(n_puntos):
    for j in range(n_puntos):
        k_ij = matriz_kij[i][j]
        h_i = puntos[i][1]
        h_j = puntos[j][1]
        D_prima_ij = matriz_D_prima[i][j]
        if D_prima_ij != 0:
            cos_tetha_ij = (h_i + h_j) / D_prima_ij
        else:
            cos_tetha_ij = 0
        if k_ij != 0:
            Q_ij = 0.5 * log(1.85138 / k_ij) + (k_ij * cos_tetha_ij) / (3 * raiz(2))     
        else:
            Q_ij = 0
        matriz_Qij[i][j] = Q_ij
print("\nMatriz Q_ij:")
for i, fila in enumerate(matriz_Qij):
    print(f" ", end="")
    for Q in fila:
        print(f"{Q:9.4f}", end="  ")
    print()
# Resistencia de cada uno de los conductores (ultimo a tierra con dato de resistencia a tierra directamente)
resistencias_conductores = [resistencia_ca/n_i, resistencia_ca/n_i, resistencia_ca/n_i, resistencia_tierra/n_i]
print("\nResistencias de los conductores (Ω/km):")
for i, R in enumerate(resistencias_conductores):
    print(f" {R:.4f} Ω/km")
# Radio de cada conductor en mm
radio_conductores_mm = [
    RMG,
    RMG,
    RMG,
    diametro_tierra/2
]
print("\nRadio de los conductores (mm):")
for i, r in enumerate(radio_conductores_mm):
    print(f" {r:.6f} mm")
# Matriz de impedancias con:
# Z_ii = R_i + j*mu0*frecuencia*((1/(4*n_i)+log(D_prima_ii/r_i)))+mu0*2*frecuencia*(P_ii + j*Q_ii) 
# Z_ij = j*mu0*frecuencia*(log(D_prima_ij/D_ij)) + mu0*2*frecuencia*(P_ij + j*Q_ij) para i != j
matriz_impedancias = [[0]*n_puntos for _ in range(n_puntos)]
for i in range(n_puntos):
    for j in range(n_puntos):
        R_i = resistencias_conductores[i]
        D_ij = matriz_distancias[i][j]
        D_prima_ij = matriz_D_prima[i][j]
        P_ij = matriz_Pij[i][j]
        Q_ij = matriz_Qij[i][j]
        r_i = radio_conductores_mm[i] / 1000  # Convertir mm a m
        if i == j:
            Z_ii = R_i + (1j * mu0 * frecuencia * ((1 / (4 * n_i)) + log(D_prima_ij / r_i)) + mu0 * 2 * frecuencia * (P_ij + 1j * Q_ij))*1000
            matriz_impedancias[i][j] = Z_ii
        else:
            Z_ij = (1j * mu0 * frecuencia * (log(D_prima_ij / D_ij)) + mu0 * 2 * frecuencia * (P_ij + 1j * Q_ij))*1000
            matriz_impedancias[i][j] = Z_ij
print("\nMatriz de impedancias (Ω/km):")
for i, fila in enumerate(matriz_impedancias):
    print(f" ", end="")
    for Z in fila:
        print(f"{Z:18.4f}", end="  ")
    print()
""" Como la linea posee cable de tierra, es necesario realizar un análisis matricial para eliminarlos y obtener una matriz 3*3
que representa las impedancias por fase
Z=[Zf, Zft; Ztf, Zt]
Zfas = Zf - Zft * inv(Zt) * Ztf
"""
Z = np.array(matriz_impedancias)
Zf = Z[0:3, 0:3]
Zt = Z[3:4, 3:4]
Zft = Z[0:3, 3:4]
Ztf = Z[3:4, 0:3]
Zt_inv = np.linalg.inv(Zt)
Zfas = Zf - Zft @ Zt_inv @ Ztf
print("\nMatriz de impedancias por fase (Ω/km):")
for i, fila in enumerate(Zfas):
    print(f" ", end="")
    for Z in fila:
        print(f"{Z:18.4f}", end="  ")
    print()


# Finalmente, multiplicar por la longitud de la línea para obtener las impedancias totales
Zfas_total = Zfas * longitud
print("\nMatriz de impedancias por fase total para la longitud dada (Ω):")
for i, fila in enumerate(Zfas_total):
    print(f" ", end="")
    for Z in fila:
        print(f"{Z:18.4f}", end="  ")
    print()

#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# IMPEDANCIAS DE SECUENCIA
#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
print("\n3. IMPEDANCIAS DE SECUENCIA")
A = np.array([[1, 1, 1],
              [complex(-0.5, -raiz(3)/2), complex(-0.5, raiz(3)/2), 1],
              [complex(-0.5, raiz(3)/2), complex(-0.5, -raiz(3)/2), 1]])
A_inv = np.linalg.inv(A)
Z_seq = A_inv @ Zfas @ A
print("\nMatriz de impedancias de secuencia (Ω/km):")
for i, fila in enumerate(Z_seq):
    print(f" ", end="")
    for Z in fila:
        print(f"{Z:18.4f}", end="  ")
    print()
# Matriz anterior de impedancias de secuencia pero con argumento y angulo
print("\nMatriz de impedancias de secuencia con magnitud y ángulo (Ω/km):")
for i, fila in enumerate(Z_seq):
    print(f"", end="")
    for Z in fila:
        magnitud = abs(Z)
        angulo = math.degrees(math.atan2(Z.imag, Z.real))
        print(f"{magnitud:12.4f} ∠ {angulo:8.2f}°", end="  ")
    print()
Z0 = Z_seq[2, 2]
print(f"\nImpedancia homopolar de la línea (Z0): {Z0:.4f} Ω")
Z1 = Z_seq[1, 1]
print(f"Impedancia directa e inversa de la línea (Z1): {Z1:.4f} Ω")
#Teniendo en cuenta la longitud de la línea
Z_seq_total = Z_seq * longitud
print("\nMatriz de impedancias de secuencia total para la longitud dada (Ω):")
for i, fila in enumerate(Z_seq_total):
    print(f"", end="")
    for Z in fila:
        print(f"{Z:18.4f}", end="  ")
    print()
# Matriz anterior de impedancias de secuencia pero con argumento y angulo
print("\nMatriz de impedancias de secuencia total con magnitud y ángulo para la longitud dada (Ω):")
for i, fila in enumerate(Z_seq_total):
    print(f" ", end="")
    for Z in fila:
        magnitud = abs(Z)
        angulo = math.degrees(math.atan2(Z.imag, Z.real))
        print(f"{magnitud:12.4f} ∠ {angulo:8.2f}°", end="  ")
    print()
Z0_total = Z_seq_total[2, 2]
print(f"\nImpedancia homopolar de la línea total (Z0): {Z0_total:.4f} Ω")
Z1_total = Z_seq_total[1, 1]
print(f"Impedancia directa e inversa de la línea total (Z1): {Z1_total:.4f} Ω")
# Resistencias R0/R1
R0 = Z0.real
R1 = Z1.real
print(f"\nR0/R1: {R0/R1:.3f}")
# Reactancias X0/X1
X0 = Z0.imag
X1 = Z1.imag
print(f"X0/X1: {X0/X1:.3f}")

#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# MATRIZ DE CAPACIDADES
#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
print("\n4. MATRIZ DE CAPACIDADES")
# Coeficientes de potencial
# C_ii = (1/(2*pi*epsilon0))*log(D_prima_ii/r_i)
# C_ij = (1/(2*pi*epsilon0))*log(D_prima_ij/D_ij) para i != j
matriz_capacidades = [[0]*n_puntos for _ in range(n_puntos)]
for i in range(n_puntos):
    for j in range(n_puntos):
        D_ij = matriz_distancias[i][j]/1000   # Convertir m a km
        D_prima_ij = matriz_D_prima[i][j]/1000   # Convertir m a km
        r_i = radio_conductores_mm[i] / 1000000  # Convertir mm a km
        if i == j:
            C_ii = ((1 / (2 * pi * epsilon0)) * log(D_prima_ij / r_i))/1000000000
            matriz_capacidades[i][j] = C_ii
        else:
            C_ij = (1 / (2 * pi * epsilon0)) * log(D_prima_ij / D_ij)/1000000000
            matriz_capacidades[i][j] = C_ij
print("\nMatriz de coeficientes de potencial de fase incluyendo cables de tierra (km/μF):")
for i, fila in enumerate(matriz_capacidades):
    print(f" ", end="")
    for C in fila:
        print(f"{C:18.4f}", end="  ")
    print()
# Matriz de capacitancias por fase (eliminando tierra)
C = np.array(matriz_capacidades)
Cf = C[0:3, 0:3]
Ct = C[3:4, 3:4]
Cft = C[0:3, 3:4]
Ctf = C[3:4, 0:3]
Ct_inv = np.linalg.inv(Ct)
Cfas = Cf - Cft @ Ct_inv @ Ctf
print("\nMatriz de coeficientes de potencial de fase (km/μF):")
for i, fila in enumerate(Cfas):
    print(f" ", end="")
    for C in fila:
        print(f"{C:18.4f}", end="  ")
    print()

#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# MATRIZ DE SUSCEPTANCIAS
#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
print("\n5. MATRIZ DE SUSCEPTANCIAS")
# Matriz de susceptancias por km (B=j*2*pi*frecuencia*C^-1) solo parte imaginaria, no mostrando la parte real
Cfas_inv = np.linalg.inv(Cfas)
Bfas = 1j * 2 * pi * frecuencia * Cfas_inv
print("\nMatriz de susceptancias por fase (μS/km):")
for i, fila in enumerate(Bfas):
    print(f" ", end="")
    for B in fila:
        print(f"{B:18.4f}", end="  ")
    print()
# Matriz de susceptancias total para la longitud dada
Bfas_total = Bfas * longitud
print("\nMatriz de susceptancias por fase total para la longitud dada (μS):")
for i, fila in enumerate(Bfas_total):
    print(f" ", end="")
    for B in fila:
        print(f"{B:18.4f}", end="  ")
    print()
# Susceptancias de secuencia
B_seq = A_inv @ Bfas @ A
print("\nMatriz de susceptancias de secuencia (μS):")
for i, fila in enumerate(B_seq):
    print(f" ", end="")
    for B in fila:
        print(f"{B:18.4f}", end="  ")
    print()
# Susceptancia homopolar de la linea (B0)
B0 = B_seq[2, 2]
print(f"\nSusceptancia homopolar de la línea (B0): {B0:.4f} μS/km")
# Susceptancia directa e inversa de la linea (B1)
B1 = B_seq[1, 1]
print(f"Susceptancia directa e inversa de la línea (B1): {B1:.4f} μS/km")
# Susceptancia homopolar de la linea total (B0)
B0_total = B_seq[2, 2] * longitud
print(f"\nSusceptancia homopolar de la línea total (B0): {B0_total:.4f} μS")
# Susceptancia directa e inversa de la linea total (B1)
B1_total = B_seq[1, 1] * longitud
print(f"Susceptancia directa e inversa de la línea total (B1): {B1_total:.4f} μS")

#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# CÁLCULO IMPEDANCIA CARACTERÍSTICA Y CONSTANTE DE PROPAGACIÓN
#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
print("\n6. IMPEDANCIA CARACTERÍSTICA Y CONSTANTE DE PROPAGACIÓN")
# Impedancia característica (Zc=raiz((R+jX)/(jB)))
Zc = cmath.sqrt(Z1 / (1 * B1*1e-6))
print(f"\nZc: {Zc:.4f}")
magnitud_Zc = abs(Zc)
angulo_Zc = math.degrees(math.atan2(Zc.imag, Zc.real))
print(f"\nImpedancia característica de secuencia directa (Zc1): {magnitud_Zc:.4f} ∠ {angulo_Zc:.2f}° Ω")
# Constante de propagación (gamma=raiz((R+jX)*(jB)))
gamma = cmath.sqrt(Z1 * (1 * B1*1e-6))
print(f"\nGamma: {gamma:.6f}")
magnitud_gamma = abs(gamma)
angulo_gamma = math.degrees(math.atan2(gamma.imag, gamma.real))
print(f"\nConstante de propagación de secuencia directa (γ1): {magnitud_gamma:.6f} ∠ {angulo_gamma:.2f}° 1/km")
# Constante de propagación total para la longitud dada
gamma_total = gamma * longitud
magnitud_gamma_total = abs(gamma_total)
angulo_gamma_total = math.degrees(math.atan2(gamma_total.imag, gamma_total.real))
print(f"\nConstante de propagación para la longitud dada: {gamma_total:.4f} = {magnitud_gamma_total:.4f} ∠ {angulo_gamma_total:.2f}°")

#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# POTENCIA CARACTERÍSTICA
#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
print("\n7. POTENCIA CARACTERÍSTICA")
# Pc = tension_nominal**2 / Zc
Pc = (tension_nominal) ** 2 / Zc
magnitud_Pc = abs(Pc)
angulo_Pc = math.degrees(math.atan2(Pc.imag, Pc.real))
print(f"\nPotencia característica: {Pc:.4f} = {magnitud_Pc:.4f} ∠ {angulo_Pc:.2f}° MVA")

#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# CAIDA DE TENSIÓN
#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
print("\n8. CAIDA DE TENSIÓN")
# AU=potencia_transportada(MW)*longitud*(R*cos(phi)+X*sin(phi))/((tension_nominal**2)*10*cos(phi))
potencia_transportada_MW = 160 
Δ𝑈 = potencia_transportada_MW * 1000 * longitud * (Z1.real * cos_phi + Z1.imag * math.sin(math.acos(cos_phi))) / ((tension_nominal ** 2) * 10 * cos_phi)
print(f"\nCaída de tensión Δ𝑈: {Δ𝑈:.4f} %")
if Δ𝑈 < 5:
    print("La caída de tensión es inferior al 5%")
else:
    print("La caída de tensión es superior al 5%")

#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# DENSIDAD MÁXIMA DE CORRIENTE E INTENSIDAD MÁXIMA POR CABLE
#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
print("\n9. DENSIDAD MÁXIMA DE CORRIENTE E INTENSIDAD MÁXIMA POR CABLE")
# Densidades de corriente
secciones = [15, 25, 35, 50, 70, 95, 125, 160, 200, 250, 300, 400, 500, 600]
densidades = [6, 5, 4.55, 4.0, 3.55, 3.2, 2.9, 2.7, 2.5, 2.3, 2.15, 1.95, 1.8, 1.65]  # A/mm2
def densidad_corriente(seccion):
    if seccion < secciones[0] or seccion > secciones[-1]:
        return None  # Fuera de rango
    for i in range(len(secciones) - 1):
        if secciones[i] <= seccion <= secciones[i + 1]:
            # Regla de tres
            densidad = densidades[i] + (densidades[i + 1] - densidades[i]) * (seccion - secciones[i]) / (secciones[i + 1] - secciones[i])
            return densidad
    return None
densidad_maxima = densidad_corriente(seccion)
print(f"\nDensidad máxima de corriente para la sección {seccion} mm²: {densidad_maxima:.4f} A/mm²")
# Coeficiente reductor
composiciones = ["30+7", "6+1", "26+7", "54+7", "45+7"]
coef_reductores = [0.916, 0.937, 0.937, 0.95, 0.97]
def coeficiente_reductor(composicion):
    if composicion in composiciones:
        index = composiciones.index(composicion)
        return coef_reductores[index]
    return None
coef_reductor = coeficiente_reductor(composicion)
print(f"Coeficiente reductor para la composición {composicion}: {coef_reductor:.4f}")

densidad_max_con_reduccion = densidad_maxima * coef_reductor
print(f"Densidad máxima de corriente con coeficiente reductor: {densidad_max_con_reduccion:.4f} A/mm²")
# Intensidad máxima
intensidad_maxima_conductor = densidad_max_con_reduccion * seccion  # A
print(f"Corriente máxima admisible del conductor: {intensidad_maxima_conductor:.2f} A")

#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# POTENCIA MÁXIMA ADMISIBLE POR INTENSIDAD
#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
print("\n10. POTENCIA MÁXIMA ADMISIBLE POR INTENSIDAD")
potencia_maxima_admisible = n_i * (intensidad_maxima_conductor * tension_nominal * raiz(3)) * cos_phi 
print(f"\nPotencia máxima admisible por intensidad: {potencia_maxima_admisible/1000:.2f} MW")
# Comparar potencia máxima admisible con potencia transportada
if potencia_maxima_admisible/1000 >= potencia_transportada:
    print("La potencia máxima admisible por intensidad es mayor que la potencia transportada.")
else:
    print("La potencia máxima admisible por intensidad es menor que la potencia transportada.")

#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# POTENCIA DE TRANSPORTE EN FUNCIÓN DE CONDICIONES METEOROLÓGICAS
#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
print("\n11. CÁLCULO POTENCIA DE TRANSPORTE EN FUNCIÓN DE CONDICIONES METEOROLÓGICAS")

# Calor aportado por radiación solar para invierno y verano Qs = coeficiente_absorcion * radiacion * diametro/1000
print("\n   11.1. CALOR APORTADO POR RADIACIÓN SOLAR")
Qs_invierno = coeficiente_absorcion * radiacion_invierno * diametro/1000 # W/m
Qs_verano = coeficiente_absorcion * radiacion_verano * diametro/1000 # W/m
print(f"\nCalor aportado por radiación solar en invierno: {Qs_invierno:.3f} W/m")
print(f"Calor aportado por radiación solar en verano: {Qs_verano:.3f} W/m")

# Calor cedido por radiación Qr = pi * diametro/1000 * emisividad_conductor * cte_boltzmann * ((Tc+273.15)^4 - (Ta+273.15)^4)
print("\n   11.2. CALOR CEDIDO POR RADIACIÓN SOLAR")
cte_boltzmann = 5.6704e-8 # W/m2K4
Qr_invierno = pi * (diametro/1000) * emisividad_conductor * cte_boltzmann * ((Tc + 273.15)**4 - (temperatura_invierno + 273.15)**4) # W/m
Qr_verano = pi * (diametro/1000) * emisividad_conductor * cte_boltzmann * ((Tc + 273.15)**4 - (temperatura_verano + 273.15)**4) # W/m
print(f"\nCalor cedido por radiación en invierno: {Qr_invierno:.3f} W/m")
print(f"Calor cedido por radiación en verano: {Qr_verano:.3f} W/m")

# Calor cedido por convección Qc = pi * conductividad_termica * (Tc - Ta) * Nu
print("\n   11.3. CALOR CEDIDO POR CONVECCIÓN")

# 1.  Convección natural: Nu = A * (Gr*Pr)^m
print("\n       11.3.1. CONVECCIÓN NATURAL")
# Valores de A y m según Gr*Pr
def coeficientes_conveccion_natural(Gr_Pr):
    if 0.1 <= Gr_Pr < 100:
        return 1.02, 0.148
    elif 100 <= Gr_Pr < 10000:
        return 0.85, 0.188
    elif 10000 <= Gr_Pr < 1e7:
        return 0.48, 0.25
    elif 1e7 <= Gr_Pr < 1e12:
        return 0.125, 0.333
    else:
        return None, None
# Gr = (diametro/1000)^3 * (Tc - Ta) * g / ((Tav + 27.15) * (viscosidad_cinematica)^2)
g = 9.81  # m/s2
# viscosidad_cinematica = 1.32e-5 + 9.5e-8 * Tav
def viscosidad_cinematica(Tav):
    return 1.32e-5 + 9.5e-8 * Tav
viscosidad_cinematica_invierno = viscosidad_cinematica((Tc + temperatura_invierno) / 2)
viscosidad_cinematica_verano = viscosidad_cinematica((Tc + temperatura_verano) / 2)
print(f"\nViscosidad cinemática en invierno: {viscosidad_cinematica_invierno:.8f} m²/s")
print(f"Viscosidad cinemática en verano: {viscosidad_cinematica_verano:.8f} m²/s")
# densidad_relativa_aire = exp(-1.16e-4* h)
def Gr(Ta, viscosidad_cinematica):
    Tav = (Tc + Ta) / 2
    return ((diametro/1000)**3 * (Tc - Ta) * g) / ((Tav + 273.15) * (viscosidad_cinematica**2))
Gr_invierno = Gr(temperatura_invierno, viscosidad_cinematica_invierno)
Gr_verano = Gr(temperatura_verano, viscosidad_cinematica_verano)
print(f"\nNúmero de Grashof en invierno: {Gr_invierno:.2e}")
print(f"Número de Grashof en verano: {Gr_verano:.2e}")
# Pr = calor_esp_aire * viscosidad_dinamica / conductividad_termica_aire
calor_esp_aire = 1005  # J/kgK
# viscosidad_dinamica = (17.239 + 4.625e-2 * Tav - 2.03e-5 * Tav **2)e-6 kg/m s
def viscosidad_dinamica(Tav):
    return (17.239 + 4.625e-2 * Tav - 2.03e-5 * Tav **2) * 1e-6  # kg/m s
viscosidad_dinamica_invierno = viscosidad_dinamica((Tc - temperatura_invierno) / 2)
viscosidad_dinamica_verano = viscosidad_dinamica((Tc - temperatura_verano) / 2)
print(f"\nViscosidad dinámica en invierno: {viscosidad_dinamica_invierno:.6e} kg/m s")
print(f"Viscosidad dinámica en verano: {viscosidad_dinamica_verano:.6e} kg/m s")
# conductividad_termica_aire = 2.368e-2 + 7.23e-5 * Tav - 2.763e-8 * Tav**2 W/Km
def conductividad_termica_aire(Tav):
    return 2.42e-2 + 7.2e-5 * Tav  # W/Km
conductividad_termica_aire_invierno = conductividad_termica_aire((Tc + temperatura_invierno) / 2)
conductividad_termica_aire_verano = conductividad_termica_aire((Tc + temperatura_verano) / 2)
print(f"\nConductividad térmica del aire en invierno: {conductividad_termica_aire_invierno:.6e} W/Km")
print(f"Conductividad térmica del aire en verano: {conductividad_termica_aire_verano:.6e} W/Km")

Pr_invierno = calor_esp_aire * viscosidad_dinamica_invierno / conductividad_termica_aire_invierno
Pr_verano = calor_esp_aire * viscosidad_dinamica_verano / conductividad_termica_aire_verano
print(f"\nNúmero de Prandtl en invierno: {Pr_invierno:.4f}")
print(f"Número de Prandtl en verano: {Pr_verano:.4f}")

Gr_Pr_invierno = Gr_invierno * Pr_invierno
Gr_Pr_verano = Gr_verano * Pr_verano
Nu_invierno = 0
Nu_verano = 0
A_invierno, m_invierno = coeficientes_conveccion_natural(Gr_Pr_invierno)
if A_invierno is not None:
    Nu_invierno = A_invierno * (Gr_Pr_invierno ** m_invierno)
A_verano, m_verano = coeficientes_conveccion_natural(Gr_Pr_verano)
if A_verano is not None:
    Nu_verano = A_verano * (Gr_Pr_verano ** m_verano)
print(f"\nNúmero de Nusselt en invierno: {Nu_invierno:.4f}")
print(f"Número de Nusselt en verano: {Nu_verano:.4f}")

# 2. Convección forzada: Nu = B1 * Re^n
print("\n       11.3.2. CONVECCIÓN FORZADA")
# Re = (diametro/1000) * velocidad_viento / viscosidad_cinematica
# Rugosidad: Rf = diametro_alambre_ext / (2*(diametro-diametro_alambe_ext))
Rf = diametro_alambre_ext / (2 * (diametro - diametro_alambre_ext))
print(f"\nRugosidad Rf: {Rf:.3f}")
def Re(viscosidad_cinematica):
    return (diametro/1000) * velocidad_viento / viscosidad_cinematica
Re_invierno = Re(viscosidad_cinematica_invierno)
Re_verano = Re(viscosidad_cinematica_verano)
print(f"\nNúmero de Reynolds en invierno: {Re_invierno:.2e}")
print(f"Número de Reynolds en verano: {Re_verano:.2e}")

# Valores para parametros B_1 y n dependiendo del rango de Re:
def coeficientes_conveccion_forzada(Re, Rf):
    if 100 < Re < 2.65e3:
        return 0.641, 0.471
    elif Rf <= 0.05 and 2.65e3 <= Re < 5e4:
        return 0.178, 0.633
    elif Rf > 0.05 and 2.65e3 <= Re < 5e4:
        return 0.048, 0.8
    else:
        return None, None
Nu_invierno_forzada = 0
Nu_verano_forzada = 0
B1_invierno, n_invierno = coeficientes_conveccion_forzada(Re_invierno, Rf)
if B1_invierno is not None:
    Nu_invierno_forzada = B1_invierno * (Re_invierno ** n_invierno)
B1_verano, n_verano = coeficientes_conveccion_forzada(Re_verano, Rf)
if B1_verano is not None:
    Nu_verano_forzada = B1_verano * (Re_verano ** n_verano)
print(f"\nNúmero de Nusselt por convección forzada en invierno: {Nu_invierno_forzada:.4f}")
print(f"Número de Nusselt por convección forzada en verano: {Nu_verano_forzada:.4f}")

Nu_45_invierno = Nu_invierno_forzada * (0.42 + 0.58*math.sin(45*pi/180)**0.9)
Nu_45_verano = Nu_verano_forzada * (0.42 + 0.58*math.sin(45*pi/180)**0.9)
print(f"\nNúmero de Nusselt corregido para ángulo de 45° en invierno: {Nu_45_invierno:.4f}")
print(f"Número de Nusselt corregido para ángulo de 45° en verano: {Nu_45_verano:.4f}")
# Qc = pi* conductividad_termica * (Tc-Ta) * Nu con Nu=max(Nu_natural, Nu_45_forzada)
Qc_invierno = pi * conductividad_termica_aire_invierno * (Tc - temperatura_invierno) * max (Nu_invierno, Nu_45_invierno)
Qc_verano = pi * conductividad_termica_aire_verano * (Tc - temperatura_verano) * max (Nu_verano, Nu_45_verano)
print(f"\nCalor cedido por convección en invierno: {Qc_invierno:.3f} W/m")
print(f"Calor cedido por convección en verano: {Qc_verano:.3f} W/m")

# Resultados corriente máxima: I = raiz((Qr+Qc-Qs)/resistencia_ca/1000)
print("\n   11.4. RESULTADOS CORRIENTE MÁXIMA")

I_max_invierno = raiz((Qr_invierno + Qc_invierno - Qs_invierno)/ (resistencia_ca/n_i * 1e-3))
I_max_verano = raiz((Qr_verano + Qc_verano - Qs_verano) / (resistencia_ca/n_i * 1e-3))
print(f"\nCorriente máxima admisible en invierno según condiciones meteorológicas: {I_max_invierno:.2f} A")
print(f"Corriente máxima admisible en verano según condiciones meteorológicas: {I_max_verano:.2f} A")

#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# POTENCIA MÁXIMA DE TRANSPORTE SEGÚN CONDICIONES METEOROLÓGICAS
#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
print("\n   11.5. POTENCIA MÁXIMA DE TRANSPORTE")
potencia_maxima_invierno = (I_max_invierno * tension_nominal * raiz(3)) * cos_phi
potencia_maxima_verano = (I_max_verano * tension_nominal * raiz(3)) * cos_phi
print(f"\nPotencia máxima de transporte en invierno según condiciones meteorológicas: {potencia_maxima_invierno/1000:.2f} MVA")
print(f"Potencia máxima de transporte en verano según condiciones meteorológicas: {potencia_maxima_verano/1000:.2f} MVA")

#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# PÉRDIDAS DE POTENCIA
#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
print("\n12. PÉRDIDAS DE POTENCIA")
perdidas_potencia = ((potencia_transportada_MW * resistencia_ca * longitud) / (tension_nominal**2 * cos_phi**2))*100
print(f"\nPérdidas de potencia en la línea: {perdidas_potencia:.5f} %")
# En valor absoluto
perdidas_potencia_valor = (potencia_transportada_MW * resistencia_ca * longitud) / (tension_nominal**2 * cos_phi**2) * potencia_transportada_MW
print(f"Pérdidas de potencia en la línea en valor absoluto: {perdidas_potencia_valor:.5f} MW")

#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# CORTOCIRCUITO MÁXIMO
#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
print("\n13. CORTOCIRCUITO MÁXIMO")
# Valores de c_conductor, densidad_conductor, cte_conductor, alpha_conductor en función de si el material es Cobre, Aluminio-Acero o Acero:
materiales = ["Cobre", "Aluminio-Acero", "Acero"]
c_conductores = [390, 910, 480]  # J/kgK
densidades_conductores = [8900, 2700, 7850]  # kg/m3
ctes_conductores = [56e6, 34.8e6, 7.25e6]  # 1/ohmio m
alphas_conductores = [0.0039, 0.004, 0.0045]  # 1/ºC
def propiedades_material(material):
    if material in materiales:
        index = materiales.index(material)
        return (c_conductores[index], densidades_conductores[index], ctes_conductores[index], alphas_conductores[index])
    return (None, None, None, None)
c_conductor, densidad_conductor, cte_conductor, alpha_conductor = propiedades_material(material)
print(f"c_conductor, densidad_conductor, cte_conductor, alpha_conductor", c_conductor, densidad_conductor, cte_conductor, alpha_conductor )
# Temperatura máxima recomendada según material
if material == "Acero":
    temperatura_max_recomendada = 300
else:
    temperatura_max_recomendada = 200
# Factor K
multiplicacion = log ((1+alpha_conductor*(temperatura_max_recomendada-20))/(1+alpha_conductor*(Tc-20)))
K = (raiz (cte_conductor*c_conductor*densidad_conductor*multiplicacion/alpha_conductor))/1000000
print(f"Factor K: {K} Araiz(s)/mm²")
# Icc max
Icc_max = K * seccion / raiz(tiempo_accionamiento_proteccion)
print(f"Icc max: {Icc_max} A")

#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# EFECTO CORONA
#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
print("\n14. EFECTO CORONA ")
presion_barometrica = 1/(log(log(76)-altitud_media/18330))*100 
print(f"Presión barométrica {presion_barometrica:.4f} cmHg")
# Factor corrección densidad aire: delta
def delta (temperatura):
    return 3.921*presion_barometrica/(273+temperatura)
delta_invierno = delta (temperatura_invierno)
delta_verano = delta (temperatura_verano)
print(f"Factor corrección densidad aire invierno (δinv) {delta_invierno:.4f}")
print(f"Factor corrección densidad aire verano (δver) {delta_verano:.4f}")

# Tensión crítica dieléctrica
Ep = 30/math.sqrt(2)
def Uc (mt,delta):
    return raiz(3) * mc * mt * Ep * delta * diametro / (20) * log(DMG/(RMG)) * n_i
Uc_invierno = Uc(mt_invierno, delta_invierno)
Uc_verano = Uc(mt_verano,delta_verano)
print(f"Tensión critica invierno {Uc_invierno:.2f} kV")
print(f"Tensión crítica verano {Uc_verano} kV")
if Uc_invierno < tension_nominal:
    print(f"Hay pérdidas por efecto corona en invierno")
    Perdidas_efecto_corona_invierno = 241/delta_invierno * (frecuencia + 25) * raiz (diametro/2/DMG) * (tension_nominal/raiz(3) - Uc_invierno/raiz(3))**2 * 1e-5
    print(f"Pérdidas por efecto corona en invierno: {Perdidas_efecto_corona_invierno:.2f} kW·km/fase")
else: print(f"No hay pérdidas por efecto corona en invierno")
if Uc_verano < tension_nominal:
    print(f"Hay pérdidas por efecto corona en verano")
    Perdidas_efecto_corona_verano = 241/delta_verano * (frecuencia + 25) * raiz (diametro/2/DMG) * (tension_nominal/raiz(3) - Uc_verano/raiz(3))**2 * 1e-5
    print(f"Pérdidas por efecto corona en verano: {Perdidas_efecto_corona_verano:.2f} kW·km/fase")
else: print(f"No hay pérdidas por efecto corona en verano")




sys.stdout = sys.__stdout__
texto_resultados = buffer.getvalue()
html = """
<html>
<head>
<meta charset="utf-8">
<style>
body { font-family: Arial; }
.mayus { font-weight: bold; }
</style>
</head>
<body>
<pre>
"""
def es_mayusculas(linea):
    letras = [c for c in linea if c.isalpha()]
    return len(letras)>0 and all(c.isupper() for c in letras)

for linea in texto_resultados.split("\n"):
    # Detectar líneas completamente en mayúsculas
    if es_mayusculas(linea):
        html += f'<span class="mayus">{linea}</span>\n'
    else:
        html += linea + "\n"

html += """
</pre>
</body>
</html>
"""
with open("resultados.html", "w", encoding="utf-8") as f:
    f.write(html)






from docx import Document
doc = Document()
doc.add_heading("Resultados", 1)
for linea in texto_resultados.split("\n"):
    doc.add_paragraph(linea)
doc.save("resultados.docx")