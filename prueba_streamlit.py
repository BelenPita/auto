# -*- coding: utf-8 -*-
import math
import numpy as np
from math import pi
from math import sqrt as raiz
from math import log
import cmath
import io
import sys
import ast
import os
import unicodedata
import streamlit as st
import pandas as pd




print = st.write

# Crear un buffer para capturar salidas de print
buffer = io.StringIO()

# Coeficientes
mu0 = 4e-7 * pi # H/m
epsilon0 = 8.854e-12  # F/m



#-----------------------------------------------------------------------------------------------------------------------------------------------
# LECTURA DE VALORES EXCEL
#-----------------------------------------------------------------------------------------------------------------------------------------------

def _normalize_key(s):
    if s is None:
        return ""
    s = str(s)
    s = unicodedata.normalize('NFKD', s)
    s = ''.join(c for c in s if not unicodedata.combining(c))
    s = ''.join(ch for ch in s if ch.isalnum())
    return s.lower()

def load_parameters_from_excel(file_name="Cálculos eléctricos.xlsx"):
    # Lee un archivo de Excel con dos columnas (nombre, valor)
    base_dir = os.path.dirname(__file__)
    candidates = [file_name, file_name + ".xlsx", file_name + ".xls"]
    path = None
    for c in candidates:
        p = os.path.join(base_dir, c)
        if os.path.exists(p):
            path = p
            break
    if path is None:
        return {}

    mapping = {
        "tensionnominal": "tension_nominal",
        "tension_nominal": "tension_nominal",
        "tensionmaselevada": "tensionmaselevada",
        "tension_mas_elevada": "tensionmaselevada",
        "ndecircuitos": "ndecircuitos",
        "n_de_circuitos": "ndecircuitos",
        "ndeconductoresporfase": "ndeconductoresporfase",
        "n_de_conductores_por_fase": "ndeconductoresporfase",
        "coeficientetemperatura": "coef_temp",
        "coeficiente_temperatura": "coef_temp",
        "coeficienttemperatura": "coef_temp",
        "frecuencia": "frecuencia",
        "longitud": "longitud",
        "resistividad": "resistividad",
        "cosphi": "cos_phi",
        "cos_phi": "cos_phi",
        "factordepotencia": "cos_phi",
        "factor_de_potencia": "cos_phi",
        "potenciaatransportar": "potencia_transportada",
        "potencia_a_transportar": "potencia_transportada",
        "tiempoaccionamientoproteccion": "tiempo_accionamiento_proteccion",
        "altitudmedia": "altitud_media",
        "tipodeconductor": "conductor",
        "tipo_de_conductor": "conductor",
        "material": "material",
        "temperaturainvierno": "temperatura_invierno",
        "temperaturaverano": "temperatura_verano",
        "temperaturaconductor": "Tc",
        "emisividadconductor": "emisividad_conductor",
        "emisividad_conductor": "emisividad_conductor",
        "coeficienteabsorcion": "coeficiente_absorcion",
        "coeficiente_absorcion": "coeficiente_absorcion",
        "radiacioninvierno": "radiacion_invierno",
        "radiacionverano": "radiacion_verano",
        "velocidadviento": "velocidad_viento",
        "mc": "mc",
        "mtinvierno": "mt_invierno",
        "mtverano": "mt_verano",
        "resistenciatierra": "resistencia_tierra",
        "secciontierra": "seccion_tierra",
        "diametrotierra": "diametro_tierra",
        "coeficientetemperaturatierra": "coef_temp_tierra",
        "coeficiente_temperatura_tierra": "coef_temp_tierra",
        "coeftemptierra": "coef_temp_tierra",
        "uno": "uno",
        "dos": "dos",
        "tres": "tres",
        "cuatro": "cuatro",
        "cinco": "cinco",
        "seis": "seis",
        "tierra": "tierra",
        "distanciaconductores": "distancia_conductores"
    }

    def _parse_value(v):
        if v is None:
            return None
        if isinstance(v, (int, float, complex)):
            return v
        if isinstance(v, str):
            s = v.strip()
            try:
                val = ast.literal_eval(s)
                return val
            except Exception:
                try:
                    if "," in s and "." not in s:
                        s = s.replace(',', '.')
                    return float(s)
                except Exception:
                    return s
        return v

    result = {}
    try:
        import pandas as pd
        df = pd.read_excel(path, header=None, usecols=[0, 1])
        for idx, row in df.iterrows():
            key = row.iloc[0]
            val = row.iloc[1]
            nk = _normalize_key(key)
            if nk in mapping:
                varname = mapping[nk]
                parsed = _parse_value(val)
                result[varname] = parsed
    except Exception:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(path, data_only=True)
            ws = wb.active
            for row in ws.iter_rows(min_row=1, max_col=2, values_only=True):
                key, val = row[0], row[1]
                nk = _normalize_key(key)
                if nk in mapping:
                    varname = mapping[nk]
                    parsed = _parse_value(val)
                    result[varname] = parsed
        except Exception:
            return {}
    return result


# Valores por defecto (sin depender del Excel)
default_params = {
    'tension_nominal': 220.00,
    'tensionmaselevada': 145.00,
    'frecuencia': 50.00,
    'longitud': 0.3131,
    'ndecircuitos': 1,
    'ndeconductoresporfase': 1,
    'distancia_conductores': 400.00,
    'conductor': 'LA-110',
    'material': 'Aluminio-Acero',
    'cos_phi': 0.928,
    'potencia_transportada': 140.00,
    'tiempo_accionamiento_proteccion': 0.50,
    'resistividad': 250.00,
    'altitud_media': 800.00,
    'coef_temp': 0.00403,
    'Tc': 85.00,
    'emisividad_conductor': 0.5,
    'coeficiente_absorcion': 0.5,
    'velocidad_viento': 0.5,
    'seccion_tierra': 155.50,
    'diametro_tierra': 18.00,
    'resistencia_tierra': 0.33,
    'coef_temp_tierra': 0.0000144,
    'temperatura_invierno': 4.00,
    'temperatura_verano': 19.00,
    'radiacion_invierno': 79.00,
    'radiacion_verano': 261.00,
    'mc': 0.83,
    'mt_invierno': 0.80,
    'mt_verano': 1.00,
    'uno_x': 2.300,
    'uno_y': 30.2833333,
    'dos_x': 4.350,
    'dos_y': 24.2266667,
    'tres_x': 0.25,
    'tres_y': 21.5066667,
    'cuatro_x': 0.000,
    'cuatro_y': 38.1166667,
    'cinco_x': 3,
    'cinco_y': 2,
    'seis_x': 3,
    'seis_y': 4,
    'tierra_x': 6,
    'tierra_y': 3,
}

# Inicializar session_state con valores por defecto
for varname, value in default_params.items():
    if varname not in st.session_state:
        st.session_state[varname] = value

# Configurar título
st.set_page_config(layout="wide", page_title="Cálculos Eléctricos")
st.title("Cálculos Eléctricos de Líneas")
st.header("Parámetros de entrada")

# DATOS DE LÍNEA
with st.expander("Datos de la Línea", expanded=True):
    tension_nominal = st.number_input(
        "Tensión nominal (kV)",
        value=st.session_state['tension_nominal'],
        format="%.2f",
        key='tension_nominal'
    )
    tensionmaselevada = st.number_input(
        "Tensión más elevada (kV)",
        value=st.session_state['tensionmaselevada'],
        format="%.2f",
        key='tensionmaselevada'
    )
    frecuencia = st.number_input(
        "Frecuencia (Hz)",
        value=st.session_state['frecuencia'],
        format="%.2f",
        key='frecuencia'
    )
    longitud = st.number_input(
        "Longitud (km)",
        value=st.session_state['longitud'],
        format="%.4f",
        key='longitud'
    )
    ndecircuitos = st.selectbox(
        "Número de circuitos",
        options=[1, 2],
        index=int(st.session_state['ndecircuitos']) - 1,
        key='ndecircuitos'
    )
    ndeconductoresporfase = st.selectbox(
        "Número de conductores por fase",
        options=[1, 2, 3],
        index=int(st.session_state['ndeconductoresporfase']) - 1,
        key='ndeconductoresporfase'
    )
    if ndeconductoresporfase == 2:
        distancia_conductores = st.number_input(
            "Distancia entre conductores (mm)",
            value=st.session_state['distancia_conductores'],
            format="%.2f",
            key='distancia_conductores'
        )
    else:
        distancia_conductores = st.session_state['distancia_conductores']
    conductor = st.selectbox(
        "Tipo de conductor",
        options=["LA-30", "LA-56", "LA-78", "LA-110", "LA-145", "LA-180", "LA-280", "LA-380", "LA-455", "LA-545", "LA-635"],
        index=["LA-30", "LA-56", "LA-78", "LA-110", "LA-145", "LA-180", "LA-280", "LA-380", "LA-455", "LA-545", "LA-635"].index(st.session_state['conductor']),
        key='conductor'
    )
    material = st.selectbox(
        "Material del conductor",
        options=["Cobre", "Aluminio-Acero", "Acero"],
        index=["Cobre", "Aluminio-Acero", "Acero"].index(st.session_state['material']),
        key='material'
    )

# DATOS DE POTENCIA Y PROTECCIÓN
with st.expander("Potencia y Protección", expanded=True):
    cos_phi = st.number_input(
        "Factor de potencia (cos φ)",
        value=st.session_state['cos_phi'],
        min_value=0.0,
        max_value=1.0,
        step=0.01,
        key='cos_phi'
    )
    potencia_transportada = st.number_input(
        "Potencia transportada (MVA)",
        value=st.session_state['potencia_transportada'],
        format="%.2f",
        key='potencia_transportada'
    )
    
    tiempo_accionamiento_proteccion = st.number_input(
        "Tiempo accionamiento protección (s)",
        value=st.session_state['tiempo_accionamiento_proteccion'],
        format="%.2f",
        key='tiempo_accionamiento_proteccion'
    )

with st.expander("Parámetros del Terreno", expanded=False):
    resistividad = st.number_input(
        "Resistividad del terreno (Ω·m)",
        value=st.session_state['resistividad'],
        format="%.2f",
        key='resistividad'
    )
    altitud_media = st.number_input(
        "Altitud media (m)",
        value=st.session_state['altitud_media'],
        format="%.2f",
        key='altitud_media'
    )
    coef_temp = st.number_input(
        "Coeficiente de temperatura (1/ºC)",
        value=st.session_state['coef_temp'],
        format="%.6f",
        key='coef_temp'
    )
with st.expander("Datos Térmicos del Conductor", expanded=False):
    Tc = st.number_input(
        "Temperatura del conductor (ºC)",
        value=st.session_state['Tc'],
        format="%.2f",
        key='Tc'
    )
    emisividad_conductor = st.number_input(
        "Emisividad del conductor",
        value=st.session_state['emisividad_conductor'],
        min_value=0.0,
        max_value=1.0,
        step=0.01,
        key='emisividad_conductor'
    )
    coeficiente_absorcion = st.number_input(
        "Coeficiente de absorción",
        value=st.session_state['coeficiente_absorcion'],
        min_value=0.0,
        max_value=1.0,
        step=0.01,
        key='coeficiente_absorcion'
    )
    velocidad_viento = st.number_input(
        "Velocidad del viento (m/s)",
        value=st.session_state['velocidad_viento'],
        format="%.2f",
        key='velocidad_viento'
    )
with st.expander("Conductor de Tierra", expanded=False):
    seccion_tierra = st.number_input(
        "Sección del conductor tierra (mm²)",
        value=st.session_state['seccion_tierra'],
        format="%.2f",
        key='seccion_tierra'
    )
    diametro_tierra = st.number_input(
        "Diámetro del conductor tierra (mm)",
        value=st.session_state['diametro_tierra'],
        format="%.2f",
        key='diametro_tierra'
    )
    resistencia_tierra = st.number_input(
        "Resistencia del conductor tierra (Ω/km)",
        value=st.session_state['resistencia_tierra'],
        format="%.6f",
        key='resistencia_tierra'
    )
    coef_temp_tierra = st.number_input(
        "Coeficiente temp. tierra (1/ºC)",
        value=st.session_state['coef_temp_tierra'],
        format="%.6e",
        key='coef_temp_tierra'
    )
with st.expander("Condiciones Locales", expanded=False):
    temperatura_invierno = st.number_input(
        "Temperatura invierno (ºC)",
        value=st.session_state['temperatura_invierno'],
        format="%.2f",
        key='temperatura_invierno'
    )
    temperatura_verano = st.number_input(
        "Temperatura verano (ºC)",
        value=st.session_state['temperatura_verano'],
        format="%.2f",
        key='temperatura_verano'
    )
    radiacion_invierno = st.number_input(
        "Radiación invierno (W/m²)",
        value=st.session_state['radiacion_invierno'],
        format="%.2f",
        key='radiacion_invierno'
    )
    radiacion_verano = st.number_input(
        "Radiación verano (W/m²)",
        value=st.session_state['radiacion_verano'],
        format="%.2f",
        key='radiacion_verano'
    )
    mc = st.number_input(
        "mc (Factor de riesgo)",
        value=st.session_state['mc'],
        step=0.01,
        key='mc'
    )
    mt_invierno = st.number_input(
        "mt invierno",
        value=st.session_state['mt_invierno'],
        step=0.01,
        key='mt_invierno'
    )
    mt_verano = st.number_input(
        "mt verano",
        value=st.session_state['mt_verano'],
        step=0.01,
        key='mt_verano'
    )
with st.expander("Ubicación de Conductores", expanded=False):
    uno_x = st.number_input("Punto 1 - X (m)", value=st.session_state['uno_x'], key='uno_x')
    uno_y = st.number_input("Punto 1 - Y (m)", value=st.session_state['uno_y'], key='uno_y')
    uno = (uno_x, uno_y)
    
    dos_x = st.number_input("Punto 2 - X (m)", value=st.session_state['dos_x'], key='dos_x')
    dos_y = st.number_input("Punto 2 - Y (m)", value=st.session_state['dos_y'], key='dos_y')
    dos = (dos_x, dos_y)
    
    tres_x = st.number_input("Punto 3 - X (m)", value=st.session_state['tres_x'], key='tres_x')
    tres_y = st.number_input("Punto 3 - Y (m)", value=st.session_state['tres_y'], key='tres_y')
    tres = (tres_x, tres_y)
    
    cuatro_x = st.number_input("Punto 4 - X (m)", value=st.session_state['cuatro_x'], key='cuatro_x')
    cuatro_y = st.number_input("Punto 4 - Y (m)", value=st.session_state['cuatro_y'], key='cuatro_y')
    cuatro = (cuatro_x, cuatro_y)
    
    if ndecircuitos == 2:
        cinco_x = st.number_input("Punto 5 - X (m)", value=st.session_state['cinco_x'], key='cinco_x')
        cinco_y = st.number_input("Punto 5 - Y (m)", value=st.session_state['cinco_y'], key='cinco_y')
        cinco = (cinco_x, cinco_y)
        
        seis_x = st.number_input("Punto 6 - X (m)", value=st.session_state['seis_x'], key='seis_x')
        seis_y = st.number_input("Punto 6 - Y (m)", value=st.session_state['seis_y'], key='seis_y')
        seis = (seis_x, seis_y)
        
        tierra_x = st.number_input("Punto tierra - X (m)", value=st.session_state['tierra_x'], key='tierra_x')
        tierra_y = st.number_input("Punto tierra - Y (m)", value=st.session_state['tierra_y'], key='tierra_y')
        tierra = (tierra_x, tierra_y)
    else:
        cinco = (0, 0)
        seis = (0, 0)
        tierra_x = st.number_input("Punto tierra - X (m)", value=st.session_state['tierra_x'], key='tierra_x')
        tierra_y = st.number_input("Punto tierra - Y (m)", value=st.session_state['tierra_y'], key='tierra_y')
        tierra = (tierra_x, tierra_y)

# Función para convertir puntos a tuplas
def _xy(pt):
    try:
        if isinstance(pt, (list, tuple)) and len(pt) >= 2:
            return float(pt[0]), float(pt[1])
        if isinstance(pt, str):
            s = pt.strip().replace(';', ',')
            try:
                v = ast.literal_eval(s)
                if isinstance(v, (list, tuple)) and len(v) >= 2:
                    return float(v[0]), float(v[1])
            except Exception:
                parts = s.split(',')
                if len(parts) >= 2:
                    return float(parts[0]), float(parts[1])
        if isinstance(pt, (int, float)):
            return float(pt), 0.0
    except Exception:
        pass
    raise ValueError(f"Formato de punto inválido: {pt}")

# Matriz de puntos para cálculos posteriores
puntos = [uno, dos, tres, cuatro, cinco, seis, tierra]

# Mostrar información del conductor seleccionado
diametro_normativa = {"LA-30": 7.14,"LA-56": 9.45,"LA-78": 11.34,"LA-110": 14.0,"LA-145": 15.75,"LA-180": 17.5,"LA-280": 21.8,"LA-380": 25.38,"LA-455": 27.72,"LA-545": 30.42,"LA-635": 32.85}
resistencia_normativa = {"LA-30": 1.0749, "LA-56": 0.6136, "LA-78": 0.4261,"LA-110": 0.3066,"LA-145": 0.2422,"LA-180": 0.1962,"LA-280": 0.1194,"LA-380": 0.0857,"LA-455": 0.0718,"LA-545": 0.0596,"LA-635": 0.0511}
diametro_alambre_ext_normativa = {"LA-30": 2.38,"LA-56": 3.15,"LA-78": 3.78,"LA-110": 2,"LA-145": 2.25,"LA-180": 2.5,"LA-280": 2.68,"LA-380": 2.82,"LA-455": 3.08,"LA-545": 3.38,"LA-635": 3.65}
seccion_normativa = {"LA-30": 31.1,"LA-56": 54.6,"LA-78": 78.6,"LA-110": 116.2,"LA-145": 147.1,"LA-180": 181.6,"LA-280": 281.1,"LA-380": 381,"LA-455": 454.5,"LA-545": 547.3,"LA-635": 636.6}
composicion_normativa = {"LA-30": "6+1","LA-56": "6+1","LA-78": "6+1","LA-110": "30+7","LA-145": "30+7","LA-180": "30+7","LA-280": "26+7","LA-380": "54+7","LA-455": "54+7","LA-545": "54+7","LA-635": "54+19"}

diametro_calculado = diametro_normativa.get(conductor, None)
resistencia_calculada = resistencia_normativa.get(conductor, None)
diametro_alambre_ext_calculado = diametro_alambre_ext_normativa.get(conductor, None)
seccion_calculada = seccion_normativa.get(conductor, None)
composicion_calculada = composicion_normativa.get(conductor, None)

st.subheader(f"Características del Conductor {conductor}")
col1, col2, col3, col4 = st.columns(4)
with col1:
    st.info(f"**Diámetro**\n{diametro_calculado} mm")
with col2:
    st.info(f"**Resistencia**\n{resistencia_calculada} Ω/km")
with col3:
    st.info(f"**Sección**\n{seccion_calculada} mm²")
with col4:
    st.info(f"**Composición**\n{composicion_calculada}")

st.write("---")

# Configurar variables para cálculos
diametro = diametro_calculado
seccion = seccion_calculada
diametro_alambre_ext = diametro_alambre_ext_calculado
composicion = composicion_calculada
resistencia = resistencia_calculada
n_i = ndeconductoresporfase


















# Botón para ejecutar cálculos
if st.button("🔄 Ejecutar Cálculos", use_container_width=True, type="primary"):
    st.session_state.ejecutar_calculos = True
else:
    if 'ejecutar_calculos' not in st.session_state:
        st.session_state.ejecutar_calculos = False

if not st.session_state.ejecutar_calculos:
    st.info("")
    st.stop()

# Si llegamos aquí, ejecutar los cálculos
sys.stdout = buffer



st.markdown("""
<style>
/* Reducir el tamaño de los headers y subheaders */
h1 { font-size: 1.5em !important; }  /* st.title() */
h2 { font-size: 1.3em !important; }  /* st.header() */
h3 { font-size: 1.1em !important; }  /* st.subheader() */

/* Opcional: ajustar st.latex() */
.mathjax, .katex, .katex-display {
    font-size: 0.95em !important;
}
</style>
""", unsafe_allow_html=True)


#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# SIMPLE CIRCUITO SIMPLEX
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
if ndecircuitos==1 and ndeconductoresporfase==1:      
        
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # CÁLCULO RESISTENCIA
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # Definir función para cálculo de resistencia a temperatura
    def resistencia_a_temp(resistencia_20C, temp):
        return resistencia_20C * (1 + coef_temp * (temp - 20))
    # Cálculo de la resistencia a 85ºC y en ca
    st.header("\n1. RESISTENCIA ELÉCTRICA DE LA LÍNEA")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )
    st.markdown(f"""La resistencia eléctrica del conductor {conductor}, por unidad de longitud, en corriente continua y a la temperatura
                máxima del conductor ({Tc:.0f}ºC), vendrá dada por la siguiente expresión:""")
    st.latex(fr"""R_{{{Tc:.0f}}} = R_{{20cc}} \cdot (1 + \alpha \cdot ({{{Tc:.0f}}} - 20))""")
    st.markdown(f"""
                Donde:
                - $R_{{{Tc:.0f}}}$ es la resistencia del conductor en corriente continua a {Tc:.0f} ºC
                - $R_{{20cc}}$ es la resistencia del conductor en corriente continua a 20ºC
                - $\\alpha$ es el coeficiente de aumento de la resistividad eléctrica con la temperatura
                """)


                
    resistencia_85C = resistencia_a_temp(resistencia, Tc)
    print(f"Resistencia a {Tc:.2f}ºC: {resistencia_85C:.6f} Ω/km")
    reactancia_pelicular_85C = 10**-7 * 8 * pi * frecuencia * (1 / (resistencia_a_temp(resistencia, Tc) / 1000))
    print(f"Reactancia por efecto pelicular a {Tc:.2f}ºC: {reactancia_pelicular_85C:.6f} Ω/km")
    ys = reactancia_pelicular_85C ** 2 / (192 + 0.8*reactancia_pelicular_85C**2)
    print(f"Factor de efecto pelicular ys: {ys:.6f}")
    resistencia_ca = resistencia_85C * (1 + ys)
    print(f"Resistencia en CA a 85ºC: {resistencia_ca:.6f} Ω/km")
    r_ca_longitud = resistencia_ca * longitud
    print(f"Resistencia en CA para {longitud} km: {r_ca_longitud:.6f} Ω")

    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # CÁLCULO MATRIZ IMPEDANCIAS
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.header("\n2. MATRIZ DE IMPEDANCIAS")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )    
    st.write("""
    El cálculo de la matriz de impedancias se realiza mediante la teoría de Carson. 
    De este modo se calcula la impedancia propia y mutua de todos los conductores 
    que forman la línea teniendo en cuenta el terreno.
    """)
    st.write("Impedancia propia de un conductor:")
    st.latex(r"""
    Z_{i,i} = R_i
    + \frac{j \mu_0}{2 \pi} \, \omega 
    \left(
    \frac{\mu_r}{4 n_i'} + 
    \ln \left( \frac{D'_{i,i}}{r_{eq,i}} \right)
    \right)
    + \frac{\mu_0}{\pi} \, \omega \left( P_{i,i} + j Q_{i,i} \right)
    """)
    st.write("Impedancia mutua entre dos conductores:")
    st.latex(r"""
    Z_{i,j} = 
    \frac{j \mu_0}{2 \pi} \, \omega 
    \ln \left( \frac{D'_{i,j}}{D_{i,j}} \right)
    + \frac{\mu_0}{\pi} \, \omega \left( P_{i,j} + j Q_{i,j} \right)
    """)
    st.write("""
    Donde:
    - $R_i$ es la resistencia eléctrica del conductor $i$  
    - $D_{i,j}$ es la distancia entre el conductor $i$ y el $j$  
    - $D'_{i,j}$ es la distancia entre el conductor $i$ y la imagen del conductor $j$ respecto al suelo  
    - $r_{eq,i}$ es el radio equivalente del conductor $i$  
    - $\omega$ es la pulsación $2 \pi f$  
    - $\mu_0$ es la permeabilidad del vacío  
    - $\mu_r$ es la permeabilidad relativa del conductor  

    Los valores de $P$ y $Q$ son una serie infinita de términos que describen 
    el comportamiento del terreno. Para cálculos a frecuencias industriales, los 
    términos de la serie a partir del segundo pueden ser despreciados. Por tanto:
    """)
    st.latex(r"""
    P_{i,j} = \frac{\pi}{8} - \frac{k_{i,j} \cdot \cos(\theta_{i,j})}{3 \sqrt{2}}
    """)
    st.latex(r"""
    Q_{i,j} = \frac{1}{2} \cdot \ln \left( \frac{1.85138}{k_{i,j}} \right)
    + \frac{k_{i,j} \cdot \cos(\theta_{i,j})}{3 \sqrt{2}}
    """)
    
    st.write("Donde:")
    st.latex(r"""
    k_{i,j} = \frac{\sqrt{2}\, D'_{i,j}}{\delta}
    """)
    st.latex(r"""
    \cos(\theta_{i,j}) = \frac{h_i + h_j}{D'_{i,j}}
    """)
    st.write("""
    Siendo $\\delta$ la penetración del terreno según la ecuación:
    """)
    st.latex(r"""
    \delta =
    \sqrt{
    \frac{\rho_t}
    {\pi \, f \, \mu_0 \, \mu_r}
    }
    """)
    st.markdown("""
    Donde:
    - $\\rho_t$ es la resistividad del terreno en $\\Omega\,m$  
    - $f$ es la frecuencia en Hz  
    - $\\mu_0$ es la permeabilidad del vacío  
    - $\\mu_r$ es la permeabilidad relativa del terreno (generalmente 1)
    """)
    penetracion_terreno = raiz(resistividad/(pi*frecuencia*mu0))
    st.latex(r"\delta = \sqrt{\frac{" + f"{resistividad:.0f}" + r"}{\pi \cdot " + f"{frecuencia:.0f}" +r" \cdot 4 \cdot \pi \cdot 10^{-7} \cdot 1}} = " + f"{penetracion_terreno:.4f}" + " \, m")
    # Matriz de distancias
    puntos = [uno, dos, tres, cuatro]
    n_puntos = len(puntos)
    matriz_distancias = [[0]*n_puntos for _ in range(n_puntos)]
    def _xy(pt):
        try:
            if isinstance(pt, (list, tuple)) and len(pt) >= 2:
                return float(pt[0]), float(pt[1])
            if isinstance(pt, str):
                s = pt.strip().replace(';', ',')
                try:
                    v = ast.literal_eval(s)
                    if isinstance(v, (list, tuple)) and len(v) >= 2:
                        return float(v[0]), float(v[1])
                except Exception:
                    parts = s.split(',')
                    if len(parts) >= 2:
                        return float(parts[0]), float(parts[1])
            if isinstance(pt, (int, float)):
                return float(pt), 0.0
        except Exception:
            pass
        raise ValueError(f"Formato de punto inválido: {pt}")
    for i in range(n_puntos):
        for j in range(n_puntos):
            x1, y1 = _xy(puntos[i])
            x2, y2 = _xy(puntos[j])
            distancia = raiz((x2-x1)**2 + (y2-y1)**2)
            matriz_distancias[i][j] = distancia
    st.markdown("""Con el promedio de posiciones de los conductores, se calculan tanto la distancia entre los cables ocmo la distancia espejo respecto al suelo.""")
    print("\nLas distancias entre todos los conductores de la línea son:")
    df=pd.DataFrame(matriz_distancias).round(4)
    latex_matrix = r"D = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\text{m}"
    st.latex(latex_matrix)
    # Matriz D_prima: distancia entre puntos y sus espejos respecto al suelo
    matriz_D_prima = [[0]*n_puntos for _ in range(n_puntos)]
    for i in range(n_puntos):
        for j in range(n_puntos):
            x1, y1 = _xy(puntos[i])
            # Espejo del punto j respecto al suelo
            x2, y2 = _xy(puntos[j])
            x2_espejo, y2_espejo = x2, -y2
            distancia = raiz((x2_espejo-x1)**2 + (y2_espejo-y1)**2)
            matriz_D_prima[i][j] = distancia
    print("\nLas distancias entre conductores y el espejo respecto al suelo de los mismos para la línea es:")
    df=pd.DataFrame(matriz_D_prima).round(4)
    latex_matrix = r"D' = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\text{m}"
    st.latex(latex_matrix)
    # kij = raiz(2)*D_prima_ij / penetracion_terreno
    matriz_kij = [[0]*n_puntos for _ in range(n_puntos)]
    for i in range(n_puntos):
        for j in range(n_puntos):
            matriz_kij[i][j] = raiz(2) * matriz_D_prima[i][j] / penetracion_terreno  
    st.write("\nEl factor k para los conductores será:")
    df=pd.DataFrame(matriz_kij).round(4)
    latex_matrix = r"k_{ij} = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    # P_ij = pi/8 - k_ij*cos(tetha_ij)/(3*raiz(2)) siendo cos(tetha_ij) = (h_i + h_j )/ D_prima_ij
    matriz_Pij = [[0]*n_puntos for _ in range(n_puntos)]
    for i in range(n_puntos):
        for j in range(n_puntos):
            h_i = _xy(puntos[i])[1]
            h_j = _xy(puntos[j])[1]
            D_prima_ij = matriz_D_prima[i][j]
            if D_prima_ij != 0:
                cos_tetha_ij = (h_i + h_j) / D_prima_ij
            else:
                cos_tetha_ij = 0
            k_ij = matriz_kij[i][j]
            P_ij = (pi / 8) - (k_ij * cos_tetha_ij) / (3 * raiz(2))
            matriz_Pij[i][j] = P_ij
    print("\nEl factor P, que representa la parte real dela impedancia del terreno, según la teoría de Carson es:")
    df=pd.DataFrame(matriz_Pij).round(4)
    latex_matrix = r"P_{ij} = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    # Q_ij = 0.5*log neperiano(1.85138/k_ij) + k_ij*cos(tetha_ij)/(3*raiz(2))
    matriz_Qij = [[0]*n_puntos for _ in range(n_puntos)]
    for i in range(n_puntos):
        for j in range(n_puntos):
            k_ij = matriz_kij[i][j]
            h_i = _xy(puntos[i])[1]
            h_j = _xy(puntos[j])[1]
            D_prima_ij = matriz_D_prima[i][j]
            if D_prima_ij != 0:
                cos_tetha_ij = (h_i + h_j) / D_prima_ij
            else:
                cos_tetha_ij = 0
            if k_ij != 0:
                Q_ij = 0.5 * log(1.85138 / k_ij) + (k_ij * cos_tetha_ij) / (3 * raiz(2))           
            else:
                Q_ij = 0
            matriz_Qij[i][j] = Q_ij
    print("\nY el factor Q, que es la parte imaginaria de la impedancia que representa el terreno:")
    df=pd.DataFrame(matriz_Qij).round(4)
    latex_matrix = r"Q_{ij} = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    # Resistencia de cada uno de los conductores (ultimo a tierra con dato de resistencia a tierra directamente)
    resistencias_conductores = [resistencia_ca, resistencia_ca, resistencia_ca, resistencia_tierra]
    print("\nLa resistencia de cada uno de los conductores considerados es:")
    df=pd.DataFrame(resistencias_conductores).round(4)
    latex_matrix = r"R = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\text{Ω/km}"
    st.latex(latex_matrix)
    # Radio de cada conductor en mm
    radio_conductores_mm = [
        diametro / 2,
        diametro / 2,
        diametro / 2,
        diametro_tierra / 2
    ]
    print("\nEl radio equivalente de cada uno de los conductores es:")
    df=pd.DataFrame(radio_conductores_mm).round(4)
    latex_matrix = r"Radio = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\text{mm}"
    st.latex(latex_matrix)
    # Matriz de impedancias con:
    # Z_ii = R_i + j*mu0*frecuencia*((1/(4*n_i)+log(D_prima_ii/r_i)))+mu0*2*frecuencia*(P_ii + j*Q_ii) 
    # Z_ij = j*mu0*frecuencia*(log(D_prima_ij/D_ij)) + mu0*2*frecuencia*(P_ij + j*Q_ij) para i != j
    matriz_impedancias = [[0]*n_puntos for _ in range(n_puntos)]
    for i in range(n_puntos):
        for j in range(n_puntos):
            R_i = resistencias_conductores[i]
            D_ij = matriz_distancias[i][j]
            D_prima_ij = matriz_D_prima[i][j]
            P_ij = matriz_Pij[i][j]
            Q_ij = matriz_Qij[i][j]
            r_i = radio_conductores_mm[i] / 1000  # Convertir mm a m
            if i == j:
                n_i = 1  # Número de conductores en fase
                Z_ii = R_i + (1j * mu0 * frecuencia * ((1 / (4 * n_i)) + log(D_prima_ij / r_i)) + mu0 * 2 * frecuencia * (P_ij + 1j * Q_ij))*1000
                matriz_impedancias[i][j] = Z_ii
            else:
                Z_ij = (1j * mu0 * frecuencia * (log(D_prima_ij / D_ij)) + mu0 * 2 * frecuencia * (P_ij + 1j * Q_ij))*1000
                matriz_impedancias[i][j] = Z_ij
    print("\nPor tanto, la matriz de impedancias de fase de la línea, incluyendo cables de tierra, será:(Ω/km):")
    df=pd.DataFrame(matriz_impedancias).round(4)
    latex_matrix = r"Z = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\text{Ω/km}"
    st.latex(latex_matrix)
    # Como la linea posee cable de tierra, es necesario realizar un análisis matricial para eliminarlos y obtener una matriz 3*3 que representa las impedancias por fase
    # Z=[Zf, Zft; Ztf, Zt]
    # Zfas = Zf - Zft * inv(Zt) * Ztf
    Z = np.array(matriz_impedancias)
    Zf = Z[0:3, 0:3]
    Zt = Z[3:4, 3:4]
    Zft = Z[0:3, 3:4]
    Ztf = Z[3:4, 0:3]
    Zt_inv = np.linalg.inv(Zt)
    Zfas = Zf - Zft @ Zt_inv @ Ztf
    print("\nComo la línea posee cable de tierra, es necesario realizar un análisis matricial para eliminarlos y obtener una matriz 3x3 que representa las impedancias por fase:")
    st.latex(r"""
    Z =
    \begin{pmatrix}
    Z_f & Z_{ft} \\
    Z_{tf} & Z_{tt}
    \end{pmatrix}
    """)
    st.latex(r"""
    Z_{fas} =
    Z_f
    -
    Z_{ft}
    Z_{tt}^{-1}
    Z_{tf}
    """)
    st.markdown("""El resultado del cálculo es la matriz de impedancias de fase:""")
    df=pd.DataFrame(Zfas).round(4)
    latex_matrix = r"Z_{fas} = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\text{Ω/km}"
    st.latex(latex_matrix)
    # Finalmente, multiplicar por la longitud de la línea para obtener las impedancias totales
    Zfas_total = Zfas * longitud
    print(f"\nCon la longitud total media de la línea de {longitud} km:")
    df=pd.DataFrame(Zfas_total).round(4)
    latex_matrix = r"Z_{fas} = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\text{Ω}"
    st.latex(latex_matrix)
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # IMPEDANCIAS DE SECUENCIA
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.header("\n3. IMPEDANCIAS DE SECUENCIA")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )   
    st.markdown("""La matriz de impedancias de secuencia será el resultado de realizar la siguiente operación matricial:""")
    st.latex(r"""
    Z_{012} = H^{-1} \, Z_{fas} \, H
    """)
    st.latex(r"""h = e^{j120^\circ} = -0.5 + j \frac{\sqrt{3}}{2}""")
    st.latex(r"""
    H =
    \begin{pmatrix}
    1 & 1 & 1 \\
    1 & h^2 & h \\
    1 & h & h^2
    \end{pmatrix}
    """)
    st.latex(r"""
    H^{-1} =
    \frac{1}{3}
    \begin{pmatrix}
    1 & 1 & 1 \\
    1 & h & h^2 \\
    1 & h^2 & h
    \end{pmatrix}
    """)
    A = np.array([[1, 1, 1],
                [complex(-0.5, -raiz(3)/2), complex(-0.5, raiz(3)/2), 1],
                [complex(-0.5, raiz(3)/2), complex(-0.5, -raiz(3)/2), 1]])
    A_inv = np.linalg.inv(A)
    Z_seq = A_inv @ Zfas @ A
    print("\nEl resultado del cálculo es:")
    df=pd.DataFrame(Z_seq).round(4)
    latex_matrix = r"Z_{012} = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\text{Ω/km}"
    st.latex(latex_matrix)
    # Matriz anterior de impedancias de secuencia pero con argumento y angulo
    latex_matrix = r"Z_{012} = \begin{pmatrix}"
    for fila in Z_seq:
        elementos = []
        for Z in fila:
            magnitud = abs(Z)
            angulo = math.degrees(math.atan2(Z.imag, Z.real))
            elementos.append(
                rf"{magnitud:.4f} \angle {angulo:.2f}^\circ"
            )
        
        latex_matrix += " & ".join(elementos)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\text{Ω/km}"
    st.latex(latex_matrix)
    # Impedancia homopolar de la linea (Z0)
    Z0 = Z_seq[2, 2]
    print(f"\nImpedancia homopolar de la línea")
    st.latex(r"Z_0 = " + f"{Z0:.4f}" + r" \, \Omega / km")
    # Impedancia directa e inversa de la linea (Z1)
    Z1 = Z_seq[1, 1]
    print(f"Impedancia directa e inversa de la línea:")
    st.latex(r"Z_1 = " + f"{Z1:.4f}" + r" \, \Omega / km")
    #Teniendo en cuenta la longitud de la línea
    Z_seq_total = Z_seq * longitud
    print(f"\nCon la longitud de la línea de {longitud} km:")
    df=pd.DataFrame(Z_seq_total).round(4)
    latex_matrix = r"Z_{012} = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\text{Ω}"
    st.latex(latex_matrix)
    # Matriz anterior de impedancias de secuencia pero con argumento y angulo
    latex_matrix = r"Z_{012} = \begin{pmatrix}"
    for fila in Z_seq_total:
        elementos = []
        for Z in fila:
            magnitud = abs(Z)
            angulo = math.degrees(math.atan2(Z.imag, Z.real))
            elementos.append(
                rf"{magnitud:.4f} \angle {angulo:.2f}^\circ"
            )
        latex_matrix += " & ".join(elementos)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\text{Ω}"
    st.latex(latex_matrix)
    Z0_total = Z_seq_total[2, 2]
    st.latex(r"Z_0 = " + f"{Z0_total:.4f}" + r" \, \Omega")
    Z1_total = Z_seq_total[1, 1]
    st.latex(r"Z_1 = " + f"{Z1_total:.4f}" + r" \, \Omega")
    # Resistencias R0/R1
    R0 = Z0.real
    R1 = Z1.real
    # Reactancias X0/X1
    X0 = Z0.imag
    X1 = Z1.imag
    st.latex(r"\frac{R_0}{R_1} = \frac{" + f"{R0:.4f}" + r"}{" + f"{R1:.4f}" + r"} = " + f"{R0/R1:.3f}")
    st.latex(r"\frac{X_0}{X_1} = \frac{" + f"{X0:.4f}" + r"}{" + f"{X1:.4f}" + r"} = " + f"{X0/X1:.3f}")


    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # MATRIZ DE CAPACIDADES
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.header("\n4. MATRIZ DE CAPACIDADES")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )   
    st.write("""
    El cálculo de la matriz de capacidades de la línea se realiza mediante la matriz de coeficientes de potencial que liga el potencial eléctrico con la carga. La ecuación general en forma matricial es:
    """)
    st.latex(r"V = P \, q")
    st.markdown("""
    Donde:
    - $V$ es el vector de potenciales de la línea  
    - $P$ es la matriz de coeficientes de potencial  
    - $q$ es el vector de cargas  
    """)
    st.latex(r"q = C \, V = P^{-1} \, V")
    st.markdown("""
    Siendo $C$ la matriz de capacidades de la línea. Por tanto, la susceptancia capacitiva de la línea será:
    """)
    st.latex(r"B = \omega \, P^{-1}")
    st.markdown("Los elementos de la matriz $P$ de coeficientes de capacidades se pueden calcular:")
    st.write("Para elementos de la diagonal principal:")
    st.latex(r"P_{i,i} = \frac{1}{2 \pi \epsilon_0} \ln \left( \frac{D_{i,i}'}{r_\mathrm{eq}} \right)")
    st.write("Para elementos fuera de la diagonal:")
    st.latex(r"P_{i,j} = \frac{1}{2 \pi \epsilon_0} \ln \left( \frac{D_{i,i}'}{D_{i,j}} \right), \quad i \neq j")
    st.markdown("""
    Donde:  
    - $D_{i,j}$ es la distancia entre el conductor i y el j  
    - $D'_{i,j}$ es la distancia entre el conductor i y la imagen del conductor j respecto al suelo  
    - $r_\mathrm{eq}$ es el radio equivalente del conductor  
    - $\\epsilon_0$ es la permitividad eléctrica del vacío  

    Las matrices de distancias entre los conductores y entre los conductores y el espejo de los mismos respecto al suelo se han evaluado en un apartado anterior.
    """)

    st.write("La matriz de coeficientes de potencial de fase de la línea incluyendo cables de tierra será:")
    # Coeficientes de potencial
    # C_ii = (1/(2*pi*epsilon0))*log(D_prima_ii/r_i)
    # C_ij = (1/(2*pi*epsilon0))*log(D_prima_ij/D_ij) para i != j
    matriz_capacidades = [[0]*n_puntos for _ in range(n_puntos)]
    for i in range(n_puntos):
        for j in range(n_puntos):
            D_ij = matriz_distancias[i][j]/1000   # Convertir m a km
            D_prima_ij = matriz_D_prima[i][j]/1000   # Convertir m a km
            r_i = radio_conductores_mm[i] / 1000000  # Convertir mm a km
            if i == j:
                C_ii = ((1 / (2 * pi * epsilon0)) * log(D_prima_ij / r_i))/1000000000
                matriz_capacidades[i][j] = C_ii
            else:
                C_ij = (1 / (2 * pi * epsilon0)) * log(D_prima_ij / D_ij)/1000000000
                matriz_capacidades[i][j] = C_ij
    df=pd.DataFrame(matriz_capacidades).round(4)
    latex_matrix = r"P = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\;\text{km}/\mu\text{F}"
    st.latex(latex_matrix)
    # Matriz de capacitancias por fase (eliminando tierra)
    st.write("""
    Como la línea posee cables de tierra es necesario realizar un análisis matricial 
    para eliminarlos y obtener una matriz 3x3 que representa los coeficientes 
    de potencial de cada fase.
    """)
    st.latex(r"""
    P =
    \begin{pmatrix}
    P_f & P_{ft} \\
    P_{tf} & P_{tt}
    \end{pmatrix}
    """)
    st.write("""
    Aplicando la eliminación matricial (reducción de Kron), la matriz equivalente 
    de coeficientes de potencial de fase viene dada por:
    """)
    # Reducción de Kron
    st.latex(r"""
    P_{fas} =
    P_f
    -
    P_{ft}
    P_{tt}^{-1}
    P_{tf}
    """)
    st.write("""
    El resultado del cálculo es la matriz reducida de coeficientes de potencial de fase es:
    """)
    C = np.array(matriz_capacidades)
    Cf = C[0:3, 0:3]
    Ct = C[3:4, 3:4]
    Cft = C[0:3, 3:4]
    Ctf = C[3:4, 0:3]
    Ct_inv = np.linalg.inv(Ct)
    Cfas = Cf - Cft @ Ct_inv @ Ctf
    df=pd.DataFrame(Cfas).round(4)
    latex_matrix = r"P_{fas} = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\;\text{km}/\mu\text{F}"
    st.latex(latex_matrix)

    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # MATRIZ DE SUSCEPTANCIAS
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.header("\n5. MATRIZ DE SUSCEPTANCIAS")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )   
    st.write("""La matriz de susceptancias de la línea será el producto de la pulsación por la inversa de la matriz de coeficientes de potencial:""")
    st.latex(r"B = jwP^{-1}")
    # Matriz de susceptancias por km (B=j*2*pi*frecuencia*C^-1) solo parte imaginaria, no mostrando la parte real
    Cfas_inv = np.linalg.inv(Cfas)
    Bfas = 1j * 2 * pi * frecuencia * Cfas_inv
    df=pd.DataFrame(Bfas).round(4)
    latex_matrix = r"B = j\begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v.imag:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\;\mu\text{S}/\text{km}"
    st.latex(latex_matrix)

    # Matriz de susceptancias total para la longitud dada
    Bfas_total = Bfas * longitud
    st.markdown(f"""Con la longitud de la línea de {longitud} km, se obtiene:""")
    df=pd.DataFrame(Bfas_total).round(4)
    latex_matrix = r"B = j\begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v.imag:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\;\mu\text{S}"
    st.latex(latex_matrix)
    # Susceptancias de secuencia
    st.write("""
    La matriz de susceptancia de secuencia será el resultado de realizar la siguiente operación matricial:
    """)
    st.latex(r"""
    B_{012} = H^{-1} \, B_{fas} \, H
    """)
    st.write("""
    La matriz $H$ se definió en el apartado anterior:
    """)
    st.latex(r"""
    H =
    \begin{pmatrix}
    1 & 1 & 1 \\
    1 & h^2 & h \\
    1 & h & h^2
    \end{pmatrix}
    """)
    st.latex(r"""
    H^{-1} =
    \frac{1}{3}
    \begin{pmatrix}
    1 & 1 & 1 \\
    1 & h & h^2 \\
    1 & h^2 & h
    \end{pmatrix}
    """)
    st.write("""
    El resultado del cálculo es:
    """)
    B_seq = A_inv @ Bfas @ A
    df=pd.DataFrame(B_seq).round(4)
    latex_matrix = r"B_{012} = j\begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v.imag:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    latex_matrix += r"\;\mu\text{S}"

    st.latex(latex_matrix)
    B0 = B_seq[2, 2]
    print(f"\nLa susceptancia homopolar de la línea es:")
    st.latex(
        rf"B_0 = {B0.imag:.4f}j\, \mu S/km"
    )
    B1 = B_seq[1, 1]
    print(f"La susceptancia directa e inversa de la línea será:")
    st.latex(
        rf"B_1 = {B1.imag:.4f}j\, \, \mu S/km"
    )
    B0_total = B_seq[2, 2] * longitud
    print(f"\nCon la longistu de la línea de {longitud} km: ")
    B1_total = B_seq[1, 1] * longitud
    st.latex(
        rf"B_{0} = {B0_total.imag:.4f}j\,\mu S"
    )
    st.latex(
        rf"B_{1} = {B1_total.imag:.4f}j\,\mu S"
    )
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # CÁLCULO IMPEDANCIA CARACTERÍSTICA Y CONSTANTE DE PROPAGACIÓN
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.header("\n6. IMPEDANCIA CARACTERÍSTICA Y CONSTANTE DE PROPAGACIÓN")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )   
    st.write("""Se define la impedancia característica de la línea como:""")
    st.latex(r"Z_c = \sqrt{\frac{R+jX}{jB}}")
    # Impedancia característica (Zc=raiz((R+jX)/(jB)))
    Zc = cmath.sqrt(Z1 / (1 * B1*1e-6))
    magnitud_Zc = abs(Zc)
    angulo_Zc = math.degrees(math.atan2(Zc.imag, Zc.real))
    st.write("""Los valores de impedancia y admitancia de secuancia directa en la línea se han calculado previamente, por tanto se obtiene:""")
    # Ahora volver a escribir la formula de Zc y luego con los valores de Z1 y B1 mas los resultados
    st.latex(
        rf"Z_c = \sqrt{{\frac{{{Z1:.4f}}}{{{B1.imag:.4f}j}}}}"
        rf"={Zc:.4f}"
        rf"={magnitud_Zc:.4f}\angle {angulo_Zc:.2f}^\circ \,\Omega"
    )
    # Constante de propagación (gamma=raiz((R+jX)*(jB)))
    gamma = cmath.sqrt(Z1 * (1 * B1*1e-6))
    magnitud_gamma = abs(gamma)
    angulo_gamma = math.degrees(math.atan2(gamma.imag, gamma.real))

    # Constante de propagación total para la longitud dada
    gamma_total = gamma * longitud
    magnitud_gamma_total = abs(gamma_total)
    angulo_gamma_total = math.degrees(math.atan2(gamma_total.imag, gamma_total.real))

    st.write("""La constante de propagación será:""")
    st.latex(
        rf"\gamma = \sqrt{{(R+jX)(jB)}}"
        rf"=\sqrt{{({Z1.real:.4f}+{Z1.imag:.4f}j)({B1.imag:.4f}j)}}"
        rf"={gamma:.6f}"
        rf"={magnitud_gamma:.6f}\angle {angulo_gamma:.2f}^\circ \,1/km"
    )
    st.write("""\nDada la longitud de la línea:""")
    st.latex(rf"\gamma \cdot L = {gamma_total:.6f} = {magnitud_gamma_total:.6f}\angle {angulo_gamma_total:.2f}^\circ")
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # POTENCIA CARACTERÍSTICA
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.header("\n7. POTENCIA CARACTERÍSTICA")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )   
    print("\nLa potencia característica de la línea es función de la tensión y de la impedancia característica a través de la siguiente expresión:")
    st.latex(r"P_c=\frac{U^2}{Z_c}")
    # Pc = tension_nominal**2 / Zc
    Pc = (tension_nominal) ** 2 / Zc
    magnitud_Pc = abs(Pc)
    angulo_Pc = math.degrees(math.atan2(Pc.imag, Pc.real))
    st.write("\nPara los valores de la línea:")
    st.latex(
        rf"P_c=\frac{{U^2}}{{Z_c}}"
        rf"=\frac{{{tension_nominal}^2}}{{{Zc:.4f}}}"
        rf"={Pc:.4f}"
        rf"={magnitud_Pc:.4f}\angle {angulo_Pc:.2f}^\circ \,\text{{MVA}}"
    )

    #-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # CAIDA DE TENSIÓN
    #-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.header("\n8. CAIDA DE TENSIÓN")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )   
    print("\nLa caída de tensión por resistencia y reactancia de una línea (despreciendo la influencia de la capacidad) viene dada por las fórmulas:")
    st.latex(r"\Delta U = \frac{P \cdot L \cdot (R \cdot \cos(\varphi) + X \cdot \sin(\varphi))}{U^2 \cdot 10 \cdot \cos(\varphi)}")
    # AU=potencia_transportada(MW)*longitud*(R*cos(phi)+X*sin(phi))/((tension_nominal**2)*10*cos(phi))
    potencia_transportada_MW = potencia_transportada * cos_phi  # Convertir MVA a MW usando el factor de potencia
    ΔU = potencia_transportada_MW * 1000 * longitud * (Z1.real * cos_phi + Z1.imag * math.sin(math.acos(cos_phi))) / ((tension_nominal ** 2) * 10 * cos_phi)
    st.latex(
        rf"\Delta U = \frac{{{potencia_transportada_MW:.2f} \cdot {longitud} \cdot ({Z1.real:.4f} \cdot {cos_phi} + {Z1.imag:.4f} \cdot {math.sin(math.acos(cos_phi)):.3f})}}{{{tension_nominal}^2 \cdot 10 \cdot {cos_phi}}}"
        rf"={ΔU:.4f} %"
    )
    print(f"\nCaída de tensión ΔU: {ΔU:.4f} %")
    # Poner si la caida de tension es inferior al 5%: La caída de tensión es inferior al 5%; y si no lo es: La caída de tensión es superior al 5%
    if ΔU < 5:
        print("La caída de tensión es inferior al 5%")
    else:
        print("La caída de tensión es superior al 5%")

    #-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # DENSIDAD MÁXIMA DE CORRIENTE E INTENSIDAD MÁXIMA POR CABLE
    #------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.header("\n9. DENSIDAD MÁXIMA DE CORRIENTE E INTENSIDAD MÁXIMA POR CABLE")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )   
    st.write("""
    La densidad máxima admisible de corriente en régimen permanente para corriente alterna 
    y frecuencia de 50 Hz se deduce de la tabla 11 recogida en el artículo 4.2.1 
    de la ITC-LAT-07.
    """)
    st.write("""
    La densidad máxima del conductor de Al-Ac, $\\delta_{LA}$, viene dada por:
    """)
    st.latex(r"""
    \delta_{LA} = \delta_L \cdot k
    """)
    st.write("""
    y la intensidad máxima por fase, $I$, para una sección total del conductor $S_{LA}$, por:
    """)
    st.latex(r"""
    I = \delta_{LA} \cdot S_{LA}
    """)
    st.write(f"""
    Los valores correspondientes al conductor {conductor}, para un $\\cos \\varphi$ =  {cos_phi}, son los que a continuación se indican:
    """)
    # Densidades de corriente
    secciones = [15, 25, 35, 50, 70, 95, 125, 160, 200, 250, 300, 400, 500, 600]
    densidades = [6, 5, 4.55, 4.0, 3.55, 3.2, 2.9, 2.7, 2.5, 2.3, 2.15, 1.95, 1.8, 1.65]  # A/mm2
    def densidad_corriente(seccion):
        if seccion < secciones[0] or seccion > secciones[-1]:
            return None  # Fuera de rango
        for i in range(len(secciones) - 1):
            if secciones[i] <= seccion <= secciones[i + 1]:
                # Regla de tres
                densidad = densidades[i] + (densidades[i + 1] - densidades[i]) * (seccion - secciones[i]) / (secciones[i + 1] - secciones[i])
                return densidad
        return None
    densidad_maxima = densidad_corriente(seccion)
    # Coeficientes reductores
    composiciones = ["30+7", "6+1", "26+7", "54+7", "45+7"]
    coef_reductores = [0.916, 0.937, 0.937, 0.95, 0.97]
    def coeficiente_reductor(composicion):
        if composicion in composiciones:
            index = composiciones.index(composicion)
            return coef_reductores[index]
        return None
    coef_reductor = coeficiente_reductor(composicion)
    densidad_max_con_reduccion = densidad_maxima * coef_reductor
    # Intensidad máxima
    intensidad_maxima_conductor = densidad_max_con_reduccion * seccion  # A

    st.latex(rf"\delta_{{L}} = {densidad_maxima:.4f} \, A/mm^2")
    st.latex(rf"k = {coef_reductor:.4f}")
    st.latex(rf"\delta_{{LA}} = {densidad_max_con_reduccion:.4f} \, A/mm^2")
    st.latex(rf"I = {densidad_max_con_reduccion:.4f} \cdot {seccion} = {intensidad_maxima_conductor:.2f} \, A")
    
    #-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # POTENCIA MÁXIMA ADMISIBLE POR INTENSIDAD
    #-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.header("\n10. POTENCIA MÁXIMA ADMISIBLE POR INTENSIDAD")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )   
    potencia_maxima_admisible = (intensidad_maxima_conductor * tension_nominal * raiz(3)) * cos_phi 
    st.write("""La potencia que puede transportar la línea está limitada por la intensidad máxima y por la aída de tensión, la cual no deberá exceder del 5%""")
    st.write("""La máxima potencia a transportar limitada por laintensidad máxima es:""")
    st.latex(rf"P_{{max}} = \sqrt{{3}} \cdot U_{{nom}} \cdot I_{{max}} \cdot \cos \varphi")
    st.write(f"""Tendremos para un $\\cos \\varphi$ = {cos_phi}:""")
    st.latex(rf"P_{{max}} = \sqrt{{3}} \cdot {tension_nominal} \cdot {intensidad_maxima_conductor:.2f} \cdot {cos_phi} = {potencia_maxima_admisible:.2f} kW")
    # Comparar potencia máxima admisible con potencia transportada
    if potencia_maxima_admisible/1000 >= potencia_transportada_MW:
        print("La potencia es mayor que la potencia a transportar.")
    else:
        print("La potencia es menor que la potencia a transportar.")

    #-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # POTENCIA DE TRANSPORTE EN FUNCIÓN DE CONDICIONES METEOROLÓGICAS
    #-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.header("\n11. CÁLCULO POTENCIA DE TRANSPORTE EN FUNCIÓN DE CONDICIONES METEOROLÓGICAS")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )   
    # Calor aportado por radiación solar para invierno y verano Qs = coeficiente_absorcion * radiacion * diametro/1000
    st.subheader("\n   11.1. CALOR APORTADO POR RADIACIÓN SOLAR")
    Qs_invierno = coeficiente_absorcion * radiacion_invierno * diametro/1000 # W/m
    Qs_verano = coeficiente_absorcion * radiacion_verano * diametro/1000 # W/m

    st.markdown("""
    La radiación solar sobre el cable tiene en cuenta tanto la componente directa como la difusa. Se puede expresar como:
    """)
    st.latex(r"""
    Q_S = \alpha \cdot S \cdot d_c
    """)
    st.markdown("""
    Donde:
    - $\\alpha$ es el coeficiente de absorción del conducto 
    - $S$ es la radiación global en $\\text{W/m}^2$
    - $d_c$ es el diámetro del conductor en metros
    """)
    st.markdown("""
    Realizando los cálculos con los valores de la línea para las dos estaciones del año:
    """)
    st.latex(rf"Q_{{s_{{inv}}}} = {coeficiente_absorcion} \cdot {radiacion_invierno:.0f} \cdot {{{diametro:.2f}}} \cdot 10^{{-3}} = {Qs_invierno:.3f} \, W/m")
    st.latex(rf"Q_{{s_{{ver}}}} = {coeficiente_absorcion} \cdot {radiacion_verano:.0f} \cdot {{{diametro:.2f}}} \cdot 10^{{-3}} = {Qs_verano:.3f} \, W/m")
    # Calor cedido por radiación Qr = pi * diametro/1000 * emisividad_conductor * cte_boltzmann * ((Tc+273.15)^4 - (Ta+273.15)^4)

    st.subheader("\n   11.2. CALOR CEDIDO POR RADIACIÓN SOLAR")
    cte_boltzmann = 5.6704e-8 # W/m2K4
    Qr_invierno = pi * (diametro/1000) * emisividad_conductor * cte_boltzmann * ((Tc + 273.15)**4 - (temperatura_invierno + 273.15)**4) # W/m
    Qr_verano = pi * (diametro/1000) * emisividad_conductor * cte_boltzmann * ((Tc + 273.15)**4 - (temperatura_verano + 273.15)**4) # W/m
    st.markdown("""
    La pérdida de calor por radiación se puede calcular con la siguiente expresión:
    """)
    st.latex(r"""
    Q_r = \pi \cdot d_c \cdot \varepsilon \cdot \sigma_b 
    \left[
    \left(T_c + 273.15\right)^4 
    -
    \left(T_{amb} + 273.15\right)^4
    \right]
    """)
    st.markdown("""
    Siendo:
    - ϵ: emisividad del conductor  
    - $σ_b$: constante de Stefan–Boltzmann de valor $5.6704 \\cdot 10^{-8} \\, \\text{W/m}^2\\text{K}^4$  
    - $T_c$: temperatura del conductor en régimen permanente (ºC)  
    - $T_{amb}$: temperatura ambiente máxima, función de la estación del año (ºC)  
    """)
    st.markdown("""
    Los resultados del cálculo para los valores de la línea se muestran a continuación:
    """)
    st.latex(rf"Q_{{r_{{inv}}}} = \pi \cdot {{{diametro:.2f}}} \cdot 10^{{-3}} \cdot {emisividad_conductor} \cdot {cte_boltzmann:.4e} \cdot \left[ \left({Tc:.0f} + 273.15\right)^4 - \left({temperatura_invierno:.0f} + 273.15\right)^4 \right] = {Qr_invierno:.3f} \, W/m")
    st.latex(rf"Q_{{r_{{ver}}}} = \pi \cdot {{{diametro:.2f}}} \cdot 10^{{-3}} \cdot {emisividad_conductor} \cdot {cte_boltzmann:.4e} \cdot \left[ \left({Tc:.0f} + 273.15\right)^4 - \left({temperatura_verano:.0f} + 273.15\right)^4 \right] = {Qr_verano:.3f} \, W/m")

    # Calor cedido por convección Qc = pi * conductividad_termica * (Tc - Ta) * Nu
    st.subheader("\n   11.3. CALOR CEDIDO POR CONVECCIÓN")
    st.markdown("""Al aumentar la temperatura del conductor el aire adyacente se calienta. 
                Dado que la densidad del aire caliente disminuye con la temperatura, provoca el ascenso de este aire. 
                Aire frío remplaza este aire caliente, eliminando calor del conductor. Este fenómeno se denomina convección natural.
                \nSi existe velocidad del viento, el aire caliente será arrastrado reemplazándose por aire más frío que elimina 
                calor del conductor.
                \nSi las velocidades de viento son pequeñas, se calculará tanto la convección forzada como convección natural y se elegirá 
                el valor mayor del coeficiente de transmisión del calor.
                \nEn cualquier caso la forma de calcular será:
    """)
    st.latex(r"""
    Q_c = \pi \cdot \lambda_f \cdot (T_c - T_{amb}) \cdot Nu
    """)
    st.markdown("""
    Siendo:

    - $\lambda_f$: conductividad térmica del aire en las proximidades del conductor  
    - $T_c$: temperatura del conductor en régimen permanente (ºC)  
    - $T_{amb}$: temperatura ambiente máxima, en función de la estación del año (ºC)  
    - $Nu$: número de Nusselt, calculado en función del tipo de convección (forzada o natural)  
    """)
    
    # 1.  Convección natural: Nu = A * (Gr*Pr)^m
    print("\n       11.3.1. CONVECCIÓN NATURAL")
    # Valores de A y m dependiendo de Gr*Pr
    def coeficientes_conveccion_natural(Gr_Pr):
        if 0.1 <= Gr_Pr < 100:
            return 1.02, 0.148
        elif 100 <= Gr_Pr < 10000:
            return 0.85, 0.188
        elif 10000 <= Gr_Pr < 1e7:
            return 0.48, 0.25
        elif 1e7 <= Gr_Pr < 1e12:
            return 0.125, 0.333
        else:
            return None, None
    # Gr = (diametro/1000)^3 * (Tc - Ta) * g / ((Tav + 27.15) * (viscosidad_cinematica)^2)
    g = 9.81  # m/s2
    # viscosidad_cinematica = 1.32e-5 + 9.5e-8 * Tav
    def viscosidad_cinematica(Tav):
        return 1.32e-5 + 9.5e-8 * Tav
    viscosidad_cinematica_invierno = viscosidad_cinematica((Tc + temperatura_invierno) / 2)
    viscosidad_cinematica_verano = viscosidad_cinematica((Tc + temperatura_verano) / 2)
    print(f"\nViscosidad cinemática en invierno: {viscosidad_cinematica_invierno:.4e} m²/s")
    print(f"Viscosidad cinemática en verano: {viscosidad_cinematica_verano:.4e} m²/s")
    # densidad_relativa_aire = exp(-1.16e-4* h)
    def Gr(Ta, viscosidad_cinematica):
        Tav = (Tc + Ta) / 2
        return ((diametro/1000)**3 * (Tc - Ta) * g) / ((Tav + 273.15) * (viscosidad_cinematica**2))
    Gr_invierno = Gr(temperatura_invierno, viscosidad_cinematica_invierno)
    Gr_verano = Gr(temperatura_verano, viscosidad_cinematica_verano)
    print(f"\nNúmero de Grashof en invierno: {Gr_invierno:.2f}")
    print(f"Número de Grashof en verano: {Gr_verano:.2f}")
    # Pr = calor_esp_aire * viscosidad_dinamica / conductividad_termica_aire
    calor_esp_aire = 1005  # J/kgK
    # viscosidad_dinamica = (17.239 + 4.625e-2 * Tav - 2.03e-5 * Tav **2)e-6 kg/m s
    def viscosidad_dinamica(Tav):
        return (17.239 + 4.625e-2 * Tav - 2.03e-5 * Tav **2) * 1e-6  # kg/m s
    viscosidad_dinamica_invierno = viscosidad_dinamica((Tc + temperatura_invierno) / 2)
    viscosidad_dinamica_verano = viscosidad_dinamica((Tc + temperatura_verano) / 2)
    print(f"\nViscosidad dinámica en invierno: {viscosidad_dinamica_invierno:.6e} kg/m s")
    print(f"Viscosidad dinámica en verano: {viscosidad_dinamica_verano:.6e} kg/m s")
    # conductividad_termica_aire = 2.368e-2 + 7.23e-5 * Tav - 2.763e-8 * Tav**2 W/K·m
    def conductividad_termica_aire(Tav):
        return 2.368e-2 + 7.23e-5 * Tav - 2.763e-8 * Tav**2  # W/K·m
    conductividad_termica_aire_invierno = conductividad_termica_aire((Tc + temperatura_invierno) / 2)
    conductividad_termica_aire_verano = conductividad_termica_aire((Tc + temperatura_verano) / 2)
    print(f"\nConductividad térmica del aire en invierno: {conductividad_termica_aire_invierno:.6e} W/K·m")
    print(f"Conductividad térmica del aire en verano: {conductividad_termica_aire_verano:.6e} W/K·m")
    Pr_invierno = calor_esp_aire * viscosidad_dinamica_invierno / conductividad_termica_aire_invierno
    Pr_verano = calor_esp_aire * viscosidad_dinamica_verano / conductividad_termica_aire_verano
    print(f"\nNúmero de Prandtl en invierno: {Pr_invierno:.4f}")
    print(f"Número de Prandtl en verano: {Pr_verano:.4f}")
    Gr_Pr_invierno = Gr_invierno * Pr_invierno
    Gr_Pr_verano = Gr_verano * Pr_verano
    print(f"\nNúmero Gr*Pr en invierno: {Gr_Pr_invierno:.2f}")
    print(f"\nNúmero Gr*Pr en verano: {Gr_Pr_verano:.2f}")
    Nu_invierno = 0
    Nu_verano = 0
    A_invierno, m_invierno = coeficientes_conveccion_natural(Gr_Pr_invierno)
    if A_invierno is not None:
        Nu_invierno = A_invierno * (Gr_Pr_invierno ** m_invierno)
    A_verano, m_verano = coeficientes_conveccion_natural(Gr_Pr_verano)
    if A_verano is not None:
        Nu_verano = A_verano * (Gr_Pr_verano ** m_verano)
    print(f"\nNúmero de Nusselt en invierno: {Nu_invierno:.4f}")
    print(f"Número de Nusselt en verano: {Nu_verano:.4f}")





    st.markdown("""El número de Nusselt para enfriamiento por convección natural depende del producto de los números de Grasshof
                y Prandtl:""")
    st.latex(r"Nu = A \cdot (Gr \cdot Pr)^m")
    st.markdown("""Los valores de las constantes $A$ y $m$ dependen del valor del producto $Gr \\cdot Pr$ según la siguiente tabla:""")
    st.latex(r"""
             \begin{array}{|c|c|c|c|}
             \hline
             &
             \mathbf{A} &
             \mathbf{m} \\
             \hline
             0.1 < Gr \cdot Pr < 100 &
             1.02 &
             0.148 \\
             \hline
             100 < Gr \cdot Pr < 10^4 &
             0.85 &
             0.188 \\
             \hline
             10^4 < Gr \cdot Pr < 10^7&
             0.48 &
             0.25 \\
             \hline
             10^7 < Gr \cdot Pr < 10^{12}&
             0.125 &
             0.333 \\
             \hline
             \end{array}
             """)
    st.markdown("""El número de Grashof se puede calcular como:""")
    st.latex(r"""Gr = \frac{d_c^3 \cdot (T_c - T_{amb}) \cdot g}{T_{av} \cdot \nu_f^2}""")
    st.markdown("""
                Donde:
                - $d_c$: Diámetro exterior del conductor en metros
                - $T_c$: Temperatura del conductor en régimen permanente
                - $T_{amb}$: Temperatura ambiente en función de la estación del año
                - $g$: Aceleración de la gravedad, con un valor de 9.81 m/s²
                - $T_{av}$: Temperatura media entre el conductor y el ambiente, calculada como $(T_c + T_{amb})/2$
                - $\\nu_f$: Viscosidad cinemática del aire en m²/s
                """)
    st.markdown("""El número de Grashof en invierno es:""")
    st.latex(rf"Gr_{{inv}} = \frac{{\left({{{diametro}}} \cdot 10^{-3}\right)^3 \cdot ({Tc} - {temperatura_invierno}) \cdot 9.81}}{{\left(\frac{{{Tc} + {temperatura_invierno}}}{{2}} + 273.15\right) \cdot ({viscosidad_cinematica_invierno:.4e})^2}} = {Gr_invierno:.2f}")
    st.markdown("""El número de Grashof en verano es:""")
    st.latex(rf"Gr_{{ver}} = \frac{{\left({{{diametro}}} \cdot 10^{-3}\right)^3 \cdot ({Tc} - {temperatura_verano}) \cdot 9.81}}{{\left(\frac{{{Tc} + {temperatura_verano}}}{{2}} + 273.15\right) \cdot ({viscosidad_cinematica_verano:.4e})^2}} = {Gr_verano:.2f}")
    st.markdown("""El número de Prandtl se puede calcular como:""")
    st.latex(r"""Pr = \frac{C_p \cdot \mu}{\lambda}""")



    # 2. Convección forzada: Nu = B1 * Re^n
    print("\n       11.3.2. CONVECCIÓN FORZADA")
    # Re = (diametro/1000) * velocidad_viento / viscosidad_cinematica
    # Rugosidad: Rf = diametro_alambre_ext / (2*(diametro-diametro_alambe_ext))
    Rf = diametro_alambre_ext / (2 * (diametro - diametro_alambre_ext))
    def Re(viscosidad_cinematica):
        return (diametro/1000) * velocidad_viento / viscosidad_cinematica
    Re_invierno = Re(viscosidad_cinematica_invierno)
    Re_verano = Re(viscosidad_cinematica_verano)
    #B1 y n según Re
    def coeficientes_conveccion_forzada(Re, Rf):
        if 100 < Re < 2.65e3:
            return 0.641, 0.471
        elif Rf <= 0.05 and 2.65e3 <= Re < 5e4:
            return 0.178, 0.633
        elif Rf > 0.05 and 2.65e3 <= Re < 5e4:
            return 0.048, 0.8
        else:
            return None, None
    Nu_invierno_forzada = 0
    Nu_verano_forzada = 0
    B1_invierno, n_invierno = coeficientes_conveccion_forzada(Re_invierno, Rf)
    if B1_invierno is not None:
        Nu_invierno_forzada = B1_invierno * (Re_invierno ** n_invierno)
    B1_verano, n_verano = coeficientes_conveccion_forzada(Re_verano, Rf)
    if B1_verano is not None:
        Nu_verano_forzada = B1_verano * (Re_verano ** n_verano)
    # Nu corregido
    Nu_45_invierno = Nu_invierno_forzada * (0.42 + 0.58*math.sin(45*pi/180)**0.9)
    Nu_45_verano = Nu_verano_forzada * (0.42 + 0.58*math.sin(45*pi/180)**0.9)
    # Qc = pi* conductividad_termica * (Tc-Ta) * Nu con Nu=max(Nu_natural, Nu_45_forzada)
    Qc_invierno = pi * conductividad_termica_aire_invierno * (Tc - temperatura_invierno) * max (Nu_invierno, Nu_45_invierno)
    Qc_verano = pi * conductividad_termica_aire_verano * (Tc - temperatura_verano) * max (Nu_verano, Nu_45_verano)

    st.markdown("""Cuando la velocidad de viento es mayor que cero, el número de Nusselt es función de Reynolds a través de la siguiente ecuación:""")
    st.latex(r"Nu = B_1 \cdot Re^n")
    st.markdown("""El número de Reynolds para una corriente de aire se puede calcular como:""")
    st.latex(r"Re = \frac{d \cdot v}{\nu}")
    st.markdown("""
        Siendo:
        - $d$: Diámetro del conductor en metros
        - $v$: Velocidad del viento en m/s
        - $\\nu$: Viscosidad cinemática del aire en m²/s""")
    st.markdown("""Los valores de las constantes $B_1$ y $n$ dependen del número de Reynolds y de la rugosidad del conductor, siendo esta última:""")
    st.latex(r"R_f = \frac{d}{2 \cdot (D - d)}")
    st.markdown("""
        Siendo:
        - $d$: Diámetro del alambre exterior del conductor
        - $D$: Diámetro del conductor""")
    st.markdown("""Los valores de $B_1$ y $n$ se muestran a continuación:""")
    st.latex(r"""
        \begin{array}{|c|c|c|c|c|}
        \hline
        \textbf{Superficie} &
        \textbf{Re} &
        \textbf{Re} &
        \mathbf{B_1} &
        \mathbf{n} \\
        \hline
        & \textbf{desde} & \textbf{hasta} & & \\
        \hline
        \text{Todas las superficies} &
        10^{2} &
        2.65 \cdot 10^{3} &
        0.641 &
        0.471 \\
        \hline
        R_f \le 0.05 &
        > 2.65 \cdot 10^{3} &
        5 \cdot 10^{4} &
        0.178 &
        0.633 \\
        \hline
        R_f > 0.05 &
        > 2.65 \cdot 10^{3} &
        5 \cdot 10^{4} &
        0.048 &
        0.800 \\
        \hline
        \end{array}
        """)
    st.markdown("""Para los datos de este conductor obtenemos la siguiente rugosidad:""")
    st.latex(rf"R_f = \frac{{{diametro_alambre_ext:.2f}}}{{2 \cdot ({diametro:.2f} - {diametro_alambre_ext:.2f})}} = {Rf:.3f}")
    st.markdown(f"""El número de Reynolds para la velocidad de viento estimada de {velocidad_viento} m/s es:""")
    st.markdown("""En invierno:""")
    st.latex(rf"Re_{{inv}} = \frac{{{{{diametro:.2f}}} \cdot 10^{{-3}} \cdot {velocidad_viento} }}{{{viscosidad_cinematica_invierno:.4e}}} = {Re_invierno:.2f}")
    st.markdown("""En verano:""")
    st.latex(rf"Re_{{ver}} = \frac{{{{{diametro:.2f}}} \cdot 10^{{-3}} \cdot {velocidad_viento} }}{{{viscosidad_cinematica_verano:.4e}}} = {Re_verano:.2f}")
    st.markdown("""Por tanto, el número de Nusselt en convección forzada será:""")
    st.markdown("""En invierno:""")
    st.latex(rf"Nu_{{90_{{inv}}}} = B_1 \cdot Re^n = {B1_invierno:.4f} \cdot {Re_invierno:.2f}^{{{n_invierno:.3f}}} = {Nu_invierno_forzada:.4f}")
    st.markdown("""En verano:""")
    st.latex(rf"Nu_{{90_{{ver}}}} = B_1 \cdot Re^n =  {B1_verano:.4f} \cdot {Re_verano:.2f}^{{{n_verano:.3f}}} = {Nu_verano_forzada:.4f}")
    st.markdown("""El cálculo anterior se basa en que el viento incide perpendicularmente a la línea. Para corregir el número de Nusselt a un ángulo de incidencia de 45° se aplica la siguiente corrección:""")
    st.markdown("""En invierno:""")
    st.latex(rf"Nu_{{45_{{inv}}}} = Nu_{{90}} \cdot (0.42 + 0.58 \cdot \sin(45^\circ)^{{0.9}}) = {Nu_invierno_forzada:.4f} \cdot (0.42 + 0.58 \cdot \sin(45^\circ)^{{0.9}}) = {Nu_45_invierno:.4f}")
    st.markdown("""En verano:""")
    st.latex(rf"Nu_{{45_{{ver}}}} = Nu_{{90}} \cdot (0.42 + 0.58 \cdot \sin(45^\circ)^{{0.9}}) = {Nu_verano_forzada:.4f} \cdot (0.42 + 0.58 \cdot \sin(45^\circ)^{{0.9}}) = {Nu_45_verano:.4f}")
    st.markdown("""Finalmente, el calor cedido por convección se calcula con la siguiente expresión, eligiendo el número de Nusselt mayor entre el calculado para convección natural y el corregido para convección forzada:""")
    st.markdown("""En invierno:""")
    st.latex(rf"Q_{{c_{{inv}}}} = \pi \cdot {{\lambda_{{f}}}} \cdot ({{T_c}} - {{T_{{amb}}}}) \cdot Nu = \pi \cdot {conductividad_termica_aire_invierno:.4f} \cdot ({Tc} - {temperatura_invierno}) \cdot {max(Nu_invierno, Nu_45_invierno):.4f} = {Qc_invierno:.3f} \, W/m")
    st.markdown("""En verano:""")
    st.latex(rf"Q_{{c_{{ver}}}} = \pi \cdot {{\lambda_{{f}}}} \cdot ({{T_c}} - {{T_{{amb}}}}) \cdot Nu = \pi \cdot {conductividad_termica_aire_verano:.4f} \cdot ({Tc} - {temperatura_verano}) \cdot {max(Nu_verano, Nu_45_verano):.4f} = {Qc_verano:.3f} \, W/m")

    # Resultados corriente máxima: I = raiz((Qr+Qc-Qs)/resistencia_ca/1000)
    st.subheader("\n   11.4. RESULTADOS CORRIENTE MÁXIMA")
    I_max_invierno = raiz((Qr_invierno + Qc_invierno - Qs_invierno)/ (resistencia_ca * 1e-3))
    I_max_verano = raiz((Qr_verano + Qc_verano - Qs_verano) / (resistencia_ca * 1e-3))
  
    st.markdown("""El calor generado por efecto Joule es:""")
    st.latex(rf"Q_J = I^2 \cdot R_{{ca}}")
    st.markdown("""Conocida la resistencia del conductor a la temperatura de trabajo en corriente alterna,
                 la incógnita será la corriente que puede circular para que en régimen permanente el calor generado sea igual al calor generado""")
    st.latex(rf"Q_J + Q_S = Q_c + Q_r")
    st.markdown("""Los resultados en invierno son:""")
    st.latex(rf"I_{{max_{{inv}}}} = \sqrt{{\frac{{Q_{{r_{{inv}}}} + Q_{{c_{{inv}}}} - Q_{{s_{{inv}}}}}}{{R_{{ca}} \cdot 10^{{-3}}}}}} = \sqrt{{\frac{{{Qr_invierno:.3f} + {Qc_invierno:.3f} - {Qs_invierno:.3f}}}{{{resistencia_ca:.6f} \cdot 10^{{-3}}}}}} = {I_max_invierno:.2f} \, A")
    st.markdown("""Los resultados en verano son:""")
    st.latex(rf"I_{{max_{{ver}}}} = \sqrt{{\frac{{Q_{{r_{{ver}}}} + Q_{{c_{{ver}}}} - Q_{{s_{{ver}}}}}}{{R_{{ca}} \cdot 10^{{-3}}}}}} = \sqrt{{\frac{{{Qr_verano:.3f} + {Qc_verano:.3f} - {Qs_verano:.3f}}}{{{resistencia_ca:.6f} \cdot 10^{{-3}}}}}} = {I_max_verano:.2f} \, A")

    #------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # POTENCIA MÁXIMA DE TRANSPORTE SEGÚN CONDICIONES METEOROLÓGICAS
    #------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.subheader("\n   11.5. POTENCIA MÁXIMA DE TRANSPORTE")
    potencia_maxima_invierno = (I_max_invierno * tension_nominal * raiz(3)) * cos_phi
    potencia_maxima_verano = (I_max_verano * tension_nominal * raiz(3)) * cos_phi
    print(f"\nPotencia máxima de transporte en invierno según condiciones meteorológicas: {potencia_maxima_invierno/1000:.2f} MW")
    print(f"Potencia máxima de transporte en verano según condiciones meteorológicas: {potencia_maxima_verano/1000:.2f} MW")

    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # PÉRDIDAS DE POTENCIA
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.header("\n12. PÉRDIDAS DE POTENCIA")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )   
    print(f"Pot transportada mw= {potencia_transportada_MW}")
    perdidas_potencia = ((potencia_transportada_MW * resistencia_ca * longitud) / (tension_nominal**2 * cos_phi**2))*100
    print(f"\nPérdidas de potencia en la línea: {perdidas_potencia:.5f} %")
    # En valor absoluto
    perdidas_potencia_valor = (potencia_transportada_MW * resistencia_ca * longitud) / (tension_nominal**2 * cos_phi**2) * potencia_transportada_MW
    print(f"Pérdidas de potencia en la línea en valor absoluto: {perdidas_potencia_valor:.5f} MW")

    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # CORTOCIRCUITO MÁXIMO
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.header("\n13. CORTOCIRCUITO MÁXIMO")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )   
    # Constantes según material
    materiales = ["Cobre", "Aluminio-Acero", "Acero"]
    c_conductores = [390, 910, 480]  # J/kgK
    densidades_conductores = [8900, 2700, 7850]  # kg/m3
    ctes_conductores = [56e6, 34.8e6, 7.25e6]  # 1/ohmio m
    alphas_conductores = [0.0039, 0.004, 0.0045]  # 1/ºC
    def propiedades_material(material):
        if material in materiales:
            index = materiales.index(material)
            return (c_conductores[index], densidades_conductores[index], ctes_conductores[index], alphas_conductores[index])
        return (None, None, None, None)
    c_conductor, densidad_conductor, cte_conductor, alpha_conductor = propiedades_material(material)
    print(f"c_conductor, densidad_conductor, cte_conductor, alpha_conductor", c_conductor, densidad_conductor, cte_conductor, alpha_conductor )
    # Temperatura máxima recomendada para el material
    if material == "Acero":
        temperatura_max_recomendada = 300
    else:
        temperatura_max_recomendada = 200
    # Factor K
    multiplicacion = log ((1+alpha_conductor*(temperatura_max_recomendada-20))/(1+alpha_conductor*(Tc-20)))
    print(f"multiplicación:", multiplicacion)
    K = (raiz (cte_conductor*c_conductor*densidad_conductor*multiplicacion/alpha_conductor))/1000000
    print(f"Factor K: {K} A*(s)^(1/2)/mm²")
    Icc_max = K * seccion / raiz(tiempo_accionamiento_proteccion)
    print(f"Icc max: {Icc_max} A")

    #------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # EFECTO CORONA
    #------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.header("\n14. EFECTO CORONA ")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )   
    presion_barometrica = 1/(log(log(76)-altitud_media/18330))*100 
    print(f"Presión barométrica {presion_barometrica} cmHg")
    # Factor corrección densidad aire: delta
    def delta (temperatura):
        return 3.92*presion_barometrica/(273+temperatura)
    delta_invierno = delta (temperatura_invierno)
    delta_verano = delta (temperatura_verano)
    print(f"Factor corrección densidad aire invierno {delta_invierno}")
    print(f"Factor corrección densidad aire verano {delta_verano}")
    # Calcula de la distancia media geométrica DMG
    DMG = 100 * ((matriz_distancias[0][1]*matriz_distancias[1][2]*matriz_distancias[2][0]*matriz_distancias[0][3]*matriz_distancias[1][3]*matriz_distancias[2][3])**(1/6))
    print(f"DMG: {DMG:.3f} cm") # con tierra
    DMG = 100 * ((matriz_distancias[0][1]*matriz_distancias[1][2]*matriz_distancias[2][0])**(1/3)) # cm
    print(f"DMG: {DMG:.1f} cm") # entre fases
    DMG = 807
    print(f"DMG: {DMG}")
    # Tensión crítica dieléctrica
    def Uc (mt,delta):
        return raiz(3)*mc*mt*(30/raiz(2))*delta*diametro/20*log(DMG/(diametro/20))
    Uc_invierno = Uc(mt_invierno, delta_invierno)
    Uc_verano = Uc(mt_verano,delta_verano)
    print(f"Tensión critica invierno {Uc_invierno:.2f} kV")
    print(f"Tensión crítica verano {Uc_verano:.2f} kV")
    if Uc_invierno < tension_nominal:
        print(f"Hay pérdidas por efecto corona en invierno")
        Perdidas_efecto_corona_invierno = 241/delta_invierno * (frecuencia + 25) * raiz (diametro/20/DMG) * ((tension_nominal - Uc_invierno)/raiz(3))**2 * 1e-5
        print(f"Pérdidas por efecto corona en invierno: {Perdidas_efecto_corona_invierno:.2f} kW·km/fase")
    else: print(f"No hay pérdidas por efecto corona en invierno")
    if Uc_verano < tension_nominal:
        print(f"Hay pérdidas por efecto corona en verano")
        Perdidas_efecto_corona_verano = 241/delta_verano * (frecuencia + 25) * raiz (diametro/20/DMG) * (tension_nominal - Uc_verano)/raiz(3) * 1e-5
        print(f"Pérdidas por efecto corona en verano: {Perdidas_efecto_corona_verano:.2f} kW·km/fase")
    else: print(f"No hay pérdidas por efecto corona en verano")

    #------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # CAMPO ELÉCTRICO
    #------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.header("\n15. CAMPO ELÉCTRICO ")
    st.markdown(
        "<hr style='margin-top: 4px; margin-bottom: 8px;'>",
        unsafe_allow_html=True
    )   
    # Matriz coeficientes de potencial
    print("\nMatriz de capacidades (km/uF):")
    df=pd.DataFrame(matriz_capacidades).round(4)
    latex_matrix = r"P = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    # Matriz coeficientes de potencial inversa
    matriz_capacidades_inversa = np.linalg.inv(matriz_capacidades) * 1000
    print("\nMatriz de capacidades inversa(nF/km):")
    df=pd.DataFrame(matriz_capacidades_inversa).round(4)
    latex_matrix = r"P^{-1} = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    # Vector de potenciales (tensiones nominales/raiz(3) (angulos 0 -120 y 120, tension tierra)) en numeros complejos y kV
    V_phase = tension_nominal / raiz(3)
    V_vector = np.array([V_phase * cmath.rect(1, math.radians(0)),
                        V_phase * cmath.rect(1, math.radians(-120)),
                        V_phase * cmath.rect(1, math.radians(120)),
                        0])
    print("\nVector de potenciales (kV):")
    df=pd.DataFrame(V_vector).round(4)
    latex_matrix = r"U = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    # Vector de cargas (matriz_capacidades_inversa @ V_vector)
    Q_vector = matriz_capacidades_inversa @ V_vector /1000
    print("\nVector de cargas (kV·nF/km):")
    df=pd.DataFrame(Q_vector).round(4)
    latex_matrix = r"q = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)

    # Valor complejo de la corriente que circula por cada conductor
    I_phase = potencia_transportada * 1e6 / (raiz(3) * tension_nominal*1e3)
    I_vector = np.array([I_phase * cmath.rect(1, math.radians(0)),
                        I_phase * cmath.rect(1, math.radians(-120)),
                        I_phase * cmath.rect(1, math.radians(120))])
    print("\nVector de corrientes (A):")
    df=pd.DataFrame(I_vector).round(4)
    latex_matrix = r"I = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
    latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)

    resistencia_ca_final = resistencia * (1+7.5*frecuencia**2*(diametro/20)**4*1e-7)
    print(f"\nResistencia:{resistencia_ca_final:.6f}")
    r=resistencia_ca_final*(1+alpha_conductor*(Tc-20))
    print(f"\nResist: {r:.6f}")


    #---------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # EJECUCIÓN PROGRAMA
    #---------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    sys.stdout = sys.__stdout__
    # Ahora redirigir al buffer para capturar el resto de los cálculos
    sys.stdout = buffer
    sys.stdout = sys.__stdout__
    texto_resultados = buffer.getvalue()
    html = """
    <html>
    <head>
    <meta charset="utf-8">
    <style>
    body { font-family: Arial; }
    .mayus { font-weight: bold; }
    </style>
    </head>
    <body>
    <pre>
    """
    def es_mayusculas(linea):
        letras = [c for c in linea if c.isalpha()]
        return len(letras)>0 and all(c.isupper() for c in letras)
    for linea in texto_resultados.split("\n"):
        # Detectar líneas completamente en mayúsculas
        if es_mayusculas(linea):
            html += f'<span class="mayus">{linea}</span>\n'
        else:
            html += linea + "\n"
    html += """
    </pre>
    </body>
    </html>
    """
    with open("resultados.html", "w", encoding="utf-8") as f:
        f.write(html)


    from docx import Document
    doc = Document()
    doc.add_heading("Resultados", 1)
    for linea in texto_resultados.split("\n"):
        doc.add_paragraph(linea)
    doc.save("resultados.docx")












#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# DOBLE CIRCUITO SIMPLEX
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

elif ndecircuitos==2 and ndeconductoresporfase==1:
    # Cálculo de la resistencia a 85ºC y en ca
    print("\n1. RESISTENCIA ELÉCTRICA DE LA LÍNEA")
    def resistencia_a_temp(resistencia_20C, temp):
        return resistencia_20C * (1 + coef_temp * (temp - 20))
    resistencia_85C = resistencia_a_temp(resistencia, Tc)
    print(f"Resistencia a {Tc}ºC: {resistencia_85C:.6f} Ω")
    reactancia_pelicular_85C = 10**-7 * 8 * pi * frecuencia * (1 / (resistencia_a_temp(resistencia, Tc) / 1000))
    print(f"Reactancia por efecto pelicular a {Tc}ºC: {reactancia_pelicular_85C:.6f} Ω")
    ys = reactancia_pelicular_85C ** 2 / (192 + 0.8*reactancia_pelicular_85C**2)
    resistencia_ca = resistencia_85C * (1 + ys)
    print(f"Resistencia en CA a {Tc}ºC: {resistencia_ca:.6f} Ω")
    r_ca_longitud = resistencia_ca * longitud
    print(f"Resistencia en CA para {longitud} km: {r_ca_longitud:.6f} Ω")
    
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # CÁLCULO MATRIZ IMPEDANCIAS
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    print("\n2. MATRIZ DE IMPEDANCIAS")
    penetracion_terreno = raiz(resistividad/(pi*frecuencia*mu0))
    print(f"Penetración terreno: {penetracion_terreno:.6f} m")
    # Matriz de distancias
    puntos = [uno, dos, tres, cuatro, cinco, seis, tierra]
    n_puntos = len(puntos)
    matriz_distancias = [[0]*n_puntos for _ in range(n_puntos)]
    def _xy(pt):
        try:
            if isinstance(pt, (list, tuple)) and len(pt) >= 2:
                return float(pt[0]), float(pt[1])
            if isinstance(pt, str):
                s = pt.strip().replace(';', ',')
                try:
                    v = ast.literal_eval(s)
                    if isinstance(v, (list, tuple)) and len(v) >= 2:
                        return float(v[0]), float(v[1])
                except Exception:
                    parts = s.split(',')
                    if len(parts) >= 2:
                        return float(parts[0]), float(parts[1])
            if isinstance(pt, (int, float)):
                return float(pt), 0.0
        except Exception:
            pass
        raise ValueError(f"Formato de punto inválido: {pt}")
    for i in range(n_puntos):
        for j in range(n_puntos):
            x1, y1 = _xy(puntos[i])
            x2, y2 = _xy(puntos[j])
            distancia = raiz((x2-x1)**2 + (y2-y1)**2)
            matriz_distancias[i][j] = distancia
    print("\nMatriz de distancias (km):")
    df = pd.DataFrame(matriz_distancias).round(4)
    latex_matrix = r"D = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
        latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)

    # Matriz D_prima: distancia entre puntos y sus espejos respecto al suelo
    matriz_D_prima = [[0]*n_puntos for _ in range(n_puntos)]
    for i in range(n_puntos):
        for j in range(n_puntos):
            x1, y1 = _xy(puntos[i])
            # Espejo del punto j respecto al suelo
            x2, y2 = _xy(puntos[j])
            x2_espejo, y2_espejo = x2, -y2
            distancia = raiz((x2_espejo-x1)**2 + (y2_espejo-y1)**2)
            matriz_D_prima[i][j] = distancia
    print("\nMatriz D' (distancias entre puntos y espejos) (km):")
    df = pd.DataFrame(matriz_D_prima).round(4)
    latex_matrix = r"D' = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
        latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    # kij = raiz(2)*D_prima_ij / penetracion_terreno
    matriz_kij = [[0]*n_puntos for _ in range(n_puntos)]
    for i in range(n_puntos):
        for j in range(n_puntos):
            matriz_kij[i][j] = raiz(2) * matriz_D_prima[i][j] / penetracion_terreno  
    print("\nMatriz k_ij:")
    df = pd.DataFrame(matriz_kij).round(4)
    latex_matrix = r"k_{ij} = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
        latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    # P_ij = pi/8 - k_ij*cos(tetha_ij)/(3*raiz(2)) siendo cos(tetha_ij) = (h_i + h_j )/ D_prima_ij
    matriz_Pij = [[0]*n_puntos for _ in range(n_puntos)]
    for i in range(n_puntos):
        for j in range(n_puntos):
            h_i = _xy(puntos[i])[1]
            h_j = _xy(puntos[j])[1]
            D_prima_ij = matriz_D_prima[i][j]
            if D_prima_ij != 0:
                cos_tetha_ij = (h_i + h_j) / D_prima_ij
            else:
                cos_tetha_ij = 0
            k_ij = matriz_kij[i][j]
            P_ij = (pi / 8) - (k_ij * cos_tetha_ij) / (3 * raiz(2))
            matriz_Pij[i][j] = P_ij
    print("\nMatriz P_ij:")
    df = pd.DataFrame(matriz_Pij).round(4)
    latex_matrix = r"P_{ij} = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
        latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    # Q_ij = 0.5*log neperiano(1.85138/k_ij) + k_ij*cos(tetha_ij)/(3*raiz(2))
    matriz_Qij = [[0]*n_puntos for _ in range(n_puntos)]
    for i in range(n_puntos):
        for j in range(n_puntos):
            k_ij = matriz_kij[i][j]
            h_i = _xy(puntos[i])[1]
            h_j = _xy(puntos[j])[1]
            D_prima_ij = matriz_D_prima[i][j]
            if D_prima_ij != 0:
                cos_tetha_ij = (h_i + h_j) / D_prima_ij
            else:
                cos_tetha_ij = 0
            if k_ij != 0:
                Q_ij = 0.5 * log(1.85138 / k_ij) + (k_ij * cos_tetha_ij) / (3 * raiz(2))              
            else:
                Q_ij = 0
            matriz_Qij[i][j] = Q_ij
    print("\nMatriz Q_ij:")
    df = pd.DataFrame(matriz_Qij).round(4)
    latex_matrix = r"Q_{ij} = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
        latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    # Resistencia de cada uno de los conductores (ultimo a tierra con dato de resistencia a tierra directamente)
    resistencias_conductores = [resistencia_ca, resistencia_ca, resistencia_ca, resistencia_ca, resistencia_ca, resistencia_ca, resistencia_tierra]
    print("\nResistencias de los conductores (Ω/km):")
    for i, R in enumerate(resistencias_conductores):
        print(f" {R:.6f} Ω/km")
    # Radio de cada conductor en mm
    radio_conductores_mm = [
        diametro / 2,
        diametro / 2,
        diametro / 2,
        diametro / 2,
        diametro / 2,
        diametro / 2,
        diametro_tierra / 2
    ]
    print("\nRadio de los conductores (mm):")
    df=pd.DataFrame(radio_conductores_mm).round(4)
    latex_matrix = r"r = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
        latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    # Matriz de impedancias con:
    # Z_ii = R_i + j*mu0*frecuencia*((1/(4*n_i)+log(D_prima_ii/r_i)))+mu0*2*frecuencia*(P_ii + j*Q_ii) 
    # Z_ij = j*mu0*frecuencia*(log(D_prima_ij/D_ij)) + mu0*2*frecuencia*(P_ij + j*Q_ij) para i != j
    matriz_impedancias = [[0]*n_puntos for _ in range(n_puntos)]
    for i in range(n_puntos):
        for j in range(n_puntos):
            R_i = resistencias_conductores[i]
            D_ij = matriz_distancias[i][j]
            D_prima_ij = matriz_D_prima[i][j]
            P_ij = matriz_Pij[i][j]
            Q_ij = matriz_Qij[i][j]
            r_i = radio_conductores_mm[i] / 1000  # Convertir mm a m
            if i == j:
                n_i = ndeconductoresporfase
                Z_ii = R_i + (1j * mu0 * frecuencia * ((1 / (4 * n_i)) + log(D_prima_ij / r_i)) + mu0 * 2 * frecuencia * (P_ij + 1j * Q_ij))*1000
                matriz_impedancias[i][j] = Z_ii
            else:
                Z_ij = (1j * mu0 * frecuencia * (log(D_prima_ij / D_ij)) + mu0 * 2 * frecuencia * (P_ij + 1j * Q_ij))*1000
                matriz_impedancias[i][j] = Z_ij
    print("\nMatriz de impedancias (Ω/km):")
    for i, fila in enumerate(matriz_impedancias):
        print(f"", end="")
        for Z in fila:
            print(f"{Z:18.4f}", end="  ")
        print()
    df = pd.DataFrame(matriz_impedancias).round(4)
    latex_matrix = r"Z = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
        latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    """ Como la linea posee cable de tierra, es necesario realizar un análisis matricial para eliminarlos y obtener una matriz 3*3 que representa las impedancias por fase
    Z=[Zf, Zft; Ztf, Zt]
    Zfas = Zf - Zft * inv(Zt) * Ztf
    """
    Z = np.array(matriz_impedancias)
    Zf = Z[0:6, 0:6]
    Zt = Z[6:7, 6:7]
    Zft = Z[0:6, 6:7]
    Ztf = Z[6:7, 0:6]
    Zt_inv = np.linalg.inv(Zt)
    Zfas = Zf - Zft @ Zt_inv @ Ztf
    print("\nMatriz de impedancias por fase (Ω/km):")
    df = pd.DataFrame(Zfas).round(4)
    latex_matrix = r"Z = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.4f}" for v in row)
        latex_matrix += r" \\ "
        latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    # Finalmente, multiplicar por la longitud de la línea para obtener las impedancias totales
    Zfas_total = Zfas * longitud
    print("\nMatriz de impedancias por fase total para la longitud dada (Ω):")
    df = pd.DataFrame(Zfas_total).round(4)
    latex_matrix = r"Z = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.3f}" for v in row)
        latex_matrix += r" \\ "
        latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    # Impedancias de secuencia
    print("\n3. IMPEDANCIAS DE SECUENCIA")
    A = np.array([[1, 1, 1, 0, 0, 0],
                [complex(-0.5, -raiz(3)/2), complex(-0.5, raiz(3)/2), 1, 0, 0, 0],
                [complex(-0.5, raiz(3)/2), complex(-0.5, -raiz(3)/2), 1, 0, 0, 0],
                [0, 0, 0, 1, 1, 1],
                [0, 0, 0, complex(-0.5, -raiz(3)/2), complex(-0.5, raiz(3)/2), 1],
                [0, 0, 0, complex(-0.5, raiz(3)/2), complex(-0.5, -raiz(3)/2), 1]])
    A_inv = np.linalg.inv(A)
    Z_seq = A_inv @ Zfas @ A
    print("\nMatriz de impedancias de secuencia (Ω):")
    df = pd.DataFrame(Z_seq).round(4)
    latex_matrix = r"Z_{012} = \begin{pmatrix}"
    for row in df.values:
        latex_matrix += " & ".join(f"{v:.3f}" for v in row)
        latex_matrix += r" \\ "
        latex_matrix += r"\end{pmatrix}"
    st.latex(latex_matrix)
    # Matriz anterior de impedancias de secuencia pero con argumento y angulo
    print("\nMatriz de impedancias de secuencia con magnitud y ángulo (Ω):")
    for i, fila in enumerate(Z_seq):
        print(f" ", end="")
        for Z in fila:
            magnitud = abs(Z)
            angulo = math.degrees(math.atan2(Z.imag, Z.real))
            print(f"{magnitud:12.4f} ∠ {angulo:8.2f}°", end="  ")
        print()
    Z0 = Z_seq[2, 2]
    print(f"\nImpedancia homopolar de la línea (Z0): {Z0:.3f} Ω/km")
    Z1 = Z_seq[1, 1]
    print(f"Impedancia directa e inversa de la línea (Z1): {Z1:.3f} Ω/km")
    #Teniendo en cuenta la longitud de la línea
    Z_seq_total = Z_seq * longitud
    print("\nMatriz de impedancias de secuencia total para la longitud dada (Ω):")
    for i, fila in enumerate(Z_seq_total):
        print(f" ", end="")
        for Z in fila:
            print(f"{Z:18.3f}", end="  ")
        print()
    # Matriz anterior de impedancias de secuencia pero con argumento y angulo
    print("\nMatriz de impedancias de secuencia total con magnitud y ángulo para la longitud dada (Ω):")
    for i, fila in enumerate(Z_seq_total):
        print(f" ", end="")
        for Z in fila:
            magnitud = abs(Z)
            angulo = math.degrees(math.atan2(Z.imag, Z.real))
            print(f"{magnitud:12.3f} ∠ {angulo:8.2f}°", end="  ")
        print()
    Z0_total = Z_seq_total[2, 2]
    print(f"\nImpedancia homopolar de la línea total (Z0): {Z0_total:.3f} Ω")
    Z1_total = Z_seq_total[1, 1]
    print(f"Impedancia directa e inversa de la línea total (Z1): {Z1_total:.3f} Ω")
    # Resistencias R0/R1
    R0 = Z0.real
    R1 = Z1.real
    print(f"\nR0/R1: {R0/R1:.3f}")
    # Reactancias X0/X1
    X0 = Z0.imag
    X1 = Z1.imag
    print(f"X0/X1: {X0/X1:.3f}")

    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # MATRIZ DE CAPACIDADES
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    print("\n4. MATRIZ DE CAPACIDADES")
    # Coeficientes de potencial
    # C_ii = (1/(2*pi*epsilon0))*log(D_prima_ii/r_i)
    # C_ij = (1/(2*pi*epsilon0))*log(D_prima_ij/D_ij) para i != j
    matriz_capacidades = [[0]*n_puntos for _ in range(n_puntos)]
    for i in range(n_puntos):
        for j in range(n_puntos):
            D_ij = matriz_distancias[i][j]/1000   # Convertir m a km
            D_prima_ij = matriz_D_prima[i][j]/1000   # Convertir m a km
            r_i = radio_conductores_mm[i] / 1000000  # Convertir mm a km
            if i == j:
                C_ii = ((1 / (2 * pi * epsilon0)) * log(D_prima_ij / r_i))/1000000000
                matriz_capacidades[i][j] = C_ii
            else:
                C_ij = (1 / (2 * pi * epsilon0)) * log(D_prima_ij / D_ij)/1000000000
                matriz_capacidades[i][j] = C_ij
    print("\nMatriz de coeficientes de potencial (km/uF):")
    for i, fila in enumerate(matriz_capacidades):
        print(f" ", end="")
        for C in fila:
            print(f"{C:18.2f}", end="  ")
        print()
    # Matriz de capacitancias por fase (eliminando tierra)
    C = np.array(matriz_capacidades)
    Cf = C[0:6, 0:6]
    Ct = C[6:7, 6:7]
    Cft = C[0:6, 6:7]
    Ctf = C[6:7, 0:6]
    Ct_inv = np.linalg.inv(Ct)
    Cfas = Cf - Cft @ Ct_inv @ Ctf
    print("\nMatriz de coeficientes de potencial por fase (km/μF):")
    for i, fila in enumerate(Cfas):
        print(f" ", end="")
        for C in fila:
            print(f"{C:18.2f}", end="  ")
        print()

    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # MATRIZ DE SUSCEPTANCIAS
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    print("\n5. MATRIZ DE SUSCEPTANCIAS")
    # Matriz de susceptancias por km (B=j*2*pi*frecuencia*C^-1) solo parte imaginaria, no mostrando la parte real
    Cfas_inv = np.linalg.inv(Cfas)
    Bfas = 1j * 2 * pi * frecuencia * Cfas_inv
    print("\nMatriz de susceptancias por fase (μS/km):")
    for i, fila in enumerate(Bfas):
        print(f"", end="")
        for B in fila:
            print(f"{B:18.3f}", end="  ")
        print()
    # Matriz de susceptancias total para la longitud dada
    Bfas_total = Bfas * longitud
    print("\nMatriz de susceptancias por fase total para la longitud dada (μS):")
    for i, fila in enumerate(Bfas_total):
        print(f" ", end="")
        for B in fila:
            print(f"{B:18.2f}", end="  ")
        print()
    # Susceptancias de secuencia
    B_seq = A_inv @ Bfas @ A
    print("\nMatriz de susceptancias de secuencia (μS):")
    for i, fila in enumerate(B_seq):
        print(f"", end="")
        for B in fila:
            print(f"{B:18.3f}", end="  ")
        print()
    # Susceptancia homopolar de la linea (B0)
    B0 = B_seq[2, 2]
    print(f"\nSusceptancia homopolar de la línea (B0): {B0:.3f} μS/km")
    # Susceptancia directa e inversa de la linea (B1)
    B1 = B_seq[1, 1]
    print(f"Susceptancia directa e inversa de la línea (B1): {B1:.3f} μS/km")
    # Susceptancia homopolar de la linea total (B0)
    B0_total = B_seq[2, 2] * longitud
    print(f"\nSusceptancia homopolar de la línea total (B0): {B0_total:.3f} μS")
    # Susceptancia directa e inversa de la linea total (B1)
    B1_total = B_seq[1, 1] * longitud
    print(f"Susceptancia directa e inversa de la línea total (B1): {B1_total:.3f} μS")

    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # CÁLCULO IMPEDANCIA CARACTERÍSTICA Y CONSTANTE DE PROPAGACIÓN
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    print("\n6. IMPEDANCIA CARACTERÍSTICA Y CONSTANTE DE PROPAGACIÓN")
    # Impedancia característica (Zc=raiz((R+jX)/(jB)))
    Zc = cmath.sqrt(Z1 / (1 * B1*1e-6))
    magnitud_Zc = abs(Zc)
    angulo_Zc = math.degrees(math.atan2(Zc.imag, Zc.real))
    print(f"\nImpedancia característica de secuencia directa (Zc1):{Zc:.3f} = {magnitud_Zc:.4f} ∠ {angulo_Zc:.2f}° Ω")
    # Constante de propagación (gamma=raiz((R+jX)*(jB)))
    gamma = cmath.sqrt(Z1 * (1 * B1*1e-6))
    magnitud_gamma = abs(gamma)
    angulo_gamma = math.degrees(math.atan2(gamma.imag, gamma.real))
    print(f"\nConstante de propagación de secuencia directa (γ1): {gamma:.4e} = {magnitud_gamma:.4e} ∠ {angulo_gamma:.2f}° 1/km")
    # Constante de propagación total para la longitud dada
    gamma_total = gamma * longitud
    magnitud_gamma_total = abs(gamma_total)
    angulo_gamma_total = math.degrees(math.atan2(gamma_total.imag, gamma_total.real))
    print(f"\nConstante de propagación para la longitud dada: {gamma_total:.4e} = {magnitud_gamma_total:.4e} ∠ {angulo_gamma_total:.2f}°")

    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # POTENCIA CARACTERÍSTICA
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    print("\n7. POTENCIA CARACTERÍSTICA")
    # Pc = tension_nominal**2 / Zc
    Pc = (tension_nominal) ** 2 / Zc
    magnitud_Pc = abs(Pc)
    angulo_Pc = math.degrees(math.atan2(Pc.imag, Pc.real))
    print(f"\nPotencia característica: {Pc:.4f} = {magnitud_Pc:.4f} ∠ {angulo_Pc:.2f}° kVA")

    #-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # CAIDA DE TENSIÓN
    #-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    print("\n8. CAIDA DE TENSIÓN")
    # ΔU=potencia_transportada(MW)*longitud*(R*cos(phi)+X*sin(phi))/((tension_nominal**2)*10*cos(phi))
    potencia_transportada_MW = potencia_transportada * cos_phi  # Convertir MVA a MW usando el factor de potencia
    ΔU = potencia_transportada_MW * 1000 * longitud * (Z1.real * cos_phi + Z1.imag * math.sin(math.acos(cos_phi))) / ((tension_nominal ** 2) * 10 * cos_phi)
    print(f"\nCaída de tensión ΔU: {ΔU:.4f} %")
    # Poner si la caida de tension es inferior al 5%: La caída de tensión es inferior al 5%; y si no lo es: La caída de tensión es superior al 5%
    if ΔU < 5:
        print("La caída de tensión es inferior al 5%")
    else:
        print("La caída de tensión es superior al 5%")

    #-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # DENSIDAD MÁXIMA DE CORRIENTE E INTENSIDAD MÁXIMA POR CABLE
    #------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    print("\n9. DENSIDAD MÁXIMA DE CORRIENTE E INTENSIDAD MÁXIMA POR CABLE")
    # Valores densidad de corriente según la sección
    secciones = [15, 25, 35, 50, 70, 95, 125, 160, 200, 250, 300, 400, 500, 600]
    densidades = [6, 5, 4.55, 4.0, 3.55, 3.2, 2.9, 2.7, 2.5, 2.3, 2.15, 1.95, 1.8, 1.65]  # A/mm2
    def densidad_corriente(seccion):
        if seccion < secciones[0] or seccion > secciones[-1]:
            return None  # Fuera de rango
        for i in range(len(secciones) - 1):
            if secciones[i] <= seccion <= secciones[i + 1]:
                # Regla de tres
                densidad = densidades[i] + (densidades[i + 1] - densidades[i]) * (seccion - secciones[i]) / (secciones[i + 1] - secciones[i])
                return densidad
        return None
    densidad_maxima = densidad_corriente(seccion)
    print(f"\nDensidad máxima de corriente para la sección {seccion} mm2: {densidad_maxima:.4f} A/mm2")
    # Valores coeficiente reductor según composición material
    composiciones = ["30+7", "6+1", "26+7", "54+7", "45+7"]
    coef_reductores = [0.916, 0.937, 0.937, 0.95, 0.97]
    def coeficiente_reductor(composicion):
        if composicion in composiciones:
            index = composiciones.index(composicion)
            return coef_reductores[index]
        return None
    coef_reductor = coeficiente_reductor(composicion)
    print(f"Coeficiente reductor para la composición {composicion}: {coef_reductor:.4f}")

    densidad_max_con_reduccion = densidad_maxima * coef_reductor
    print(f"Densidad máxima de corriente con coeficiente reductor: {densidad_max_con_reduccion:.4f} A/mm2")
    # Intensidad máxima admisible
    intensidad_maxima_conductor = densidad_max_con_reduccion * seccion  # A
    print(f"Corriente máxima admisible del conductor: {intensidad_maxima_conductor:.2f} A")

    #-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # POTENCIA MÁXIMA ADMISIBLE POR INTENSIDAD
    #-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    print("\n10. POTENCIA MÁXIMA ADMISIBLE POR INTENSIDAD")
    potencia_maxima_admisible = (intensidad_maxima_conductor * tension_nominal * raiz(3)) * cos_phi 
    print(f"\nPotencia máxima admisible por intensidad: {potencia_maxima_admisible/1000:.4f} MW")
    # Comparar potencia máxima admisible con potencia transportada
    if potencia_maxima_admisible/1000 >= potencia_transportada_MW:
        print("La potencia máxima admisible por intensidad es mayor que la potencia transportada.")
    else:
        print("La potencia máxima admisible por intensidad es menor que la potencia transportada.")

    #-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # POTENCIA DE TRANSPORTE EN FUNCIÓN DE CONDICIONES METEOROLÓGICAS
    #-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    print("\n11. CÁLCULO POTENCIA DE TRANSPORTE EN FUNCIÓN DE CONDICIONES METEOROLÓGICAS")
    # Calor aportado por radiación solar para invierno y verano Qs = coeficiente_absorcion * radiacion * diametro/1000
    print("\n   11.1. CALOR APORTADO POR RADIACIÓN SOLAR")
    Qs_invierno = coeficiente_absorcion * radiacion_invierno * diametro/1000 # W/m
    Qs_verano = coeficiente_absorcion * radiacion_verano * diametro/1000 # W/m
    print(f"\nCalor aportado por radiación solar en invierno: {Qs_invierno:.3f} W/m")
    print(f"Calor aportado por radiación solar en verano: {Qs_verano:.3f} W/m")
    # Calor cedido por radiación Qr = pi * diametro/1000 * emisividad_conductor * cte_boltzmann * ((Tc+273.15)^4 - (Ta+273.15)^4)
    print("\n   11.2. CALOR CEDIDO POR RADIACIÓN SOLAR")
    cte_boltzmann = 5.6704e-8 # W/m2K4
    Qr_invierno = pi * (diametro/1000) * emisividad_conductor * cte_boltzmann * ((Tc + 273.15)**4 - (temperatura_invierno + 273.15)**4) # W/m
    Qr_verano = pi * (diametro/1000) * emisividad_conductor * cte_boltzmann * ((Tc + 273.15)**4 - (temperatura_verano + 273.15)**4) # W/m
    print(f"\nCalor cedido por radiación en invierno: {Qr_invierno:.3f} W/m")
    print(f"Calor cedido por radiación en verano: {Qr_verano:.3f} W/m")

    # Calor cedido por convección Qc = pi * conductividad_termica * (Tc - Ta) * Nu
    print("\n   11.3. CALOR CEDIDO POR CONVECCIÓN")
    # 1.  Convección natural: Nu = A * (Gr*Pr)^m
    print("\n       11.3.1. CONVECCIÓN NATURAL")
    """ Valores para A y m dependiendo del rango de Gr*Pr:
    0.1 <= Gr*Pr < 100: A=1.02, m=0.148
    100 <= Gr*Pr < 10000: A=0.85, m=0.188
    10000 <= Gr*Pr < 1e7: A=0.48, m=0.25
    1e7 <= Gr*Pr < 1e12: A=0.125, m=0.333
    """
    def coeficientes_conveccion_natural(Gr_Pr):
        if 0.1 <= Gr_Pr < 100:
            return 1.02, 0.148
        elif 100 <= Gr_Pr < 10000:
            return 0.85, 0.188
        elif 10000 <= Gr_Pr < 1e7:
            return 0.48, 0.25
        elif 1e7 <= Gr_Pr < 1e12:
            return 0.125, 0.333
        else:
            return None, None
    # Gr = (diametro/1000)^3 * (Tc - Ta) * g / ((Tav + 27.15) * (viscosidad_cinematica)^2)
    g = 9.81  # m/s2
    # viscosidad_cinematica = 1.32e-5 + 9.5e-8 * Tav
    def viscosidad_cinematica(Tav):
        return 1.32e-5 + 9.5e-8 * Tav
    viscosidad_cinematica_invierno = viscosidad_cinematica((Tc + temperatura_invierno) / 2)
    viscosidad_cinematica_verano = viscosidad_cinematica((Tc + temperatura_verano) / 2)
    print(f"\nViscosidad cinemática en invierno: {viscosidad_cinematica_invierno:.3e} m²/s")
    print(f"Viscosidad cinemática en verano: {viscosidad_cinematica_verano:.3e} m²/s")
    # densidad_relativa_aire = exp(-1.16e-4* h)
    def Gr(Ta, viscosidad_cinematica):
        Tav = (Tc + Ta) / 2
        return ((diametro/1000)**3 * (Tc - Ta) * g) / ((Tav + 273.15) * (viscosidad_cinematica**2))
    Gr_invierno = Gr(temperatura_invierno, viscosidad_cinematica_invierno)
    Gr_verano = Gr(temperatura_verano, viscosidad_cinematica_verano)
    print(f"\nNúmero de Grashof en invierno: {Gr_invierno:.2f}")
    print(f"Número de Grashof en verano: {Gr_verano:.2f}")
    # Pr = calor_esp_aire * viscosidad_dinamica / conductividad_termica_aire
    calor_esp_aire = 1005  # J/kgK
    # viscosidad_dinamica = (17.239 + 4.625e-2 * Tav - 2.03e-5 * Tav **2)e-6 kg/m s
    def viscosidad_dinamica(Tav):
        return (17.239 + 4.625e-2 * Tav - 2.03e-5 * Tav **2) * 1e-6  # kg/m s
    viscosidad_dinamica_invierno = viscosidad_dinamica((Tc + temperatura_invierno) / 2)
    viscosidad_dinamica_verano = viscosidad_dinamica((Tc + temperatura_verano) / 2)
    print(f"\nViscosidad dinámica en invierno: {viscosidad_dinamica_invierno:.3e} kg/m s")
    print(f"Viscosidad dinámica en verano: {viscosidad_dinamica_verano:.3e} kg/m s")
    # conductividad_termica_aire = 2.368e-2 + 7.23e-5 * Tav - 2.763e-8 * Tav**2 W/Km
    def conductividad_termica_aire(Tav):
        return 2.368e-2 + 7.23e-5 * Tav - 2.763e-8 * Tav**2  # W/Km
    conductividad_termica_aire_invierno = conductividad_termica_aire((Tc + temperatura_invierno) / 2)
    conductividad_termica_aire_verano = conductividad_termica_aire((Tc + temperatura_verano) / 2)
    print(f"\nConductividad térmica del aire en invierno: {conductividad_termica_aire_invierno:.3e} W/K·m")
    print(f"Conductividad térmica del aire en verano: {conductividad_termica_aire_verano:.3e} W/K·m")

    Pr_invierno = calor_esp_aire * viscosidad_dinamica_invierno / conductividad_termica_aire_invierno
    Pr_verano = calor_esp_aire * viscosidad_dinamica_verano / conductividad_termica_aire_verano
    print(f"\nNúmero de Prandtl en invierno: {Pr_invierno:.4f}")
    print(f"Número de Prandtl en verano: {Pr_verano:.4f}")
    Gr_Pr_invierno = Gr_invierno * Pr_invierno
    Gr_Pr_verano = Gr_verano * Pr_verano
    Nu_invierno = 0
    Nu_verano = 0
    A_invierno, m_invierno = coeficientes_conveccion_natural(Gr_Pr_invierno)
    if A_invierno is not None:
        Nu_invierno = A_invierno * (Gr_Pr_invierno ** m_invierno)
    A_verano, m_verano = coeficientes_conveccion_natural(Gr_Pr_verano)
    if A_verano is not None:
        Nu_verano = A_verano * (Gr_Pr_verano ** m_verano)
    print(f"\nNúmero de Nusselt en invierno: {Nu_invierno:.4f}")
    print(f"Número de Nusselt en verano: {Nu_verano:.4f}")

    # 2. Convección forzada: Nu = B1 * Re^n
    print("\n       11.3.2. CONVECCIÓN FORZADA")
    # Re = (diametro/1000) * velocidad_viento / viscosidad_cinematica
    # Rugosidad: Rf = diametro_alambre_ext / (2*(diametro-diametro_alambe_ext))
    Rf = diametro_alambre_ext / (2 * (diametro - diametro_alambre_ext))
    print(f"\nRugosidad Rf: {Rf:.3f}")

    def Re(viscosidad_cinematica):
        return (diametro/1000) * velocidad_viento / viscosidad_cinematica
    Re_invierno = Re(viscosidad_cinematica_invierno)
    Re_verano = Re(viscosidad_cinematica_verano)
    print(f"\nNúmero de Reynolds en invierno: {Re_invierno:.2f}")
    print(f"Número de Reynolds en verano: {Re_verano:.2f}")
    # B1 y n según Re
    def coeficientes_conveccion_forzada(Re, Rf):
        if 100 < Re < 2.65e3:
            return 0.641, 0.471
        elif Rf <= 0.05 and 2.65e3 <= Re < 5e4:
            return 0.178, 0.633
        elif Rf > 0.05 and 2.65e3 <= Re < 5e4:
            return 0.048, 0.8
        else:
            return None, None
    Nu_invierno_forzada = 0
    Nu_verano_forzada = 0
    B1_invierno, n_invierno = coeficientes_conveccion_forzada(Re_invierno, Rf)
    if B1_invierno is not None:
        Nu_invierno_forzada = B1_invierno * (Re_invierno ** n_invierno)
    B1_verano, n_verano = coeficientes_conveccion_forzada(Re_verano, Rf)
    if B1_verano is not None:
        Nu_verano_forzada = B1_verano * (Re_verano ** n_verano)
    print(f"\nNúmero de Nusselt por convección forzada en invierno: {Nu_invierno_forzada:.4f}")
    print(f"Número de Nusselt por convección forzada en verano: {Nu_verano_forzada:.4f}")
    Nu_45_invierno = Nu_invierno_forzada * (0.42 + 0.58*math.sin(45*pi/180)**0.9)
    Nu_45_verano = Nu_verano_forzada * (0.42 + 0.58*math.sin(45*pi/180)**0.9)
    print(f"\nNúmero de Nusselt corregido para ángulo de 45° en invierno: {Nu_45_invierno:.4f}")
    print(f"Número de Nusselt corregido para ángulo de 45° en verano: {Nu_45_verano:.4f}")
    # Qc = pi* conductividad_termica * (Tc-Ta) * Nu con Nu=max(Nu_natural, Nu_45_forzada)
    Qc_invierno = pi * conductividad_termica_aire_invierno * (Tc - temperatura_invierno) * max (Nu_invierno, Nu_45_invierno)
    Qc_verano = pi * conductividad_termica_aire_verano * (Tc - temperatura_verano) * max (Nu_verano, Nu_45_verano)
    print(f"\nCalor cedido por convección en invierno: {Qc_invierno:.3f} W/m")
    print(f"Calor cedido por convección en verano: {Qc_verano:.3f} W/m")

    # Resultados corriente máxima: I = raiz((Qr+Qc-Qs)/resistencia_ca/1000)
    print("\n   11.4. RESULTADOS CORRIENTE MÁXIMA")
    print(f"\nResistencia del conductor a la temperatura de cálculo: {resistencia_ca:.6f} Ω/km")
    I_max_invierno = raiz((Qr_invierno + Qc_invierno - Qs_invierno)/ (resistencia_ca * 1e-3))
    I_max_verano = raiz((Qr_verano + Qc_verano - Qs_verano) / (resistencia_ca * 1e-3))
    print(f"\nCorriente máxima admisible en invierno según condiciones meteorológicas: {I_max_invierno:.2f} A")
    print(f"Corriente máxima admisible en verano según condiciones meteorológicas: {I_max_verano:.2f} A")

    #------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # POTENCIA MÁXIMA DE TRANSPORTE SEGÚN CONDICIONES METEOROLÓGICAS
    #------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    print("\n   11.5. POTENCIA MÁXIMA DE TRANSPORTE")
    potencia_maxima_invierno = (I_max_invierno * tension_nominal * raiz(3)) * cos_phi
    potencia_maxima_verano = (I_max_verano * tension_nominal * raiz(3)) * cos_phi
    print(f"\nPotencia máxima de transporte en invierno según condiciones meteorológicas: {potencia_maxima_invierno/1000:.2f} MW")
    print(f"Potencia máxima de transporte en verano según condiciones meteorológicas: {potencia_maxima_verano/1000:.2f} MW")

    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # PÉRDIDAS DE POTENCIA
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    print("\n12. PÉRDIDAS DE POTENCIA")
    perdidas_potencia = ((potencia_transportada_MW * resistencia_ca * longitud) / (tension_nominal**2 * cos_phi**2))*100
    print(f"\nPérdidas de potencia en la línea: {perdidas_potencia:.5f} %")
    # En valor absoluto
    perdidas_potencia_valor = (potencia_transportada_MW * resistencia_ca * longitud) / (tension_nominal**2 * cos_phi**2) * potencia_transportada_MW
    print(f"Pérdidas de potencia en la línea en valor absoluto: {perdidas_potencia_valor:.5f} MW")

    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # CORTOCIRCUITO MÁXIMO
    #----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    print("\n13. CORTOCIRCUITO MÁXIMO")
    # Contantes para cálculo de factor K en función del material del conductor
    materiales = ["Cobre", "Aluminio-Acero", "Acero"]
    c_conductores = [390, 910, 480]  # J/kgK
    densidades_conductores = [8900, 2700, 7850]  # kg/m3
    ctes_conductores = [56e6, 34.8e6, 7.25e6]  # 1/ohmio m
    alphas_conductores = [0.0039, 0.004, 0.0045]  # 1/ºC
    def propiedades_material(material):
        if material in materiales:
            index = materiales.index(material)
            return (c_conductores[index], densidades_conductores[index], ctes_conductores[index], alphas_conductores[index])
        return (None, None, None, None)
    c_conductor, densidad_conductor, cte_conductor, alpha_conductor = propiedades_material(material)
    print(f"c_conductor, densidad_conductor, cte_conductor, alpha_conductor", c_conductor, densidad_conductor, cte_conductor, alpha_conductor )
    # Temperatura máxima recomendada en función del material (ºC)
    if material == "Acero":
        temperatura_max_recomendada = 300
    else:
        temperatura_max_recomendada = 200
    # Factor K
    multiplicacion = log ((1+alpha_conductor*(temperatura_max_recomendada-20))/(1+alpha_conductor*(Tc-20)))
    K = (raiz (cte_conductor*c_conductor*densidad_conductor*multiplicacion/alpha_conductor))/1000000
    print(f"Factor K: {K:.4f} Araiz(s)/mm2")
    # Corriente máxima de cortocircuito
    Icc_max = K * seccion / raiz(tiempo_accionamiento_proteccion)
    print(f"Icc max: {Icc_max:.4f} A")

    #------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # EFECTO CORONA
    #------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    print("\n14. EFECTO CORONA ")
    # Presión barométrica
    presion_barometrica = 1/(log(log(76)-altitud_media/18330))*100 
    print(f"\nPresión barométrica h = {presion_barometrica:.3f} cmHg")
    # Factor corrección densidad aire: delta
    def delta (temperatura):
        return 3.92*presion_barometrica/(273+temperatura)
    delta_invierno = delta (temperatura_invierno)
    delta_verano = delta (temperatura_verano)
    print(f"\nFactor corrección densidad aire invierno (δinv) = {delta_invierno:.4f}")
    print(f"Factor corrección densidad aire verano (δver) = {delta_verano:.4f}")


    DMG = 100*(matriz_distancias[0][1]*matriz_distancias[1][2]*matriz_distancias[2][0])**(1/6)
    print(f"DMG: {DMG}")
    DMG = 807
    print(f"DMG: {DMG}")

    # Tensión crítica dieléctrica
    def Uc (mt,delta):
        return raiz(3)*mc*mt*(30/raiz(2))*delta*(diametro/20)*log(DMG/(diametro/20))
    Uc_invierno = Uc(mt_invierno, delta_invierno)
    Uc_verano = Uc(mt_verano,delta_verano)
    print(f"\nTensión critica invierno {Uc_invierno:.4f} kV")
    print(f"Tensión crítica verano {Uc_verano:.4f} kV")

    # Pérdidas por efecto corona
    if Uc_invierno < tension_nominal:
        Perdidas_efecto_corona_invierno = 241/delta_invierno * (frecuencia + 25) * raiz (diametro/20/DMG) * (tension_nominal - Uc_invierno)/raiz(3) * 1e-5
        print(f"\nHay pérdidas por efecto corona en invierno: {Perdidas_efecto_corona_invierno:.4f} kW·km/fase")
    else: print(f"\nNo hay pérdidas por efecto corona en invierno")

    if Uc_verano < tension_nominal:
        Perdidas_efecto_corona_verano = 241/delta_verano * (frecuencia + 25) * raiz (diametro/20/DMG) * (tension_nominal - Uc_verano)/raiz(3) * 1e-5
        print(f"Hay pérdidas por efecto corona en verano: {Perdidas_efecto_corona_verano:.4f} kW·km/fase")
    else: print(f"No hay pérdidas por efecto corona en verano")


    #------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # CAMPO ELÉCTRICO
    #------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    print("\n15. CAMPO ELÉCTRICO ")
    # Matriz coeficientes de potencial
    print("\nMatriz de capacidades (km/uF):")
    for i, fila in enumerate(matriz_capacidades):
        print(f"", end="")
        for C in fila:
            print(f"{C:18.2f}", end="  ")
        print()       
    # Matriz coeficientes de potencial inversa
    matriz_capacidades_inversa = np.linalg.inv(matriz_capacidades) * 1000
    print("\nMatriz de capacidades inversa(nF/km):")
    for i, fila in enumerate(matriz_capacidades_inversa):
        print(f"", end="")
        for C in fila:
            print(f"{C:18.2f}", end="  ")
        print()
    # Vector de potenciales (tensiones nominales/raiz(3) (angulos 0 -120 y 120, tension tierra)) en numeros complejos y kV
    V_phase = tension_nominal / raiz(3)
    V_vector = np.array([V_phase * cmath.rect(1, math.radians(0)),
                        V_phase * cmath.rect(1, math.radians(-120)),
                        V_phase * cmath.rect(1, math.radians(120)),
                        V_phase * cmath.rect(1, math.radians(0)),
                        V_phase * cmath.rect(1, math.radians(-120)),
                        V_phase * cmath.rect(1, math.radians(120)),
                        0])
    print("\nVector de potenciales (kV):")
    for i, V in enumerate(V_vector):
        print(f"{V:.3f}")
    # Vector de cargas (matriz_capacidades_inversa @ V_vector)
    Q_vector = matriz_capacidades_inversa @ V_vector /1000
    print("\nVector de cargas (kV·nF/km):")
    for i, Q in enumerate(Q_vector):
        print(f"{Q:.3f} mC/km")




    #---------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # EJECUCIÓN PROGRAMA
    #---------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    sys.stdout = sys.__stdout__
    texto_resultados = buffer.getvalue()
    html = """
    <html>
    <head>
    <meta charset="utf-8">
    <style>
    body { font-family: Arial; }
    .mayus { font-weight: bold; }
    </style>
    </head>
    <body>
    <pre>
    """
    def es_mayusculas(linea):
        letras = [c for c in linea if c.isalpha()]
        return len(letras)>0 and all(c.isupper() for c in letras)

    for linea in texto_resultados.split("\n"):
        # Detectar líneas completamente en mayúsculas
        if es_mayusculas(linea):
            html += f'<span class="mayus">{linea}</span>\n'
        else:
            html += linea + "\n"
    html += """
    </pre>
    </body>
    </html>
    """
    with open("resultados.html", "w", encoding="utf-8") as f:
        f.write(html)






    from docx import Document

    doc = Document()
    doc.add_heading("Resultados", 1)

    for linea in texto_resultados.split("\n"):
        doc.add_paragraph(linea)

    doc.save("resultados.docx")




#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# SIMPLE CIRCUITO DUPLEX
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
